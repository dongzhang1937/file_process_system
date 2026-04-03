"""
需求分析服务模块
实现从上传文件中提取需求，并进行智能匹配和解答

三阶段处理流水线：
1. 精准数据库匹配（chapters表标题/内容匹配）
2. SQL语法提取与6种数据库验证（Function Calling）
3. 网络搜索 + LLM归纳总结
"""
import os
import re
import json
import tempfile
from datetime import datetime
from docx import Document
from difflib import SequenceMatcher
from config.db_config import fetch_one, fetch_all
from config.logging_config import logger
from .llm_service import get_llm_service
from .web_search import WebSearchService


class RequirementNode:
    """
    需求树节点
    
    用于构建多级编号层级结构的需求树。
    通过判断是否有子节点来识别叶子节点（"最后一层"），
    而非简单按编号深度。
    """
    
    def __init__(self, number='', title='', content='', level=0, parent=None):
        """
        Args:
            number: 编号字符串，如 '1.1.1'
            title: 标题（编号后的文本）
            content: 内容文本
            level: 层级深度（从1开始）
            parent: 父节点引用
        """
        self.number = number
        self.title = title
        self.content = content
        self.level = level
        self.parent = parent
        self.children = []
        self.is_leaf = True  # 默认是叶子节点，添加子节点时更新
    
    def add_child(self, child):
        """添加子节点"""
        child.parent = self
        self.children.append(child)
        self.is_leaf = False  # 有子节点则不是叶子
    
    def get_full_path(self):
        """获取从根到当前节点的完整路径"""
        path = []
        node = self
        while node and node.number:
            path.insert(0, {'number': node.number, 'title': node.title})
            node = node.parent
        return path
    
    def get_full_title(self):
        """获取完整编号标题"""
        if self.number:
            return f"{self.number} {self.title}"
        return self.title
    
    def get_leaf_nodes(self):
        """获取所有叶子节点"""
        leaves = []
        if self.is_leaf and self.number:  # 根节点除外
            leaves.append(self)
        for child in self.children:
            leaves.extend(child.get_leaf_nodes())
        return leaves
    
    def to_dict(self):
        """转为字典"""
        return {
            'number': self.number,
            'title': self.title,
            'content': self.content,
            'level': self.level,
            'is_leaf': self.is_leaf,
            'full_path': self.get_full_path(),
            'children': [c.to_dict() for c in self.children]
        }
    
    def __repr__(self):
        return f"RequirementNode({self.number} {self.title[:30]}, leaf={self.is_leaf}, children={len(self.children)})"


class RequirementAnalyzer:
    """需求分析器"""
    
    # 相似度阈值
    EXACT_MATCH_THRESHOLD = 0.95  # 精确匹配阈值
    FUZZY_MATCH_THRESHOLD = 0.6   # 模糊匹配阈值
    SEMANTIC_MATCH_THRESHOLD = 0.5  # 语义匹配阈值
    
    # 管理员用户名常量
    ADMIN_USERNAME = 'asd'
    
    def __init__(self, llm_config_id=None, username=None):
        """
        初始化需求分析器
        
        Args:
            llm_config_id: LLM配置ID，为None则使用默认配置
            username: 用户名，用于获取用户级别的配置
        """
        self.llm_config_id = llm_config_id
        self.username = username or self.ADMIN_USERNAME
        self.llm_service = None
        self.web_search_service = WebSearchService()
    
    def _get_llm_service(self):
        """延迟加载LLM服务"""
        if self.llm_service is None:
            try:
                self.llm_service = get_llm_service(self.llm_config_id)
            except Exception as e:
                logger.warning(f"LLM服务初始化失败: {e}")
                return None
        return self.llm_service
    
    def parse_requirements_from_file(self, file_path, section_filter=None, use_llm=False):
        """
        从文件中解析需求列表
        
        Args:
            file_path: 文件路径（支持 .docx, .txt, .pdf）
            section_filter: 章节过滤列表，如 ['1.4.1', '1.4.2']，为None则解析所有章节
            use_llm: 是否使用LLM智能识别章节结构（仅对PDF有效）
        
        Returns:
            需求列表 [{'index': 1, 'content': '需求内容', 'title': '需求标题', 'section': {...}}, ...]
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.docx':
            return self._parse_docx_requirements(file_path, section_filter)
        elif ext == '.txt':
            return self._parse_txt_requirements(file_path)
        elif ext == '.pdf':
            return self._parse_pdf_requirements(file_path, section_filter, use_llm=use_llm)
        else:
            raise ValueError(f"不支持的文件格式: {ext}")
    
    def _parse_docx_requirements(self, file_path, section_filter=None):
        """
        解析Word文档中的需求
        
        支持三种格式：
        1. 编号列表格式：1. 要求内容  2. 要求内容
        2. 表格格式：序号 | 技术要求 | 技术规格
        3. 纯段落格式：无编号的段落文本，每个段落视为一条需求
        
        Args:
            file_path: 文档路径
            section_filter: 章节过滤列表，如 ['1.4.1', '1.4.2']，为None则解析所有章节
        """
        requirements = []
        doc = Document(file_path)
        
        current_section = None  # 当前章节
        current_requirement = None
        index = 0
        pending_paragraphs = []  # 收集章节内的纯段落文本
        
        # 辅助函数：检查当前章节是否在过滤列表中
        def is_section_allowed(section):
            if section_filter is None:
                return True
            if section is None:
                return False
            section_num = section.get('number', '')
            # 检查是否完全匹配或是子章节
            for allowed in section_filter:
                if section_num == allowed or section_num.startswith(allowed + '.'):
                    return True
            return False
        
        # 标记是否曾经进入过允许的章节
        has_entered_allowed_section = False
        # 标记是否刚解析完表格（用于检测表格后的隐式章节边界）
        just_finished_table = False
        
        for element in self._iter_block_items(doc):
            if hasattr(element, 'text'):
                # 段落
                text = element.text.strip()
                if not text:
                    continue
                
                # 检测是否是章节标题
                # 支持格式：1.4.1 xxx, 一、xxx, 二、xxx, (一)xxx, 第一章 xxx 等
                # 先清理文本中的特殊空白字符
                clean_text = re.sub(r'[\t\u3000]+', ' ', text)  # 将tab和全角空格替换为普通空格
                clean_text = re.sub(r'\s+', ' ', clean_text).strip()  # 合并多个空白
                
                section_match = re.match(r'^(\d+(?:\.\d+)+)\s*(.+)', clean_text)
                # 中文数字章节标题（如 "一、文档交付要求", "（一）总体要求"）
                # 必须有明确的分隔符：、 ） ) ． . 或空格
                # 避免误匹配"一个"、"一种"等普通文本
                chinese_section_match = re.match(
                    r'^[（(]?([一二三四五六七八九十百]+)[）)、．.][ \t]*(.+)', clean_text
                ) if not section_match else None
                # "第X章/节/部分" 格式
                chapter_match = re.match(
                    r'^第([一二三四五六七八九十百\d]+)[章节部分条款]\s*(.+)', clean_text
                ) if not section_match and not chinese_section_match else None
                
                # 【关键逻辑】表格后的隐式章节边界检测
                # 如果刚解析完表格，遇到一个短段落（不是编号需求也不是章节标题），
                # 很可能是新的章节标题（如"文档交付要求"），将current_section设为None
                implicit_section_break = False
                if just_finished_table and not section_match and not chinese_section_match and not chapter_match:
                    is_req_item = self._is_requirement_item(text)
                    # 短文本（<=20字）且不是需求项 -> 可能是隐式章节标题
                    if len(clean_text) <= 20 and not is_req_item:
                        implicit_section_break = True
                        logger.info(f"[隐式章节边界] 表格后遇到短段落: '{clean_text}'")
                        # 保存之前的需求
                        if current_requirement and is_section_allowed(current_section):
                            requirements.append(current_requirement)
                        current_requirement = None
                        pending_paragraphs = []
                        # 设置为未知章节（不在过滤列表中）
                        current_section = {
                            'number': '',
                            'title': clean_text
                        }
                        just_finished_table = False
                        continue
                
                # 重置表格后状态
                just_finished_table = False
                
                if section_match or chinese_section_match or chapter_match:
                    # 保存之前的需求（如果章节允许）
                    if current_requirement and is_section_allowed(current_section):
                        requirements.append(current_requirement)
                    current_requirement = None
                    
                    # 处理上一章节收集的纯段落（如果没有编号需求且章节允许）
                    if pending_paragraphs and current_section and is_section_allowed(current_section):
                        para_reqs = self._process_pending_paragraphs(pending_paragraphs, current_section, index)
                        for req in para_reqs:
                            index += 1
                            req['index'] = index
                            requirements.append(req)
                    pending_paragraphs = []
                    
                    # 根据不同匹配类型设置章节信息
                    if section_match:
                        # 阿拉伯数字格式: 1.4.1 xxx
                        current_section = {
                            'number': section_match.group(1),
                            'title': section_match.group(2).strip()
                        }
                    elif chinese_section_match:
                        # 中文数字格式: 一、xxx 或 （一）xxx
                        chinese_num = chinese_section_match.group(1)
                        current_section = {
                            'number': chinese_num,  # 保留中文数字
                            'title': chinese_section_match.group(2).strip()
                        }
                    elif chapter_match:
                        # "第X章" 格式
                        chapter_num = chapter_match.group(1)
                        current_section = {
                            'number': chapter_num,
                            'title': chapter_match.group(2).strip()
                        }
                    
                    # 检查是否进入/离开了允许的章节
                    if is_section_allowed(current_section):
                        has_entered_allowed_section = True
                    elif has_entered_allowed_section:
                        # 曾经在允许的章节中，现在离开了 -> 可以提前结束（优化性能）
                        # 但不强制退出，因为可能后面还有其他允许的章节
                        pass
                    
                    continue
                
                # 检测是否是新的需求项（不是章节标题）
                is_new_requirement = self._is_requirement_item(text)
                
                if is_new_requirement:
                    # 有编号需求时，清空pending_paragraphs（它们可能是引导语）
                    pending_paragraphs = []
                    
                    # 保存之前的需求（如果章节允许）
                    if current_requirement and is_section_allowed(current_section):
                        requirements.append(current_requirement)
                    
                    # 只有当前章节允许时才创建新需求
                    if is_section_allowed(current_section):
                        index += 1
                        # 提取需求编号和内容
                        req_index, req_content = self._extract_requirement_content(text)
                        current_requirement = {
                            'index': index,
                            'req_index': req_index,
                            'title': req_content[:50] + '...' if len(req_content) > 50 else req_content,
                            'content': req_content,
                            'raw_text': text,
                            'section': current_section.copy() if current_section else None,
                            'type': 'list'
                        }
                    else:
                        current_requirement = None
                elif current_requirement:
                    # 追加到当前需求的内容（多行需求）
                    current_requirement['content'] += '\n' + text
                else:
                    # 【新增】收集无编号的段落，可能是纯段落格式的需求
                    if current_section and len(text) >= 20 and is_section_allowed(current_section):
                        pending_paragraphs.append(text)
            
            elif hasattr(element, 'rows'):
                # 表格
                # 遇到表格时，清空pending_paragraphs（表格前的内容可能是引导语）
                pending_paragraphs = []
                
                # 解析表格，表格内部会自行处理章节过滤
                # 同时返回检测到的新章节（如果有）
                table_reqs, new_section = self._parse_table_requirements(element, current_section, section_filter)
                for req in table_reqs:
                    index += 1
                    req['index'] = index
                    requirements.append(req)
                
                # 如果表格内检测到新的章节（如"二、文档交付要求"），更新current_section
                if new_section:
                    # 先保存之前的需求
                    if current_requirement and is_section_allowed(current_section):
                        requirements.append(current_requirement)
                    current_requirement = None
                    current_section = new_section
                else:
                    # 保存之前的需求（如果章节允许）
                    if current_requirement and is_section_allowed(current_section):
                        requirements.append(current_requirement)
                    current_requirement = None
                
                # 标记刚解析完表格，用于后续隐式章节边界检测
                just_finished_table = True
        
        # 添加最后一个需求
        if current_requirement and is_section_allowed(current_section):
            requirements.append(current_requirement)
        
        # 处理最后一个章节的pending_paragraphs
        if pending_paragraphs and current_section and is_section_allowed(current_section):
            para_reqs = self._process_pending_paragraphs(pending_paragraphs, current_section, index)
            for req in para_reqs:
                index += 1
                req['index'] = index
                requirements.append(req)
        
        return requirements
    
    def _process_pending_paragraphs(self, paragraphs, section, start_index):
        """
        处理收集的纯段落文本，将其转换为需求项
        
        Args:
            paragraphs: 段落文本列表
            section: 当前章节信息
            start_index: 起始索引
        
        Returns:
            需求列表
        """
        if not paragraphs:
            return []
        
        requirements = []
        
        # 过滤掉引导性语句
        skip_patterns = [
            r'.*需要满足以下要求.*',
            r'.*如下要求.*',
            r'.*具体要求.*',
            r'.*满足如下.*',
        ]
        
        for i, text in enumerate(paragraphs):
            # 检查是否是引导语（跳过）
            is_intro = False
            for pattern in skip_patterns:
                if re.match(pattern, text):
                    is_intro = True
                    break
            
            if is_intro:
                continue
            
            # 跳过太短的段落
            if len(text) < 20:
                continue
            
            requirements.append({
                'index': 0,  # 稍后设置
                'req_index': str(len(requirements) + 1),
                'title': text[:50] + '...' if len(text) > 50 else text,
                'content': text,
                'raw_text': text,
                'section': section.copy() if section else None,
                'type': 'paragraph'
            })
        
        return requirements
    
    def _iter_block_items(self, doc):
        """按文档流顺序遍历段落与表格"""
        from docx.document import Document as DocxDocument
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        
        parent_elm = doc.element.body
        
        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield Table(child, doc)
    
    def _is_requirement_item(self, text):
        """判断是否是需求项（而非章节标题）"""
        # 排除章节标题格式（如 1.4.1 xxx）
        if re.match(r'^\d+\.\d+', text):
            return False
        
        # 常见的需求编号格式
        patterns = [
            r'^(\d+)[\.、\)]\s+',           # 1. 或 1、或 1) 后面有内容
            r'^[（\(](\d+)[）\)]\s*',        # (1) 或 （1）
            r'^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮])\s*',  # 圆圈数字
            r'^[★▲●○◆]\s*',                 # 特殊符号开头
        ]
        
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        
        # 表格行样式（PDF 表格常见）："1 安全可靠测评 产品应当..."
        if re.match(r'^(\d+)\s+[^\d\s].{2,}', text):
            # 再加一层关键词判断，减少误判
            if re.search(r'(应|需|必须|支持|提供|符合|具备|满足|采用)', text):
                return True
        
        return False
    
    def _extract_requirement_content(self, text):
        """从需求文本中提取编号和内容"""
        patterns = [
            r'^(\d+)[\.、\)]\s*(.+)',        # 1. xxx 或 1、xxx 或 1) xxx
            r'^[（\(](\d+)[）\)]\s*(.+)',    # (1) xxx 或 （1）xxx
            r'^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮])\s*(.+)',  # ① xxx
            r'^[★▲●○◆]\s*(.+)',              # ★ xxx
            r'^(\d+)\s+(.+)',                # 1 xxx（PDF 表格常见）
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                groups = match.groups()
                if len(groups) == 2:
                    return groups[0], groups[1]
                elif len(groups) == 1:
                    return '', groups[0]
        
        return '', text
    
    def _parse_table_requirements(self, table, current_section, section_filter=None):
        """
        从表格中解析需求
        
        Args:
            table: 表格对象
            current_section: 当前章节信息
            section_filter: 章节过滤列表
            
        Returns:
            tuple: (requirements列表, 检测到的新章节或None)
        """
        requirements = []
        detected_new_section = None  # 用于返回检测到的新章节
        
        if not table.rows:
            return requirements, detected_new_section
        
        # 获取表头
        headers = []
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
            headers.append(cell_text)
        
        # 识别是否是需求表格
        header_keywords = ['序号', '技术要求', '技术规格', '功能要求', '参数', '指标', '规格', '要求']
        is_requirement_table = any(kw in ''.join(headers) for kw in header_keywords)
        
        if not is_requirement_table:
            return requirements, detected_new_section
        
        # 识别列映射
        header_map = {'index': None, 'requirement': None, 'spec': None, 'section': None}
        
        for i, h in enumerate(headers):
            h_clean = h.strip()
            if h_clean in ['序号', '编号', '项', 'No', 'No.', '#']:
                header_map['index'] = i
            elif any(kw in h_clean for kw in ['技术要求', '功能要求', '要求', '功能', '项目', '名称']):
                header_map['requirement'] = i
            elif any(kw in h_clean for kw in ['技术规格', '规格', '参数', '参数值', '指标']):
                header_map['spec'] = i
            # 检测章节列（如"章节"、"条款"等）
            elif any(kw in h_clean for kw in ['章节', '条款', '节号', '条目']):
                header_map['section'] = i
        
        # 如果没有找到明确的要求列，使用启发式方法
        if header_map['requirement'] is None and len(headers) >= 2:
            header_map['requirement'] = 1 if header_map['index'] == 0 else 0
        if header_map['spec'] is None and len(headers) >= 3:
            for i in range(len(headers)):
                if i != header_map['index'] and i != header_map['requirement']:
                    header_map['spec'] = i
                    break
        
        # 辅助函数：检查章节是否允许
        def is_section_allowed_for_row(row_section):
            if section_filter is None:
                return True
            if row_section is None:
                return False
            section_num = row_section.get('number', '')
            for allowed in section_filter:
                if section_num == allowed or section_num.startswith(allowed + '.'):
                    return True
            return False
        
        # 跟踪表格内的当前章节（用于合并单元格的情况）
        table_current_section = current_section.copy() if current_section else None
        
        # 预扫描表格，检测是否有章节分隔行
        # 如果表格第一行的要求文本包含章节编号，可能整个表格跨越多章节
        has_section_markers = False
        for row in table.rows[1:]:
            cells_text = []
            for cell in row.cells:
                cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                cells_text.append(cell_text)
            all_text = ' '.join(cells_text)
            if re.search(r'^\d+\.\d+(?:\.\d+)*\s', all_text):
                has_section_markers = True
                break
        
        # 解析数据行
        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_data = {}
            all_cells_text = []  # 收集所有单元格文本用于章节检测
            for col_idx, cell in enumerate(row.cells):
                cell_text = '\n'.join(p.text.strip() for p in cell.paragraphs if p.text.strip())
                header = headers[col_idx] if col_idx < len(headers) else f'col_{col_idx}'
                row_data[header] = cell_text
                all_cells_text.append(cell_text)
            
            # 提取字段
            req_index = ''
            if header_map['index'] is not None:
                idx_key = headers[header_map['index']]
                req_index = row_data.get(idx_key, str(row_idx))
            else:
                req_index = str(row_idx)
            
            req_text = ''
            if header_map['requirement'] is not None:
                req_key = headers[header_map['requirement']]
                req_text = row_data.get(req_key, '')
            
            spec_text = ''
            if header_map['spec'] is not None:
                spec_key = headers[header_map['spec']]
                spec_text = row_data.get(spec_key, '')
            
            # 清理特殊字符
            req_index = str(req_index).replace('↵', '').replace('←', '').strip()
            req_text = req_text.replace('↵', '\n').replace('←', '').strip()
            spec_text = spec_text.replace('↵', '\n').replace('←', '').strip()
            
            # 检查行内是否有章节编号（可能在序号列、要求列或专门的章节列）
            row_section = table_current_section
            is_section_header_row = False  # 标记是否是章节标题行
            
            # 尝试从序号列检测章节（如 "1.4.1" 或 "1.4.2-1"）
            section_match = re.match(r'^(\d+(?:\.\d+)+)', req_index)
            if section_match:
                detected_section_num = section_match.group(1)
                # 只有在检测到至少2级章节时才认为是章节编号（如1.4, 1.4.1等）
                if detected_section_num.count('.') >= 1:
                    row_section = {
                        'number': detected_section_num,
                        'title': ''
                    }
                    table_current_section = row_section
            
            # 尝试从要求文本开头检测章节（如 "1.4.1 xxx要求：..."）
            if not section_match and req_text:
                text_section_match = re.match(r'^(\d+(?:\.\d+)+)\s+(.*)$', req_text)
                if text_section_match:
                    detected_section_num = text_section_match.group(1)
                    remaining_text = text_section_match.group(2).strip()
                    # 只有在检测到至少2级章节时才更新
                    if detected_section_num.count('.') >= 1:
                        row_section = {
                            'number': detected_section_num,
                            'title': remaining_text[:50] if remaining_text else ''
                        }
                        table_current_section = row_section
                        # 如果这是纯章节标题行（如"1.4.2 国产中间件技术要求"），跳过不作为需求
                        if len(remaining_text) < 50 and not spec_text:
                            is_section_header_row = True
            
            # 检查整行文本是否是章节标题（合并单元格的情况）
            if not section_match:
                full_row_text = ' '.join(all_cells_text).strip()
                # 清理特殊空白字符
                clean_full_row = re.sub(r'[\t\u3000]+', ' ', full_row_text)
                clean_full_row = re.sub(r'\s+', ' ', clean_full_row).strip()
                
                full_row_match = re.match(r'^(\d+(?:\.\d+)+)\s+(.*)$', clean_full_row)
                if full_row_match:
                    detected_section_num = full_row_match.group(1)
                    remaining_text = full_row_match.group(2).strip()
                    if detected_section_num.count('.') >= 1:
                        row_section = {
                            'number': detected_section_num,
                            'title': remaining_text[:50] if remaining_text else ''
                        }
                        table_current_section = row_section
                        # 如果行内没有具体的技术规格，可能是章节分隔行
                        if len(remaining_text) < 80 and not spec_text:
                            is_section_header_row = True
                
                # 【新增】检查是否是中文数字章节标题（如"二、文档交付要求"）
                if not full_row_match:
                    chinese_section_match = re.match(
                        r'^[（(]?([一二三四五六七八九十百]+)[）)、．.][ \t]*(.+)', clean_full_row
                    )
                    if chinese_section_match:
                        chinese_num = chinese_section_match.group(1)
                        section_title = chinese_section_match.group(2).strip()
                        # 这是一个新的主章节，需要返回给调用者
                        detected_new_section = {
                            'number': chinese_num,
                            'title': section_title
                        }
                        # 【关键】更新表格内的当前章节，后续行将使用新章节
                        table_current_section = detected_new_section
                        row_section = detected_new_section
                        is_section_header_row = True
                        logger.info(f"[表格内检测到中文章节] {chinese_num} {section_title}")
            
            # 跳过章节标题行
            if is_section_header_row:
                continue
            
            # 检查章节过滤
            if section_filter and not is_section_allowed_for_row(row_section):
                continue
            
            if req_text or spec_text:
                # 组合技术要求和技术规格作为完整内容
                full_content = req_text
                if spec_text:
                    full_content += f"\n【技术规格】{spec_text}"
                
                requirements.append({
                    'index': 0,  # 稍后设置
                    'req_index': req_index,
                    'title': req_text[:50] + '...' if len(req_text) > 50 else req_text,
                    'content': full_content,
                    'raw_text': f"{req_text} | {spec_text}",
                    'section': row_section.copy() if row_section else None,
                    'type': 'table',
                    'spec': spec_text
                })
        
        return requirements, detected_new_section
    
    def _parse_txt_requirements(self, file_path):
        """
        解析TXT文件中的需求 - 层级编号树结构解析器
        
        支持多级编号格式（如 1, 1.1, 1.1.1, 1.1.1.1），
        自动删除标题中的空格，构建树结构，
        精确识别叶子节点（无子节点的最后一层）作为待匹配需求项。
        
        Returns:
            list: 需求列表（仅叶子节点），每项包含完整路径信息
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        # 构建需求树
        root, all_nodes = self._build_requirement_tree(lines)
        
        # 提取叶子节点作为待匹配需求
        leaves = root.get_leaf_nodes()
        
        requirements = []
        for i, leaf in enumerate(leaves):
            requirements.append({
                'index': i + 1,
                'req_index': leaf.number,
                'title': leaf.title,
                'content': leaf.content if leaf.content else leaf.title,
                'raw_text': f"{leaf.number} {leaf.title}",
                'section': {
                    'number': leaf.number,
                    'title': leaf.title,
                    'path': leaf.get_full_path()
                },
                'type': 'tree_leaf',
                'level': leaf.level,
                'full_path': leaf.get_full_path()
            })
        
        logger.info(f"[TXT树解析] 共解析 {len(all_nodes)} 个节点，{len(leaves)} 个叶子节点")
        return requirements
    
    def _build_requirement_tree(self, lines):
        """
        从文本行构建需求树
        
        Args:
            lines: 文本行列表
        
        Returns:
            tuple: (root_node, all_nodes_list)
        """
        # 编号行正则：匹配 1, 1.1, 1.1.1 等格式
        # 编号后可跟 点号、空格、制表符
        number_pattern = re.compile(r'^(\d+(?:\.\d+)*)[.\s\t]*(.+)')
        
        root = RequirementNode(number='', title='ROOT', level=0)
        all_nodes = []
        
        # 第一遍：解析所有编号行
        current_content_lines = []  # 累积非编号行的内容
        last_node = None
        
        for line in lines:
            text = line.strip()
            if not text:
                # 空行：如果有当前节点，追加空行到内容
                if last_node and current_content_lines:
                    current_content_lines.append('')
                continue
            
            match = number_pattern.match(text)
            if match:
                # 在创建新节点前，将累积的内容保存到上一个节点
                if last_node and current_content_lines:
                    content = '\n'.join(current_content_lines).strip()
                    if content:
                        last_node.content = (last_node.content + '\n' + content).strip() if last_node.content else content
                    current_content_lines = []
                
                number_str = match.group(1)
                # 删除标题中的多余空格
                title = re.sub(r'\s+', '', match.group(2).strip())
                
                level = number_str.count('.') + 1
                
                node = RequirementNode(
                    number=number_str,
                    title=title,
                    content='',
                    level=level
                )
                all_nodes.append(node)
                last_node = node
            else:
                # 非编号行：作为上一个节点的内容
                current_content_lines.append(text)
        
        # 处理最后一个节点的剩余内容
        if last_node and current_content_lines:
            content = '\n'.join(current_content_lines).strip()
            if content:
                last_node.content = (last_node.content + '\n' + content).strip() if last_node.content else content
        
        # 第二遍：构建树结构（基于编号的父子关系）
        node_map = {}  # number -> node
        for node in all_nodes:
            node_map[node.number] = node
        
        for node in all_nodes:
            parent_number = self._get_parent_number(node.number)
            if parent_number and parent_number in node_map:
                node_map[parent_number].add_child(node)
            else:
                # 没有找到父节点，直接挂到根
                root.add_child(node)
        
        # 日志输出树结构
        self._log_tree_structure(root, all_nodes)
        
        return root, all_nodes
    
    def _get_parent_number(self, number):
        """
        获取父编号
        例如: '1.1.1' -> '1.1', '1.1' -> '1', '1' -> ''
        """
        parts = number.split('.')
        if len(parts) <= 1:
            return ''
        return '.'.join(parts[:-1])
    
    def _log_tree_structure(self, root, all_nodes):
        """打印树结构日志"""
        logger.info(f"[TXT树解析] ========== 需求树结构 ==========")
        logger.info(f"[TXT树解析] 总节点数: {len(all_nodes)}")
        
        def _print_node(node, indent=0):
            prefix = '  ' * indent
            leaf_marker = '🍃' if node.is_leaf else '📁'
            content_preview = node.content[:50].replace('\n', ' ') if node.content else ''
            if node.number:
                logger.info(f"[TXT树解析] {prefix}{leaf_marker} {node.number} {node.title}"
                          f" (内容: {len(node.content)}字{'...' + content_preview if content_preview else ''})")
            for child in node.children:
                _print_node(child, indent + 1)
        
        _print_node(root)
        
        leaves = root.get_leaf_nodes()
        logger.info(f"[TXT树解析] 叶子节点: {len(leaves)} 个")
        for leaf in leaves:
            logger.info(f"[TXT树解析]   -> {leaf.number} {leaf.title}")
        logger.info(f"[TXT树解析] ========== 树结构结束 ==========")
    
    def _parse_pdf_requirements(self, file_path, section_filter=None, use_llm=False):
        """
        解析PDF文档中的需求
        
        使用 pdf_parser 解析 PDF 内容，支持章节过滤
        
        Args:
            file_path: PDF文件路径
            section_filter: 章节过滤列表，支持多种格式：
                - 数字格式: ['1.4.1', '5.2']
                - 中文格式: ['第五章', '技术规格']
                - 混合格式: ['第5章', '2.技术规格']
            use_llm: 是否使用LLM智能识别章节结构（默认False）
        """
        requirements = []
        
        logger.info(f"=" * 60)
        logger.info(f"[_parse_pdf_requirements] 开始解析PDF")
        logger.info(f"  文件: {file_path}")
        logger.info(f"  章节过滤: {section_filter}")
        logger.info(f"  使用LLM: {use_llm}")
        logger.info(f"=" * 60)
        
        try:
            from .pdf_parser import PDFParser
            parser = PDFParser(file_path)
            if use_llm:
                logger.info("[_parse_pdf_requirements] 调用 parse_with_llm()...")
                chapters, full_text = parser.parse_with_llm()
            else:
                logger.info("[_parse_pdf_requirements] 调用 parse() (正则模式)...")
                chapters, full_text = parser.parse()
        except ImportError as e:
            raise ImportError(f"PDF解析库未安装，请安装 pdfplumber: {e}")
        except Exception as e:
            raise ValueError(f"PDF解析失败: {e}")
        
        logger.info(f"=" * 80)
        logger.info(f"[PDF章节树] PDF解析完成，共 {len(chapters)} 个章节")
        logger.info(f"[PDF章节树] 章节过滤条件: {section_filter}")
        logger.info(f"[PDF章节树] ========== 完整章节结构 ==========")
        for i, ch in enumerate(chapters):
            level = ch.get('level', 0)
            title = ch.get('title', '')
            content_len = len(ch.get('content', ''))
            page = ch.get('page_start', ch.get('page', 0))
            parent = ch.get('parent', '')
            indent = "  " * level
            parent_display = parent[:30] if parent else '无'
            logger.info(f"[PDF章节树] {i+1:3d}. {indent}[L{level}] {title[:60]} (内容:{content_len}字, 页:{page}, 父:{parent_display})")
        logger.info(f"[PDF章节树] ========== 章节结构结束 ==========")
        
        # 如果没有章节结构，将整个内容作为需求处理
        if not chapters:
            # 按行分割处理
            lines = full_text.split('\n')
            index = 0
            for line in lines:
                text = line.strip()
                if not text or len(text) < 5:  # 跳过太短的行
                    continue
                
                # 跳过章节标题
                if re.match(r'^\d+\.\d+', text) or re.match(r'^第[一二三四五六七八九十\d]+章', text):
                    continue
                
                index += 1
                req_index, content = self._extract_requirement_content(text)
                requirements.append({
                    'index': index,
                    'req_index': req_index,
                    'title': content[:50] + '...' if len(content) > 50 else content,
                    'content': content,
                    'raw_text': text,
                    'type': 'list'
                })
            return requirements
        
        # 中文数字到阿拉伯数字的映射
        chinese_num_map = {
            '零': 0, '一': 1, '二': 2, '三': 3, '四': 4,
            '五': 5, '六': 6, '七': 7, '八': 8, '九': 9, '十': 10,
            '十一': 11, '十二': 12, '十三': 13, '十四': 14, '十五': 15,
            '十六': 16, '十七': 17, '十八': 18, '十九': 19, '二十': 20
        }
        
        def chinese_to_arabic(chinese_str):
            """将中文数字转换为阿拉伯数字"""
            if chinese_str in chinese_num_map:
                return str(chinese_num_map[chinese_str])
            # 简单处理 "十X" 格式
            if chinese_str.startswith('十') and len(chinese_str) == 2:
                second = chinese_num_map.get(chinese_str[1], 0)
                return str(10 + second)
            if chinese_str.startswith('二十') and len(chinese_str) == 3:
                third = chinese_num_map.get(chinese_str[2], 0)
                return str(20 + third)
            return chinese_str
        
        def normalize_section_number(title):
            """从标题中提取并规范化章节编号"""
            # 1. 匹配 "第X章" 格式
            match = re.match(r'^第([一二三四五六七八九十\d]+)[章节篇部]', title)
            if match:
                num = match.group(1)
                if num.isdigit():
                    return num
                return chinese_to_arabic(num)
            
            # 2. 匹配纯数字格式: "1." "1.2" "1.2.3"
            match = re.match(r'^(\d+(?:\.\d+)*)[\.、\s\:：]', title)
            if match:
                return match.group(1)
            
            # 3. 匹配中文数字格式: "一、" "（一）"
            match = re.match(r'^[（\(]?([一二三四五六七八九十]+)[）\)、\.\s]', title)
            if match:
                return chinese_to_arabic(match.group(1))
            
            return ''
        
        def is_section_allowed(chapter_title, chapter_level, parent_chapter_num=None, parent_title=''):
            """
            检查章节是否被允许
            
            支持的过滤格式：
            - "5" 或 "第五章" 或 "第5章" → 匹配第5章及其所有子章节
            - "5.2" 或 "2.技术规格" → 匹配5.2章节
            - "技术规格" → 匹配标题包含"技术规格"的章节
            - "第五章::2.技术规格" → 精确匹配父章节下的子章节
            
            智能组合识别：
            - 如果过滤条件同时包含章号(如"第5章")和子章节(如"2.技术规格")
            - 则只匹配该章下的指定子章节
            """
            if not section_filter:
                return True
            
            section_num = normalize_section_number(chapter_title)
            
            # 处理 "父章节::子章节" 精确过滤
            has_scoped_filter = any('::' in s for s in section_filter)
            if has_scoped_filter:
                for allowed in section_filter:
                    if '::' not in allowed:
                        continue
                    parent_filter, child_filter = [p.strip() for p in allowed.split('::', 1)]
                    if not parent_filter or not child_filter:
                        continue
                    
                    # 父章节匹配
                    parent_ok = False
                    if parent_title and parent_filter in parent_title:
                        parent_ok = True
                    
                    # 通过编号匹配父章节
                    if not parent_ok and parent_chapter_num:
                        if parent_filter.isdigit():
                            parent_ok = parent_chapter_num == parent_filter or parent_chapter_num.startswith(parent_filter + '.')
                        else:
                            parent_num = normalize_section_number(parent_filter)
                            if parent_num:
                                parent_ok = parent_chapter_num == parent_num or parent_chapter_num.startswith(parent_num + '.')
                    
                    # 子章节匹配
                    child_ok = False
                    # 数字前缀匹配
                    child_num_match = re.match(r'^(\d+(?:\.\d+)*)', child_filter)
                    if child_num_match:
                        target_num = child_num_match.group(1)
                        title_num_match = re.match(r'^(\d+)[\.．、\s\:：]?', chapter_title)
                        if title_num_match and title_num_match.group(1) == target_num:
                            child_ok = True
                        if section_num == target_num or section_num.endswith('.' + target_num) or section_num.startswith(target_num + '.'):
                            child_ok = True
                    
                    # 关键词匹配
                    if not child_ok:
                        if child_filter in chapter_title:
                            child_ok = True
                        else:
                            clean_child = re.sub(r'\s+', '', child_filter)
                            clean_title = re.sub(r'\s+', '', chapter_title)
                            if clean_child and clean_child in clean_title:
                                child_ok = True
                    
                    if parent_ok and child_ok:
                        return True
                # 存在 scoped 过滤但未匹配，直接过滤
                return False
            
            # 智能分析过滤条件：识别是否有"章+子章节"组合
            chapter_nums = []  # 顶级章节号列表，如 ['5']
            sub_sections = []  # 子章节条件，如 ['2', '2.技术规格', '技术规格']
            
            for allowed in section_filter:
                allowed_clean = allowed.strip()
                # 识别 "第X章" 格式
                chapter_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节篇部]?$', allowed_clean)
                if chapter_match:
                    allowed_num = chapter_match.group(1)
                    if allowed_num.isdigit():
                        chapter_nums.append(allowed_num)
                    else:
                        chapter_nums.append(chinese_to_arabic(allowed_num))
                    continue
                
                # 识别单独的数字（可能是章号或子章节号）
                if re.match(r'^\d+$', allowed_clean):
                    # 如果是单独数字，可能是章号也可能是子章节号
                    # 根据上下文判断：如果已有章号，则当作子章节
                    if chapter_nums:
                        sub_sections.append(allowed_clean)
                    else:
                        chapter_nums.append(allowed_clean)
                    continue
                
                # 其他都视为子章节条件
                sub_sections.append(allowed_clean)
            
            logger.info(f"=" * 60)
            logger.info(f"[章节过滤] ===== 开始检查 =====")
            logger.info(f"[章节过滤] 章节标题: '{chapter_title}'")
            logger.info(f"[章节过滤] section_num: '{section_num}'")
            logger.info(f"[章节过滤] parent_chapter_num: '{parent_chapter_num}'")
            logger.info(f"[章节过滤] 原始过滤条件: {section_filter}")
            logger.info(f"[章节过滤] 解析后: chapter_nums={chapter_nums}, sub_sections={sub_sections}")
            
            # 组合条件模式：章号 + 子章节
            if chapter_nums and sub_sections:
                logger.info(f"  [组合匹配] 检查章节: '{chapter_title[:40]}', section_num='{section_num}', parent='{parent_chapter_num}'")
                # 必须满足：在指定章节下 且 匹配子章节条件
                in_target_chapter = False
                
                # 检查是否在目标章节下
                for target_ch in chapter_nums:
                    if parent_chapter_num == target_ch:
                        in_target_chapter = True
                        logger.info(f"    -> 在目标章节 {target_ch} 下")
                        break
                    if section_num == target_ch:
                        # 如果是章节本身，也返回 True（显示章节标题）
                        logger.info(f"    -> 是目标章节本身 {target_ch}")
                        return True
                
                if not in_target_chapter:
                    logger.info(f"    -> 不在任何目标章节下，过滤")
                    return False
                
                # 检查子章节条件
                for sub in sub_sections:
                    logger.info(f"    -> 检查子章节条件: '{sub}' vs 章节标题: '{chapter_title}'")
                    # 数字格式子章节
                    num_match = re.match(r'^(\d+(?:\.\d+)*)', sub)
                    if num_match:
                        target_num = num_match.group(1)
                        logger.info(f"      -> 从过滤条件提取数字: '{target_num}'")
                        # 匹配 "2.技术规格" 或 "2、技术规格" 或 "2 技术规格" 或 "2．技术规格"（全角点）
                        # 检查章节标题是否以这个数字开头
                        title_num_match = re.match(r'^(\d+)[\.．、\s\:：]?', chapter_title)
                        if title_num_match:
                            logger.info(f"      -> 从章节标题提取数字: '{title_num_match.group(1)}'")
                            if title_num_match.group(1) == target_num:
                                logger.info(f"      -> 数字前缀匹配成功!")
                                return True
                        if section_num == target_num or section_num.startswith(target_num + '.'):
                            logger.info(f"      -> section_num 匹配成功")
                            return True
                    
                    # 关键词匹配（更宽松：去掉空格后匹配）
                    keyword_match = re.match(r'^[\d\.、\s]*(.+)$', sub)
                    if keyword_match:
                        keyword = keyword_match.group(1).strip()
                        if keyword:
                            # 标准匹配
                            if keyword in chapter_title:
                                logger.info(f"      -> 关键词匹配成功: '{keyword}'")
                                return True
                            # 去空格后匹配
                            clean_keyword = re.sub(r'\s+', '', keyword)
                            clean_title = re.sub(r'\s+', '', chapter_title)
                            if clean_keyword in clean_title:
                                logger.info(f"      -> 去空格关键词匹配成功: '{clean_keyword}'")
                                return True
                
                logger.info(f"    -> 未匹配任何子章节条件")
                return False
            
            # 非组合模式：原有逻辑
            for allowed in section_filter:
                allowed_clean = allowed.strip()
                
                # 1. 处理 "第X章" 格式的过滤条件
                chapter_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节篇部]?', allowed_clean)
                if chapter_match:
                    allowed_num = chapter_match.group(1)
                    if allowed_num.isdigit():
                        target_num = allowed_num
                    else:
                        target_num = chinese_to_arabic(allowed_num)
                    
                    # 匹配该章及其子章节
                    if section_num == target_num:
                        return True
                    if section_num.startswith(target_num + '.'):
                        return True
                    # 如果父章节匹配，子章节也允许
                    if parent_chapter_num == target_num:
                        return True
                    continue
                
                # 2. 处理数字格式: "5.2" 或 "2"
                num_match = re.match(r'^(\d+(?:\.\d+)*)', allowed_clean)
                if num_match:
                    target_num = num_match.group(1)
                    if section_num == target_num:
                        return True
                    if section_num.startswith(target_num + '.'):
                        return True
                    if section_num.endswith('.' + target_num) or section_num == target_num:
                        return True
                    # 子章节编号匹配（如过滤条件是"2"，匹配"5.2"中的子章节"2.技术规格"）
                    sub_match = re.match(r'^' + re.escape(target_num) + r'[\.、\s]', chapter_title)
                    if sub_match:
                        return True
                    continue
                
                # 3. 处理包含关键词的过滤条件: "技术规格" "2.技术规格"
                # 提取关键词（去掉数字前缀）
                keyword_match = re.match(r'^[\d\.、\s]*(.+)$', allowed_clean)
                if keyword_match:
                    keyword = keyword_match.group(1).strip()
                    if keyword and keyword in chapter_title:
                        return True
            
            return False
        
        # 构建章节层级关系，跟踪父章节
        current_parent_num = None  # 当前的顶级章节编号
        index = 0
        
        logger.info(f"开始遍历章节，共 {len(chapters)} 个章节，过滤条件: {section_filter}")
        
        for chapter in chapters:
            chapter_title = chapter.get('title', '')
            chapter_content = chapter.get('content', '')
            chapter_level = chapter.get('level', 1)
            
            # 提取章节编号
            section_num = normalize_section_number(chapter_title)
            
            # 【关键】从 chapter 中获取 Gemini 返回的 parent 字段
            chapter_parent = chapter.get('parent', '')
            
            # 如果是顶级章节（第X章），更新父章节
            if re.match(r'^第[一二三四五六七八九十\d]+[章节篇部]', chapter_title):
                current_parent_num = section_num
                logger.info(f"[父章节更新] 进入新的顶级章节: {chapter_title[:40]}, parent_num={current_parent_num}")
            # 【新增】如果 level=1 且不是"第X章"格式，也视为顶级章节
            elif chapter_level == 1:
                current_parent_num = section_num or chapter_title[:20]
                logger.info(f"[父章节更新] 进入L1章节: {chapter_title[:40]}, parent_num={current_parent_num}")
            # 【关键修复】如果 Gemini 返回了 parent 字段，尝试从中提取章节号
            elif chapter_parent and not current_parent_num:
                # 从 parent 标题中提取章节号
                parent_num = normalize_section_number(chapter_parent)
                if parent_num:
                    current_parent_num = parent_num
                    logger.info(f"[父章节更新] 从 Gemini parent 字段推断: {chapter_parent[:30]} -> {current_parent_num}")
            
            # 检查章节是否被允许
            is_allowed = is_section_allowed(chapter_title, chapter_level, current_parent_num, chapter_parent)
            logger.info(f"章节过滤检查: '{chapter_title[:40]}' (num={section_num}, level={chapter_level}, parent={current_parent_num}, parent_title={chapter_parent[:20] if chapter_parent else ''}, content_len={len(chapter_content)}) -> {'允许' if is_allowed else '过滤'}")
            
            if not is_allowed:
                continue
            
            logger.info(f"✓ 处理章节: {chapter_title[:50]}... (level={chapter_level}, num={section_num}, content_len={len(chapter_content)})")
            
            # 【调试】打印章节内容的前200字，检查编号是否正确
            logger.info(f"  [内容预览] {chapter_content[:200].replace(chr(10), ' | ')}...")
            
            # 构建章节信息
            current_section = {
                'number': section_num,
                'title': chapter_title,
                'parent_num': current_parent_num
            }
            
            # 【关键修复】如果章节内容为空或太短，检查是否整个章节标题本身就是需求
            if len(chapter_content.strip()) < 10:
                logger.warning(f"[内容过短] 章节 '{chapter_title[:40]}' 内容长度只有 {len(chapter_content)}，检查是否有子章节或标题本身是需求")
                # 如果章节标题看起来像需求项（如"2. 技术规格：..."），将标题作为需求
                if self._is_requirement_item(chapter_title) or len(chapter_title) > 20:
                    index += 1
                    req_index, req_content = self._extract_requirement_content(chapter_title)
                    if not req_content:
                        req_content = chapter_title
                    requirements.append({
                        'index': index,
                        'req_index': req_index or str(index),
                        'title': req_content[:50] + '...' if len(req_content) > 50 else req_content,
                        'content': req_content + ('\n' + chapter_content if chapter_content.strip() else ''),
                        'raw_text': chapter_title,
                        'section': current_section.copy(),
                        'type': 'chapter_title'
                    })
                    logger.info(f"  -> 将章节标题作为需求添加: {req_content[:50]}")
                continue
            
            # 解析章节内容中的需求
            # 【关键】先检查内容中是否包含表格（[表格开始]...[表格结束]）
            table_pattern = r'\[表格开始\](.*?)\[表格结束\]'
            table_matches = list(re.finditer(table_pattern, chapter_content, re.DOTALL))
            
            if table_matches:
                # 内容中包含表格，需要分段处理：表格前的文本 + 表格 + 表格后的文本
                last_end = 0
                for table_match in table_matches:
                    # 处理表格前的普通文本
                    before_text = chapter_content[last_end:table_match.start()].strip()
                    if before_text:
                        index = self._parse_text_line_requirements(
                            before_text, current_section, requirements, index
                        )
                    
                    # 处理表格内容
                    table_text = table_match.group(1).strip()
                    index = self._parse_table_text_requirements(
                        table_text, current_section, requirements, index
                    )
                    
                    last_end = table_match.end()
                
                # 处理最后一个表格后的文本
                after_text = chapter_content[last_end:].strip()
                if after_text:
                    index = self._parse_text_line_requirements(
                        after_text, current_section, requirements, index
                    )
            else:
                # 没有表格标记，但可能内容本身是表格行（| col1 | col2 | 格式）
                lines = chapter_content.split('\n')
                pipe_lines = [l for l in lines if l.strip().startswith('|') and l.strip().endswith('|')]
                
                if len(pipe_lines) >= 2:
                    # 看起来是管道符格式的表格（可能表格标记丢失了）
                    logger.info(f"  [表格检测] 检测到 {len(pipe_lines)} 行管道符格式表格行")
                    index = self._parse_table_text_requirements(
                        '\n'.join(pipe_lines), current_section, requirements, index
                    )
                    # 处理非表格行
                    non_pipe_lines = [l for l in lines if not (l.strip().startswith('|') and l.strip().endswith('|'))]
                    non_pipe_text = '\n'.join(non_pipe_lines).strip()
                    if non_pipe_text:
                        index = self._parse_text_line_requirements(
                            non_pipe_text, current_section, requirements, index
                        )
                else:
                    # 纯文本内容，按原来的逻辑处理
                    index = self._parse_text_line_requirements(
                        chapter_content, current_section, requirements, index
                    )
        
        logger.info(f"=" * 60)
        logger.info(f"PDF解析完成，共提取 {len(requirements)} 条需求")
        if len(requirements) == 0:
            logger.warning(f"[警告] 未提取到任何需求！请检查：")
            logger.warning(f"  1. 过滤条件是否正确: {section_filter}")
            logger.warning(f"  2. 章节内容是否为空（Gemini可能只返回了标题没有内容）")
            logger.warning(f"  3. 内容格式是否符合需求项识别规则（需要编号如 1. 2. ① 等）")
        logger.info(f"=" * 60)
        return requirements
    
    def _parse_text_line_requirements(self, text_content, current_section, requirements, index):
        """
        解析普通文本中的需求项（非表格内容）。
        与原来的逐行解析逻辑相同。
        
        Returns:
            更新后的 index
        """
        lines = text_content.split('\n')
        current_requirement = None
        
        for line in lines:
            text = line.strip()
            if not text:
                continue
            
            # 跳过表格标记行
            if text in ('[表格开始]', '[表格结束]'):
                continue
            
            is_new_requirement = self._is_requirement_item(text)
            
            if is_new_requirement:
                if current_requirement:
                    requirements.append(current_requirement)
                
                index += 1
                req_index, req_content = self._extract_requirement_content(text)
                current_requirement = {
                    'index': index,
                    'req_index': req_index,
                    'title': req_content[:50] + '...' if len(req_content) > 50 else req_content,
                    'content': req_content,
                    'raw_text': text,
                    'section': current_section.copy(),
                    'type': 'list'
                }
            elif current_requirement:
                current_requirement['content'] += '\n' + text
                current_requirement['raw_text'] += '\n' + text
        
        if current_requirement:
            requirements.append(current_requirement)
        
        return index
    
    def _parse_table_text_requirements(self, table_text, current_section, requirements, index):
        """
        解析文本格式的表格内容（来自 [表格开始]...[表格结束] 或管道符格式）。
        
        表格格式示例:
            | 序号 | 重要性 | 指标项名称 | 技术指标要求 |
            | 1 | ★ | 安全可靠测评 | 产品应当符合... |
            | 2 | # | 事务处理机制 | 支持事务隔离... |
            |   |   |              | 支持分布式关系数据库MVCC... |
        
        关键处理：
        1. 识别表头，建立列映射
        2. 合并同一指标的多行内容（序号为空的行属于上一条）
        3. 保留指标项名称 + 技术指标要求的完整信息
        
        Returns:
            更新后的 index
        """
        lines = table_text.strip().split('\n')
        if len(lines) < 2:
            return index
        
        # 解析所有行为列数组
        parsed_rows = []
        for line in lines:
            line = line.strip()
            if not line:
                continue
            # 去掉首尾的 | 符号，按 | 分割
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]
            cols = [c.strip() for c in line.split('|')]
            if cols:
                parsed_rows.append(cols)
        
        if len(parsed_rows) < 2:
            return index
        
        # 第一行作为表头
        headers = parsed_rows[0]
        data_rows = parsed_rows[1:]
        
        logger.info(f"  [表格解析] 表头: {headers}, 数据行: {len(data_rows)} 行")
        
        # 识别列映射
        col_map = {
            'index': None,       # 序号列
            'importance': None,  # 重要性列
            'name': None,        # 指标项名称列
            'spec': None,        # 技术指标要求列
        }
        
        for i, h in enumerate(headers):
            h_clean = h.strip()
            if h_clean in ['序号', '编号', '项', 'No', 'No.', '#']:
                col_map['index'] = i
            elif h_clean in ['重要性', '重要程度', '优先级']:
                col_map['importance'] = i
            elif any(kw in h_clean for kw in ['指标项名称', '指标名称', '项目名称', '功能名称', 
                                               '指标项', '名称', '项目', '功能']):
                col_map['name'] = i
            elif any(kw in h_clean for kw in ['技术指标要求', '技术指标', '指标要求', '技术要求',
                                               '功能要求', '要求', '规格', '参数', '说明', '描述']):
                col_map['spec'] = i
        
        # 如果没找到明确的列映射，用启发式方法
        if col_map['name'] is None and col_map['spec'] is None:
            if len(headers) >= 4:
                # 假设是 序号|重要性|名称|要求 格式
                col_map['index'] = 0
                col_map['importance'] = 1
                col_map['name'] = 2
                col_map['spec'] = 3
            elif len(headers) >= 3:
                col_map['index'] = 0
                col_map['name'] = 1
                col_map['spec'] = 2
            elif len(headers) >= 2:
                col_map['name'] = 0
                col_map['spec'] = 1
        
        logger.info(f"  [表格解析] 列映射: {col_map}")
        
        # 解析数据行，合并同一指标的多行
        # 关键逻辑：如果某行序号列为空，说明它是上一条指标的续行
        current_item = None
        
        for row in data_rows:
            # 补齐列数
            while len(row) < len(headers):
                row.append('')
            
            # 获取各列值
            row_index = row[col_map['index']].strip() if col_map['index'] is not None and col_map['index'] < len(row) else ''
            importance = row[col_map['importance']].strip() if col_map['importance'] is not None and col_map['importance'] < len(row) else ''
            name = row[col_map['name']].strip() if col_map['name'] is not None and col_map['name'] < len(row) else ''
            spec = row[col_map['spec']].strip() if col_map['spec'] is not None and col_map['spec'] < len(row) else ''
            
            # 判断是否是序号行（新的需求项）
            has_index = bool(row_index and re.match(r'^\d+$', row_index))
            
            if has_index:
                # 保存上一条
                if current_item:
                    index += 1
                    self._add_table_requirement(
                        current_item, current_section, requirements, index
                    )
                
                # 开始新条目
                current_item = {
                    'row_index': row_index,
                    'importance': importance,
                    'name': name,
                    'specs': [spec] if spec else [],
                }
            elif current_item:
                # 续行：序号为空，追加到当前条目
                # 名称可能也需要追加（如果名称列有新内容）
                if name and name != current_item['name']:
                    current_item['name'] += ' ' + name
                if spec:
                    current_item['specs'].append(spec)
            else:
                # 没有当前条目也没有序号，可能是独立行
                if spec or name:
                    current_item = {
                        'row_index': '',
                        'importance': importance,
                        'name': name,
                        'specs': [spec] if spec else [],
                    }
        
        # 保存最后一条
        if current_item:
            index += 1
            self._add_table_requirement(
                current_item, current_section, requirements, index
            )
        
        return index
    
    def _add_table_requirement(self, item, current_section, requirements, index):
        """将表格中解析出的一条指标添加到需求列表"""
        name = item.get('name', '')
        importance = item.get('importance', '')
        specs = item.get('specs', [])
        row_index = item.get('row_index', '')
        
        # 构建完整的需求标题（包含指标项名称）
        is_key = importance in ('★', '▲')
        importance_prefix = f"{'★' if is_key else '#'} " if importance else ''
        
        # 标题：用指标项名称
        title = f"{importance_prefix}{name}" if name else f"{importance_prefix}{specs[0][:50]}" if specs else ''
        
        # 内容：指标项名称 + 所有技术指标要求（合并多行）
        content_parts = []
        if name:
            content_parts.append(f"【指标项名称】{name}")
        if importance:
            content_parts.append(f"【重要性】{importance}")
        if specs:
            spec_text = '\n'.join(f"  {i+1}) {s}" for i, s in enumerate(specs)) if len(specs) > 1 else specs[0]
            content_parts.append(f"【技术指标要求】{spec_text}")
        
        content = '\n'.join(content_parts)
        
        # 原始文本：保留管道符格式
        raw_parts = [f"序号: {row_index}"]
        if importance:
            raw_parts.append(f"重要性: {importance}")
        if name:
            raw_parts.append(f"指标项: {name}")
        if specs:
            raw_parts.append(f"要求: {'；'.join(specs)}")
        raw_text = ' | '.join(raw_parts)
        
        requirements.append({
            'index': index,
            'req_index': row_index or str(index),
            'title': title[:80] + '...' if len(title) > 80 else title,
            'content': content,
            'raw_text': raw_text,
            'section': current_section.copy() if current_section else None,
            'type': 'table',
            'is_key': is_key,
            'indicator_name': name,
            'specs': specs,
        })
        
        logger.info(f"  [表格需求] #{row_index} {importance} {name[:30]} -> {len(specs)} 条要求")
    
    def _extract_section_number(self, title):
        """从标题中提取章节编号"""
        # 匹配阿拉伯数字格式: 1.4.1 xxx
        match = re.match(r'^(\d+(?:\.\d+)*)\s*[\.、\s]', title)
        if match:
            return match.group(1)
        
        # 匹配中文数字格式: 一、xxx
        match = re.match(r'^([一二三四五六七八九十]+)[、\s]', title)
        if match:
            return match.group(1)
        
        # 匹配 "第X章" 格式
        match = re.match(r'^第([一二三四五六七八九十\d]+)章', title)
        if match:
            return match.group(1)
        
        return ''
    
    def _extract_title_content(self, text):
        """从文本中提取标题和内容（保留兼容性）"""
        req_index, content = self._extract_requirement_content(text)
        title = content[:50] + '...' if len(content) > 50 else content
        return title, content
    
    def analyze_requirement(self, requirement, user_id, document_ids=None,
                            enable_web_search=True, enable_sql_validation=True,
                            sql_db_types=None):
        """
        分析单个需求并生成答案 - 三阶段流水线
        
        流程：
        1. 精准数据库匹配（chapters表标题/内容匹配）
        2. SQL语法提取 + 6种数据库验证（Function Calling）
        3. 网络搜索 + LLM归纳总结
        
        注：一个需求可能同时包含SQL要求和其他要求，需要分别处理后合并
        
        Args:
            requirement: 需求内容（字符串或字典）
            user_id: 用户ID
            document_ids: 指定搜索的文档ID列表
            enable_web_search: 是否启用网络搜索
            enable_sql_validation: 是否启用SQL验证
            sql_db_types: 指定验证的数据库类型列表
        
        Returns:
            {
                'requirement': '原始需求',
                'answer': '回答内容',
                'match_type': 'exact/sql_validation/web_search/combined/none',
                'source': {...},
                'confidence': 0.95,
                'sql_results': {...},  # SQL验证结果（如有）
                'web_results': {...},  # 搜索结果（如有）
            }
        """
        # 处理输入
        if isinstance(requirement, dict):
            req_content = requirement.get('content', requirement.get('title', ''))
            req_title = requirement.get('title', '')
            full_path = requirement.get('full_path', [])
        else:
            req_content = str(requirement)
            req_title = req_content[:50] if len(req_content) > 50 else req_content
            full_path = []
        
        result = {
            'requirement': req_content,
            'requirement_title': req_title,
            'full_path': full_path,
            'answer': None,
            'match_type': 'none',
            'source': None,
            'confidence': 0,
            'sql_results': None,
            'web_results': None
        }
        
        # ============ 阶段1: 精准数据库匹配 ============
        exact_match = self._exact_match_in_documents(req_title, req_content, user_id, document_ids)
        if exact_match:
            content = exact_match['content']
            result['answer'] = content
            result['match_type'] = 'exact'
            images = self._get_images_from_content(content)
            if not images:
                logger.info(f"[三阶段] _get_images_from_content 未提取到图片，尝试 _get_chapter_images (chapter_id={exact_match.get('chapter_id')})")
                images = self._get_chapter_images(exact_match.get('chapter_id'))
            logger.info(f"[三阶段] 精确匹配图片: 共 {len(images)} 张, ids={[img.get('id') for img in images]}")
            result['source'] = {
                'type': 'document',
                'filename': exact_match.get('filename', ''),
                'chapter_id': exact_match.get('chapter_id'),
                'chapter_title': exact_match.get('chapter_title', ''),
                'path': exact_match.get('path', []),
                'images': images
            }
            result['confidence'] = exact_match.get('similarity', 1.0)
            logger.info(f"[三阶段] 阶段1命中: {req_title[:30]}... -> 精确匹配 (置信度: {result['confidence']:.2f})")
            return result
        
        # ============ 阶段2 & 3: 对未匹配需求进行分类处理 ============
        unmatched_result = self._classify_and_process_unmatched(
            req_content, req_title,
            enable_web_search=enable_web_search,
            enable_sql_validation=enable_sql_validation,
            sql_db_types=sql_db_types
        )
        
        if unmatched_result:
            result.update(unmatched_result)
        else:
            result['answer'] = '抱歉，未能找到相关答案。'
            result['match_type'] = 'none'
            result['confidence'] = 0
        
        return result
    
    def _classify_and_process_unmatched(self, content, title,
                                         enable_web_search=True,
                                         enable_sql_validation=True,
                                         sql_db_types=None):
        """
        对未匹配需求进行分类和处理
        
        一个需求可能同时包含SQL要求和其他要求，需要分别处理：
        - SQL部分 -> 阶段2（SQL提取 + 多数据库验证）
        - 非SQL部分 -> 阶段3（网络搜索 + LLM归纳）
        
        Returns:
            dict: 更新后的结果字段，或 None
        """
        has_sql = False
        sql_result_data = None
        web_result_data = None
        
        # ============ 阶段2: SQL语法提取与验证 ============
        # 不再用正则预筛选，直接交给 LLM 判断是否需要生成 SQL
        logger.info(f"[三阶段] enable_sql_validation={enable_sql_validation}, sql_db_types={sql_db_types}")
        if enable_sql_validation:
            try:
                logger.info(f"[三阶段] 阶段2: 直接进入SQL提取验证（由LLM判断）: {title[:50]}")
                
                # 使用LLM + Function Calling 提取并验证SQL
                sql_result_data = self._process_sql_requirement(content, sql_db_types)
                if sql_result_data:
                    has_sql = True
                    logger.info(f"[三阶段] 阶段2完成: SQL验证结果已获取, "
                               f"answer长度={len(sql_result_data.get('answer', '') or '')}, "
                               f"has_sql_results={bool(sql_result_data.get('sql_results'))}, "
                               f"has_test_cases={bool(sql_result_data.get('test_cases'))}")
                else:
                    logger.warning(f"[三阶段] 阶段2返回None! 需求: {title[:80]}... "
                                  f"这意味着SQL验证流程未产出有效结果，将回退到阶段3网络搜索")
            except Exception as e:
                logger.error(f"[三阶段] 阶段2异常: {e}", exc_info=True)
        else:
            logger.info(f"[三阶段] SQL验证已关闭 (enable_sql_validation=False)，跳过阶段2")
        
        # ============ 阶段3: RAG向量检索 / 网络搜索 + LLM归纳 ============
        if enable_web_search:
            try:
                # 优先使用 RAG 向量检索（替代网络搜索）
                web_result_data = self._process_rag_search(content)
                
                if not web_result_data:
                    # RAG 无结果，回退到网络搜索
                    logger.info(f"[三阶段] RAG 无结果，回退到网络搜索")
                    db_doc_sites = None
                    if has_sql:
                        db_doc_sites = [
                            'dev.mysql.com',
                            'postgresql.org',
                            'cloud.tencent.com/document/product/557',
                            'cloud.tencent.com/document/product/1129'
                        ]
                    web_result_data = self._process_web_search_requirement(content, search_sites=db_doc_sites)
                
                if web_result_data:
                    logger.info(f"[三阶段] 阶段3完成: 搜索归纳结果已获取")
            except Exception as e:
                logger.error(f"[三阶段] 阶段3异常: {e}")
        
        # 收集 process_logs
        process_logs = []
        if sql_result_data and sql_result_data.get('process_log'):
            process_logs.append(sql_result_data['process_log'])
        if web_result_data and web_result_data.get('process_log'):
            process_logs.append(web_result_data['process_log'])
        
        # 合并结果
        # 提取 test_cases、intent、extracted_points、doc_proof_links（自定义Prompt模式B的产物）
        test_cases = sql_result_data.get('test_cases') if sql_result_data else None
        intent = sql_result_data.get('intent') if sql_result_data else None
        extracted_points = sql_result_data.get('extracted_points') if sql_result_data else None
        doc_proof_links = sql_result_data.get('doc_proof_links') if sql_result_data else None
        
        if has_sql and web_result_data:
            # 同时有SQL结果和搜索结果 -> combined
            answer_parts = []
            if sql_result_data.get('answer'):
                answer_parts.append(f"【SQL验证结果】\n{sql_result_data['answer']}")
            if web_result_data.get('answer'):
                answer_parts.append(f"【综合分析】\n{web_result_data['answer']}")
            
            return {
                'answer': '\n\n'.join(answer_parts),
                'match_type': 'combined',
                'confidence': max(
                    sql_result_data.get('confidence', 0),
                    web_result_data.get('confidence', 0)
                ),
                'sql_results': sql_result_data.get('sql_results'),
                'test_cases': test_cases,
                'intent': intent,
                'extracted_points': extracted_points,
                'doc_proof_links': doc_proof_links,
                'web_results': web_result_data.get('sources'),
                'source': {'type': 'combined'},
                'process_logs': process_logs
            }
        elif has_sql:
            return {
                'answer': sql_result_data.get('answer', ''),
                'match_type': 'sql_validation',
                'confidence': sql_result_data.get('confidence', 0.7),
                'sql_results': sql_result_data.get('sql_results'),
                'test_cases': test_cases,
                'intent': intent,
                'extracted_points': extracted_points,
                'doc_proof_links': doc_proof_links,
                'source': {'type': 'sql_validation'},
                'process_logs': process_logs
            }
        elif web_result_data:
            return {
                'answer': web_result_data.get('answer', ''),
                'match_type': 'web_search',
                'confidence': web_result_data.get('confidence', 0.5),
                'web_results': web_result_data.get('sources'),
                'source': {
                    'type': 'web',
                    'search_results': web_result_data.get('sources', [])
                },
                'process_logs': process_logs
            }
        
        # 最后尝试：使用LLM直接回答
        llm = self._get_llm_service()
        if llm:
            try:
                # 获取 Skills 模板
                system_prompt = '你是一个专业的数据库技术顾问，请根据你的知识回答用户的需求。如果不确定，请说明。'
                try:
                    from .mcp_skills_config import SkillsConfigManager
                    template = SkillsConfigManager.get_prompt_for_scene('general')
                    if template:
                        rendered = SkillsConfigManager.render_prompt(template, {
                            'task': content,
                            'role': '数据库技术顾问',
                            'requirement': content  # 兼容: 如果模板用了 {{requirement}}
                        })
                        if rendered:
                            system_prompt = rendered
                except Exception:
                    pass
                
                user_msg = f"请回答以下需求：\n{content}"
                llm_result = llm.chat_completion([
                    {'role': 'system', 'content': system_prompt},
                    {'role': 'user', 'content': user_msg}
                ])
                
                llm_process_log = {
                    'stage': 'llm_direct',
                    'stage_name': 'LLM直接回答',
                    'system_prompt': system_prompt,
                    'user_message': user_msg,
                    'llm_raw_output': llm_result['content'],
                    'error': None
                }
                process_logs.append(llm_process_log)
                
                return {
                    'answer': llm_result['content'],
                    'match_type': 'llm_generated',
                    'confidence': 0.3,
                    'source': {'type': 'llm'},
                    'process_logs': process_logs
                }
            except Exception as e:
                logger.error(f"[三阶段] LLM直接回答失败: {e}")
        
        return None
    
    def _process_sql_requirement(self, content, sql_db_types=None):
        """
        阶段2: SQL测试验证（多数据库 Prompt 分发架构）
        
        流程：
        1. 按数据库大类分组（mysql/pg/oracle）
        2. 对每个大类，先查拦截规则（静态特征词匹配，跳过LLM）
        3. 未命中拦截规则的大类，用对应的专用 Prompt 让 LLM 生成 SQL
        4. MySQL 的 Prompt 优先做意图识别（SQL_TEST / DOC_PROOF / UI_PROOF）
        5. 如果是 SQL_TEST，逐大类执行对应 SQL 测试
        6. 如果模式B未启用（无自定义Prompt），回退到原有模式A（Function Calling）
        
        Args:
            content: 需求文本
            sql_db_types: 指定验证的数据库类型列表
            
        Returns:
            dict: {answer, sql_results, confidence, test_cases, process_log} 或 None
        """
        llm = self._get_llm_service()
        if not llm:
            logger.error("[阶段2] LLM服务不可用 (_get_llm_service返回None)，无法执行SQL验证")
            return None
        
        # 构建数据库类型说明
        db_types_desc = ', '.join(sql_db_types) if sql_db_types else '全部已启用的数据库'
        
        # ========== 1. 按数据库大类分组 ==========
        from .simple_rules_checker import get_db_category, check_simple_rules, check_simple_rules_for_db_types
        
        # 确定要测试的数据库大类及其包含的具体 db_type
        db_categories = {}  # {'mysql': ['mysql_centralized', ...], 'pg': [...]}
        if sql_db_types:
            for dt in sql_db_types:
                cat = get_db_category(dt)
                if cat not in db_categories:
                    db_categories[cat] = []
                db_categories[cat].append(dt)
        else:
            # 未指定时，获取所有已启用的
            from .mcp_skills_config import SQLDBConfigManager
            enabled = SQLDBConfigManager.get_enabled_db_types(owner_username=self.username)
            for dt in enabled:
                cat = get_db_category(dt)
                if cat not in db_categories:
                    db_categories[cat] = []
                db_categories[cat].append(dt)
        
        if not db_categories:
            logger.warning("[阶段2] 没有可用的数据库类型，跳过SQL验证")
            return None
        
        logger.info(f"[阶段2] 数据库大类分组: {dict((k, v) for k, v in db_categories.items())}")
        
        # ========== 2. 拦截规则检查 ==========
        intercepted = check_simple_rules_for_db_types(content, sql_db_types or sum(db_categories.values(), []))
        if intercepted:
            logger.info(f"[阶段2] 拦截规则命中: {list(intercepted.keys())}")
        
        # ========== 3. 尝试加载自定义 Prompt（模式B） ==========
        # 按优先级排序：mysql 优先做意图识别
        category_order = ['mysql', 'pg', 'oracle']
        ordered_categories = sorted(db_categories.keys(), key=lambda c: category_order.index(c) if c in category_order else 99)
        
        scene_type_map = {
            'mysql': 'sql_extraction_mysql',
            'pg': 'sql_extraction_pg',
            'oracle': 'sql_extraction_oracle',
        }
        
        # 检查是否有任何自定义 Prompt 存在（决定走模式A还是模式B）
        use_custom_prompt = False
        custom_prompts = {}  # {category: rendered_prompt}
        
        # 内置默认 Prompt（数据库无记录时的 fallback）
        builtin_prompts = self._get_builtin_sql_prompts()
        
        try:
            from .mcp_skills_config import SkillsConfigManager
            for cat in ordered_categories:
                scene_type = scene_type_map.get(cat)
                if not scene_type:
                    continue
                template = SkillsConfigManager.get_prompt_for_scene(scene_type)
                if template:
                    rendered = SkillsConfigManager.render_prompt(template, {
                        'requirement': content,
                    })
                    if rendered:
                        custom_prompts[cat] = rendered
                        use_custom_prompt = True
                        logger.info(f"[阶段2] {cat} Prompt 已加载 (scene={scene_type}, 长度={len(rendered)})")
                else:
                    # 数据库无记录，使用内置默认 Prompt
                    builtin = builtin_prompts.get(cat)
                    if builtin:
                        rendered = builtin.replace('{{requirement}}', content)
                        custom_prompts[cat] = rendered
                        use_custom_prompt = True
                        logger.info(f"[阶段2] {cat} 使用内置默认 Prompt (长度={len(rendered)})")
        except Exception as e:
            logger.warning(f"[阶段2] 加载自定义 Prompt 异常: {e}")
            # 异常时也用内置默认
            for cat in ordered_categories:
                if cat not in custom_prompts:
                    builtin = builtin_prompts.get(cat)
                    if builtin:
                        custom_prompts[cat] = builtin.replace('{{requirement}}', content)
                        use_custom_prompt = True
        
        if not use_custom_prompt:
            # ========== 模式A回退：无自定义Prompt → Function Calling ==========
            logger.info(f"[阶段2] 未找到任何自定义 SQL Prompt，走模式A（Function Calling）")
            default_system_prompt = (
                '你是一个资深的数据库SQL语法专家和测试工程师。\n'
                '用户会给你一段数据库技术需求，你需要：\n\n'
                '## 第一步：需求分析\n'
                '深入理解需求中涉及的数据库特性或SQL语法要求。\n'
                '例如："支持自动更新时间字段" → 涉及 TIMESTAMP 类型 + ON UPDATE CURRENT_TIMESTAMP 特性\n\n'
                '## 第二步：设计测试SQL\n'
                '根据需求语义，设计一组测试SQL语句来验证该特性。\n'
                '**重要规则**：\n'
                '- 所有测试表名**必须**以 `_mcp_test_` 前缀开头\n'
                '- 测试表会在执行后自动清理\n\n'
                '## 第三步：调用工具验证\n'
                '调用 `execute_sql_test` 工具，传入：\n'
                '- `sql_statements`：测试SQL语句列表\n'
                '- `expected_behavior`：简短描述预期行为\n'
                f'- `db_types`：{db_types_desc}\n\n'
                '## 第四步：分析总结\n'
                '给出各数据库的支持情况汇总表\n'
            )
            return self._process_sql_with_function_calling(llm, default_system_prompt, content, db_types_desc, sql_db_types)
        
        # ========== 模式B：多数据库 Prompt 分发执行 ==========
        return self._process_sql_multi_db(llm, content, db_categories, ordered_categories,
                                          custom_prompts, intercepted, db_types_desc, sql_db_types)
    
    @staticmethod
    def _get_builtin_sql_prompts():
        """
        获取三种数据库的内置默认 Prompt（数据库无记录时的 fallback）
        """
        common_part1 = (
            '【第一部分：意图分类规则】\n'
            '请将需求严格归类为以下三种之一：\n'
            '1. "SQL_TEST" (代码验证类)：涉及具体的 SQL 语法、数据类型、内置函数等。可以通过在终端执行 SQL 语句来证明。\n'
            '2. "DOC_PROOF" (文档证明类)：涉及宏观行业标准（如 SQL2003）、产品架构或整体协议兼容性。需引用官方文档证明。\n'
            '3. "UI_PROOF" (界面操作类)：涉及产品实例购买、版本选择、运维监控台、字符集配置等图形化控制台操作。需界面截图证明。\n\n'
        )
        common_part3 = (
            '【第三部分：输出格式要求（极其重要）】\n'
            '你必须且只能输出一个合法的 JSON 对象，绝对不要使用 markdown 代码块包裹，不要包含任何解释文字。\n\n'
            'JSON 结构说明：\n'
            '{\n'
            '  "intent": "选定的分类标签",\n'
            '  "extracted_points": ["列出你从原文中提取出的所有独立考点，用于防止遗漏"],\n'
            '  "test_cases": [\n'
            '    {\n'
            '      "test_point": "测试点名称",\n'
            '      "setup_sql": "前置环境 SQL",\n'
            '      "verify_sql": "核心验证 SQL",\n'
            '      "expected_behavior": "期望执行成功的描述"\n'
            '    }\n'
            '  ]\n'
            '}\n\n'
            '【实际输入】\n'
            '{{requirement}}'
        )
        return {
            'mysql': (
                '你是一个资深的数据库招投标技术专家和 数据库 测试工程师。\n'
                '你的任务是阅读用户提供的"数据库技术要求"，判断其证明方式（意图分类），并在需要时拆解并生成对应的 SQL 测试脚本。\n\n'
                + common_part1 +
                '【第二部分：SQL 生成规则（仅当分类为 SQL_TEST 时适用）】\n'
                '1. 拆分原则：如果需求包含多个特性（如要求了 5 个函数），必须拆分为 5 个独立的测试点。\n'
                '2. 语句闭环：必须包含建表/准备数据的 setup_sql，以及用于查询验证的 verify_sql。无需准备数据时 setup_sql 留空。\n'
                '3. 语法要求：必须严格兼容 MySQL 8.0。\n'
                '4. 数据类型处理：如果技术要求是支持多种"数据类型"，请生成一个包含所有要求类型的 CREATE TABLE 语句作为 verify_sql。\n\n'
                + common_part3
            ),
            'pg': (
                '你是一个资深的数据库招投标技术专家和 PostgreSQL 测试工程师。\n'
                '你的任务是阅读用户提供的"数据库技术要求"，判断其证明方式（意图分类），并在需要时拆解并生成对应的 PostgreSQL 测试脚本。\n\n'
                + common_part1 +
                '【第二部分：SQL 生成规则（仅当分类为 SQL_TEST 时适用）】\n'
                '1. 拆分原则：如果需求包含多个特性（如要求了 5 个函数），必须拆分为 5 个独立的测试点。\n'
                '2. 语句闭环：必须包含建表/准备数据的 setup_sql，以及用于查询验证的 verify_sql。无需准备数据时 setup_sql 留空。\n'
                '3. 语法要求：必须严格兼容 PostgreSQL 15。无论原文本中使用了哪种数据库的专有名词，必须将其等价转换为 PG 语法'
                '（例如：自增列使用 SERIAL，大文本统一用 TEXT 替代 CLOB/LONGTEXT，空值替换使用 COALESCE，数字截断使用 TRUNC）。'
                '测试标量函数时直接 SELECT 即可，绝对不要加 FROM DUAL。\n'
                '4. 数据类型处理：如果技术要求是支持多种"数据类型"，请生成一个包含所有要求类型的 CREATE TABLE 语句作为 verify_sql。\n\n'
                + common_part3
            ),
            'oracle': (
                '你是一个资深的数据库招投标技术专家和 Oracle 测试工程师。\n'
                '你的任务是阅读用户提供的"数据库技术要求"，判断其证明方式（意图分类），并在需要时拆解并生成对应的 Oracle 测试脚本。\n\n'
                + common_part1 +
                '【第二部分：SQL 生成规则（仅当分类为 SQL_TEST 时适用）】\n'
                '1. 拆分原则：如果需求包含多个特性（如要求了 5 个函数），必须拆分为 5 个独立的测试点。\n'
                '2. 语句闭环：必须包含建表/准备数据的 setup_sql，以及用于查询验证的 verify_sql。无需准备数据时 setup_sql 留空。\n'
                '3. 语法要求：必须严格兼容 Oracle 19c。无论原文本中使用了哪种数据库的专有名词，必须将其等价转换为 Oracle 语法'
                '（例如：字符类型使用 VARCHAR2，数字类型使用 NUMBER，大文本使用 CLOB，获取当前时间使用 SYSDATE，空值替换使用 NVL，数字截断使用 TRUNC）。'
                '任何无实表的 SELECT 查询必须带有 FROM DUAL。\n'
                '4. 数据类型处理：如果技术要求是支持多种"数据类型"，请生成一个包含所有要求类型的 CREATE TABLE 语句作为 verify_sql。\n\n'
                + common_part3
            ),
        }
    
    def _process_sql_multi_db(self, llm, content, db_categories, ordered_categories,
                               custom_prompts, intercepted, db_types_desc, sql_db_types):
        """
        模式B 多数据库分发：按数据库大类分别生成SQL并执行
        
        流程：
        1. 第一个有 Prompt 的大类（默认 mysql）做意图识别
        2. 如果非 SQL_TEST（DOC_PROOF/UI_PROOF），直接返回
        3. 如果是 SQL_TEST，逐大类执行：
           - 先检查拦截规则命中的测试用例
           - 未命中则用 LLM 生成
           - 在对应的数据库子类型上执行
        4. 汇总所有数据库的结果
        """
        process_log = {
            'stage': 'sql_validation',
            'stage_name': 'SQL测试验证（多数据库Prompt）',
            'system_prompt': '',
            'user_message': content,
            'tool_calls': [],
            'llm_raw_output': '',
            'error': None
        }
        
        all_executed_cases = []  # 所有大类的执行结果
        all_sql_results = {}
        all_extracted_points = []
        intent = None
        
        for cat in ordered_categories:
            cat_db_types = db_categories[cat]
            logger.info(f"[阶段2-多DB] 处理大类: {cat}, db_types={cat_db_types}")
            
            # ---- 4a. 检查拦截规则 ----
            if cat in intercepted:
                intercepted_result = intercepted[cat]
                logger.info(f"[阶段2-多DB] {cat} 命中拦截规则，直接使用预设测试用例")
                cat_test_cases = intercepted_result.get('test_cases', [])
                cat_extracted_points = intercepted_result.get('extracted_points', [])
                cat_intent = intercepted_result.get('intent', 'SQL_TEST')
                
                # 如果还没确定 intent，用拦截结果的
                if intent is None:
                    intent = cat_intent
                
                if cat_intent == 'SQL_TEST' and cat_test_cases:
                    executed = self._execute_test_cases(cat_test_cases, cat_db_types)
                    all_executed_cases.extend(executed)
                    all_extracted_points.extend(cat_extracted_points)
                    
                    # 收集 sql_results
                    for tc in executed:
                        for db_type, db_result in tc.get('db_results', {}).items():
                            if db_type not in all_sql_results:
                                all_sql_results[db_type] = db_result
                            else:
                                all_sql_results[db_type]['results'].extend(db_result.get('results', []))
                    continue
            
            # ---- 4b. LLM 生成 ----
            cat_prompt = custom_prompts.get(cat)
            if not cat_prompt:
                # 该大类没有自定义 Prompt，跳过 LLM 生成
                # 但如果有其他大类已经产出了 test_cases，尝试用那些 SQL 直接在此大类上跑（兜底）
                logger.warning(f"[阶段2-多DB] {cat} 没有自定义 Prompt，跳过LLM生成")
                continue
            
            try:
                logger.info(f"[阶段2-多DB] 使用 {cat} Prompt 调用LLM...")
                messages = [
                    {'role': 'system', 'content': cat_prompt},
                    {'role': 'user', 'content': content}
                ]
                result = llm.chat_completion(messages, max_tokens=None)
                
                if isinstance(result, str):
                    raw_output = result
                else:
                    raw_output = result.get('content') or ''
                
                process_log['system_prompt'] = cat_prompt if not process_log['system_prompt'] else process_log['system_prompt']
                process_log['llm_raw_output'] += f"\n\n--- [{cat}] ---\n{raw_output}"
                
                logger.info(f"[阶段2-多DB] {cat} LLM返回: 长度={len(raw_output)}")
                
                parsed = self._parse_llm_json_response(raw_output)
                if not parsed:
                    logger.warning(f"[阶段2-多DB] {cat} JSON解析失败")
                    continue
                
                cat_intent = parsed.get('intent', 'SQL_TEST')
                cat_test_cases = parsed.get('test_cases', [])
                cat_extracted_points = parsed.get('extracted_points', [])
                
                # 第一个成功解析的大类确定 intent
                if intent is None:
                    intent = cat_intent
                
                # 如果是非SQL意图，直接返回（第一个大类的结果就能决定）
                if cat_intent in ('DOC_PROOF', 'UI_PROOF') and intent in ('DOC_PROOF', 'UI_PROOF'):
                    return self._handle_non_sql_intent(
                        cat_intent, cat_extracted_points, content, process_log
                    )
                
                # SQL_TEST: 在对应大类的数据库上执行
                if cat_test_cases:
                    all_extracted_points.extend(cat_extracted_points)
                    executed = self._execute_test_cases(cat_test_cases, cat_db_types)
                    all_executed_cases.extend(executed)
                    
                    for tc in executed:
                        for db_type, db_result in tc.get('db_results', {}).items():
                            if db_type not in all_sql_results:
                                all_sql_results[db_type] = db_result
                            else:
                                all_sql_results[db_type]['results'].extend(db_result.get('results', []))
                
            except Exception as e:
                logger.error(f"[阶段2-多DB] {cat} 处理异常: {e}", exc_info=True)
                process_log['error'] = (process_log.get('error') or '') + f'; {cat}: {e}'
                continue
        
        # ========== 汇总结果 ==========
        if not all_executed_cases and not all_sql_results:
            # 所有大类都没产出有效结果
            if intent in ('DOC_PROOF', 'UI_PROOF') and all_extracted_points:
                return self._handle_non_sql_intent(
                    intent, all_extracted_points, content, process_log
                )
            logger.warning("[阶段2-多DB] 所有大类均未产出有效SQL测试结果")
            return None
        
        # 记录工具调用到 process_log
        for i, tc in enumerate(all_executed_cases):
            process_log['tool_calls'].append({
                'round': i + 1,
                'tool_name': f"execute_test: {tc.get('test_point', f'测试点{i+1}')}",
                'arguments': json.dumps({
                    'test_point': tc.get('test_point', ''),
                    'setup_sql': tc.get('setup_sql', ''),
                    'verify_sql': tc.get('verify_sql', ''),
                    'expected_behavior': tc.get('expected_behavior', '')
                }, ensure_ascii=False),
                'result_preview': json.dumps(
                    {db: {'success': r.get('success'), 'supported': r.get('supported')} 
                     for db, r in tc.get('db_results', {}).items()},
                    ensure_ascii=False
                )[:2000]
            })
        
        # 构建答案
        answer_parts = []
        for i, tc in enumerate(all_executed_cases):
            tp = tc.get('test_point', f'测试点{i+1}')
            answer_parts.append(f"\n### {i+1}. {tp}")
            
            sql_display = ''
            if tc.get('setup_sql'):
                sql_display += f"```sql\n{tc['setup_sql']}\n```\n"
            if tc.get('verify_sql'):
                sql_display += f"```sql\n{tc['verify_sql']}\n```\n"
            answer_parts.append(sql_display)
            
            if tc.get('expected_behavior'):
                answer_parts.append(f"**预期行为**: {tc['expected_behavior']}")
            
            answer_parts.append("\n| 数据库 | 状态 | 结果 |")
            answer_parts.append("|--------|------|------|")
            
            for db_type, db_result in tc.get('db_results', {}).items():
                db_name = db_result.get('db_name', db_type)
                status = '✅' if db_result.get('success') else '❌'
                results_list = db_result.get('results', [])
                last_result = results_list[-1] if results_list else {}
                result_text = last_result.get('result', '') or last_result.get('error', '-')
                if len(str(result_text)) > 100:
                    result_text = str(result_text)[:100] + '...'
                answer_parts.append(f"| {db_name} | {status} | {result_text} |")
        
        answer = '\n'.join(answer_parts)
        
        return {
            'answer': answer,
            'sql_results': all_sql_results if all_sql_results else None,
            'test_cases': all_executed_cases,
            'intent': intent or 'SQL_TEST',
            'extracted_points': all_extracted_points,
            'confidence': 0.8 if all_sql_results else 0.4,
            'process_log': process_log
        }
    
    def _handle_sql_test_intent(self, test_cases, extracted_points, sql_db_types, 
                                 db_types_desc, process_log):
        """
        处理 SQL_TEST 意图：逐条执行测试用例并汇总结果
        
        Args:
            test_cases: LLM 生成的测试用例列表
            extracted_points: 提取的考点列表
            sql_db_types: 指定的数据库类型
            db_types_desc: 数据库类型描述文本
            process_log: 处理日志对象
            
        Returns:
            dict: {answer, sql_results, test_cases, intent, extracted_points, confidence, process_log}
        """
        # 1. 逐条执行测试用例
        executed_cases = self._execute_test_cases(test_cases, sql_db_types)
        
        # 2. 记录到 process_log
        for i, tc in enumerate(executed_cases):
            process_log['tool_calls'].append({
                'round': i + 1,
                'tool_name': f"execute_test: {tc.get('test_point', f'测试点{i+1}')}",
                'arguments': json.dumps({
                    'test_point': tc.get('test_point', ''),
                    'setup_sql': tc.get('setup_sql', ''),
                    'verify_sql': tc.get('verify_sql', ''),
                    'expected_behavior': tc.get('expected_behavior', '')
                }, ensure_ascii=False),
                'result_preview': json.dumps(
                    {db: {'success': r.get('success'), 'supported': r.get('supported')} 
                     for db, r in tc.get('db_results', {}).items()},
                    ensure_ascii=False
                )[:2000]
            })
        
        # 3. 构建答案和汇总 sql_results
        answer_parts = []
        all_sql_results = {}
        
        for i, tc in enumerate(executed_cases):
            tp = tc.get('test_point', f'测试点{i+1}')
            answer_parts.append(f"\n### {i+1}. {tp}")
            
            sql_display = ''
            if tc.get('setup_sql'):
                sql_display += f"```sql\n{tc['setup_sql']}\n```\n"
            if tc.get('verify_sql'):
                sql_display += f"```sql\n{tc['verify_sql']}\n```\n"
            answer_parts.append(sql_display)
            
            if tc.get('expected_behavior'):
                answer_parts.append(f"**预期行为**: {tc['expected_behavior']}")
            
            answer_parts.append("\n| 数据库 | 状态 | 结果 |")
            answer_parts.append("|--------|------|------|")
            
            for db_type, db_result in tc.get('db_results', {}).items():
                db_name = db_result.get('db_name', db_type)
                status = '✅' if db_result.get('success') else '❌'
                results_list = db_result.get('results', [])
                last_result = results_list[-1] if results_list else {}
                result_text = last_result.get('result', '') or last_result.get('error', '-')
                if len(result_text) > 100:
                    result_text = result_text[:100] + '...'
                answer_parts.append(f"| {db_name} | {status} | {result_text} |")
                
                result_key = f"{db_type}"
                if result_key not in all_sql_results:
                    all_sql_results[result_key] = db_result
                else:
                    all_sql_results[result_key]['results'].extend(db_result.get('results', []))
                    if not all_sql_results[result_key]['success'] and db_result.get('success'):
                        all_sql_results[result_key]['success'] = True
                        all_sql_results[result_key]['supported'] = True
        
        answer = '\n'.join(answer_parts)
        
        return {
            'answer': answer,
            'sql_results': all_sql_results if all_sql_results else None,
            'test_cases': executed_cases,
            'intent': 'SQL_TEST',
            'extracted_points': extracted_points,
            'confidence': 0.8 if all_sql_results else 0.4,
            'process_log': process_log
        }
    
    def _handle_non_sql_intent(self, intent, extracted_points, content, process_log):
        """
        处理非 SQL 意图（DOC_PROOF / UI_PROOF）
        
        DOC_PROOF: 去腾讯云私有云"数据库与存储"子页面做定向搜索，
                   找到语义相关的文档链接作为证明材料来源。
        UI_PROOF: 返回界面截图建议。
        
        Args:
            intent: 意图分类 ('DOC_PROOF' 或 'UI_PROOF')
            extracted_points: 提取的考点列表
            content: 原始需求文本
            process_log: 处理日志对象
            
        Returns:
            dict: {answer, sql_results, test_cases, intent, extracted_points,
                   confidence, process_log, doc_proof_links}
        """
        intent_labels = {
            'DOC_PROOF': '📄 文档证明',
            'UI_PROOF': '🖥️ 界面操作截图'
        }
        intent_desc = {
            'DOC_PROOF': '此需求涉及行业标准、产品架构或协议兼容性，需要引用官方文档或产品白皮书作为证明材料。',
            'UI_PROOF': '此需求涉及产品控制台操作（如实例购买、版本选择、运维监控等），需要通过管理控制台界面截图作为证明材料。'
        }
        
        label = intent_labels.get(intent, intent)
        desc = intent_desc.get(intent, '')
        
        answer_parts = [
            f"**证明方式**: {label}\n",
            f"{desc}\n"
        ]
        
        if extracted_points:
            answer_parts.append("**需要证明的考点**:\n")
            for i, pt in enumerate(extracted_points, 1):
                answer_parts.append(f"{i}. {pt}")
            answer_parts.append("")
        
        # DOC_PROOF: 搜索腾讯云私有云文档
        doc_proof_links = []
        if intent == 'DOC_PROOF':
            doc_proof_links = self._search_privatecloud_docs(content, extracted_points, process_log)
            
            if doc_proof_links:
                answer_parts.append("**相关证明材料来源**:")
                for link in doc_proof_links:
                    answer_parts.append(f"- [{link['title']}]({link['url']})")
                    if link.get('snippet'):
                        answer_parts.append(f"  _{link['snippet'][:100]}_")
            else:
                answer_parts.append("**建议的证明材料来源**:")
                answer_parts.append("- 产品官方文档（https://cloud.tencent.com/privatecloud）")
                answer_parts.append("- 产品白皮书 / 技术规格说明书")
                answer_parts.append("- 行业标准认证文件（如 SQL:2003 兼容性声明）")
        elif intent == 'UI_PROOF':
            answer_parts.append("**建议的证明方式**:")
            answer_parts.append("- 管理控制台功能截图")
            answer_parts.append("- 操作步骤录屏或截图序列")
            answer_parts.append("- 产品配置界面截图")
        
        answer = '\n'.join(answer_parts)
        
        logger.info(f"[阶段2-B] 非SQL意图处理完成: intent={intent}, "
                   f"考点={len(extracted_points)}个, "
                   f"文档链接={len(doc_proof_links)}个")
        
        return {
            'answer': answer,
            'sql_results': None,
            'test_cases': None,
            'intent': intent,
            'extracted_points': extracted_points,
            'doc_proof_links': doc_proof_links if doc_proof_links else None,
            'confidence': 0.7,
            'process_log': process_log
        }
    
    def _search_privatecloud_docs(self, content, extracted_points, process_log):
        """
        在腾讯云私有云"数据库与存储"子页面中搜索语义相关的文档链接
        
        搜索范围限定在 cloud.tencent.com/privatecloud 下的数据库相关页面
        
        Args:
            content: 原始需求文本
            extracted_points: 提取的考点列表
            process_log: 处理日志对象
            
        Returns:
            list: [{title, url, snippet}, ...] 搜索到的相关文档链接
        """
        if not self.web_search_service.config:
            logger.info("[DOC_PROOF] 网络搜索未配置，跳过文档搜索")
            return []
        
        try:
            # 构建搜索关键词：从考点中提取核心关键词
            search_keywords = self._extract_search_keywords(content)
            if extracted_points:
                # 把考点也加入搜索词（取前3个考点的关键词）
                for pt in extracted_points[:3]:
                    pt_keywords = self._extract_search_keywords(pt)
                    if pt_keywords:
                        search_keywords += ' ' + pt_keywords
            
            # 限定搜索范围到腾讯云私有云的数据库与存储相关页面
            privatecloud_sites = [
                'cloud.tencent.com/privatecloud'
            ]
            
            site_query = ' OR '.join([f'site:{s}' for s in privatecloud_sites])
            full_query = f'{search_keywords} 数据库 ({site_query})'
            
            logger.info(f"[DOC_PROOF] 搜索腾讯云私有云文档: query={full_query}")
            
            # 记录到 process_log
            if process_log:
                if 'doc_search_query' not in process_log:
                    process_log['doc_search_query'] = full_query
            
            results = self.web_search_service.search(full_query, num_results=5)
            
            if not results:
                logger.info("[DOC_PROOF] 未搜索到相关文档")
                return []
            
            # 过滤：只保留 privatecloud 相关的链接
            filtered = []
            for r in results:
                url = r.get('url', '')
                # 只保留腾讯云私有云或数据库相关的链接
                if ('privatecloud' in url or 
                    'cloud.tencent.com/document' in url or
                    'cloud.tencent.com/product' in url):
                    filtered.append({
                        'title': r.get('title', ''),
                        'url': url,
                        'snippet': r.get('snippet', '')
                    })
            
            # 如果过滤后没有结果，放宽条件取所有结果
            if not filtered:
                filtered = [{
                    'title': r.get('title', ''),
                    'url': r.get('url', ''),
                    'snippet': r.get('snippet', '')
                } for r in results[:3]]
            
            logger.info(f"[DOC_PROOF] 搜索到 {len(filtered)} 个相关文档链接")
            return filtered[:5]  # 最多返回5个
            
        except Exception as e:
            logger.warning(f"[DOC_PROOF] 搜索腾讯云私有云文档失败: {e}")
            return []
    
    def _parse_llm_json_response(self, raw_output):
        """
        从LLM返回的文本中解析JSON响应
        
        支持两种格式：
        - 新格式（推荐）: {"intent": "SQL_TEST/DOC_PROOF/UI_PROOF", "extracted_points": [...], "test_cases": [...]}
        - 旧格式（兼容）: [{test_point, setup_sql, verify_sql, expected_behavior}, ...]
        
        Returns:
            dict: {
                'intent': 'SQL_TEST' | 'DOC_PROOF' | 'UI_PROOF' | 'UNKNOWN',
                'extracted_points': [...],
                'test_cases': [...],   # test_case 列表
            }
            或空 dict {} 表示解析失败
        """
        import re
        
        text = raw_output.strip()
        if not text:
            logger.warning("[阶段2-B] JSON解析: 输入为空")
            return {}
        
        def _clean_json_text(s):
            """清理 JSON 文本中的注释和控制字符"""
            s = re.sub(r'(?m)^\s*//.*$', '', s)
            s = re.sub(r',\s*//[^\n]*', ',', s)
            s = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', s)
            return s
        
        def _try_parse_json(s, label=""):
            """尝试解析 JSON 字符串，返回 (parsed, None) 或 (None, error_str)"""
            try:
                parsed = json.loads(s)
                return parsed, None
            except json.JSONDecodeError as e:
                return None, f"JSONDecodeError: {e}"
            except TypeError as e:
                return None, f"TypeError: {e}"
        
        def _normalize_result(parsed, label=""):
            """将解析结果标准化为 {intent, extracted_points, test_cases} 格式"""
            if isinstance(parsed, dict):
                # 新格式: {intent, extracted_points, test_cases}
                intent = parsed.get('intent', 'UNKNOWN')
                extracted_points = parsed.get('extracted_points', [])
                test_cases = parsed.get('test_cases', [])
                if isinstance(test_cases, list):
                    # 过滤掉非dict元素（LLM偶尔会返回字符串元素）
                    valid_cases = [tc for tc in test_cases if isinstance(tc, dict)]
                    if len(valid_cases) < len(test_cases):
                        logger.warning(f"[阶段2-B] JSON解析({label}): test_cases中有 "
                                      f"{len(test_cases) - len(valid_cases)} 个非dict元素被过滤")
                    logger.info(f"[阶段2-B] JSON解析成功({label}): 意图={intent}, "
                               f"考点={len(extracted_points)}个, 测试用例={len(valid_cases)}个")
                    return {
                        'intent': intent,
                        'extracted_points': extracted_points,
                        'test_cases': valid_cases
                    }
            elif isinstance(parsed, list):
                # 旧格式: [{test_case}, ...] → 转为标准结构
                valid_cases = [tc for tc in parsed if isinstance(tc, dict)]
                if len(valid_cases) > 0:
                    logger.info(f"[阶段2-B] JSON解析成功({label}): 旧格式数组, {len(valid_cases)} 个测试用例")
                    return {
                        'intent': 'SQL_TEST',
                        'extracted_points': [tc.get('test_point', '') for tc in valid_cases],
                        'test_cases': valid_cases
                    }
            return None
        
        def _try_and_normalize(s, label):
            """尝试解析并标准化"""
            parsed, err = _try_parse_json(s, label)
            if parsed is not None:
                result = _normalize_result(parsed, label)
                if result:
                    return result
            # 清理后重试
            cleaned = _clean_json_text(s)
            if cleaned != s:
                parsed2, err2 = _try_parse_json(cleaned, f"{label}-清理后")
                if parsed2 is not None:
                    result2 = _normalize_result(parsed2, f"{label}-清理后")
                    if result2:
                        return result2
            return None
        
        # ========== 尝试1：直接解析 ==========
        result = _try_and_normalize(text, "直接解析")
        if result:
            return result
        
        # ========== 尝试2：提取 ```json ... ``` 代码块 ==========
        json_block_match = re.search(r'```(?:json)?\s*\n([\s\S]*?)\n```', text)
        if json_block_match:
            result = _try_and_normalize(json_block_match.group(1).strip(), "代码块")
            if result:
                return result
        
        # ========== 尝试3：提取 { 到最后的 } (JSON对象) ==========
        first_brace = text.find('{')
        last_brace = text.rfind('}')
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            result = _try_and_normalize(text[first_brace:last_brace + 1], "花括号提取")
            if result:
                return result
        
        # ========== 尝试4：提取 [ 到最后的 ] (JSON数组) ==========
        first_bracket = text.find('[')
        last_bracket = text.rfind(']')
        if first_bracket != -1 and last_bracket != -1 and last_bracket > first_bracket:
            result = _try_and_normalize(text[first_bracket:last_bracket + 1], "方括号提取")
            if result:
                return result
        
        # ========== 尝试5：修复被截断的 JSON ==========
        # 找到第一个 { 或 [ 开始的片段
        start_pos = min(
            first_brace if first_brace != -1 else len(text),
            first_bracket if first_bracket != -1 else len(text)
        )
        if start_pos < len(text):
            json_fragment = _clean_json_text(text[start_pos:])
            
            # 策略A：如果以 { 开头（JSON对象被截断）
            # 典型场景：{"intent":"SQL_TEST","test_cases":[{...},{...},{截断...
            # 策略：找到最后一个完整的 }，然后自动补全缺失的 ] 和 }
            if json_fragment.startswith('{'):
                last_b = json_fragment.rfind('}')
                if last_b > 0:
                    fragment = json_fragment[:last_b + 1]
                    # 直接尝试
                    result = _try_and_normalize(fragment, "对象截断修复")
                    if result:
                        logger.warning(f"[阶段2-B] JSON对象被截断但成功修复")
                        return result
                    
                    # 自动补全缺失的闭合括号
                    # 计算未闭合的 [ 和 { 数量
                    open_braces = fragment.count('{') - fragment.count('}')
                    open_brackets = fragment.count('[') - fragment.count(']')
                    if open_brackets > 0 or open_braces > 0:
                        suffix = ']' * open_brackets + '}' * open_braces
                        fixed = fragment + suffix
                        result = _try_and_normalize(fixed, "对象截断+补全闭合")
                        if result:
                            logger.warning(f"[阶段2-B] JSON被截断，自动补全 '{suffix}' 后修复成功")
                            return result
            
            # 策略B：如果以 [ 开头（JSON数组被截断），找最后一个 } 补 ]
            if json_fragment.startswith('['):
                last_b = json_fragment.rfind('}')
                if last_b > 0:
                    truncated_fix = json_fragment[:last_b + 1] + ']'
                    result = _try_and_normalize(truncated_fix, "数组截断修复")
                    if result:
                        logger.warning(f"[阶段2-B] JSON数组被截断但成功修复")
                        return result
            
            # 策略C：逐个提取顶层 {...} 对象
            objects = []
            brace_depth = 0
            obj_start = -1
            for idx, ch in enumerate(json_fragment):
                if ch == '{':
                    if brace_depth == 0:
                        obj_start = idx
                    brace_depth += 1
                elif ch == '}':
                    brace_depth -= 1
                    if brace_depth == 0 and obj_start >= 0:
                        obj_str = json_fragment[obj_start:idx + 1]
                        parsed, _ = _try_parse_json(obj_str)
                        if isinstance(parsed, dict):
                            # 检查是不是完整的 intent 结构
                            if 'intent' in parsed:
                                result = _normalize_result(parsed, "逐对象-intent结构")
                                if result:
                                    return result
                            elif 'test_point' in parsed:
                                objects.append(parsed)
                        obj_start = -1
            
            if objects:
                logger.warning(f"[阶段2-B] JSON整体解析失败，逐对象提取成功: {len(objects)} 个测试用例")
                return {
                    'intent': 'SQL_TEST',
                    'extracted_points': [o.get('test_point', '') for o in objects],
                    'test_cases': objects
                }
        
        logger.warning(f"[阶段2-B] JSON解析全部失败! "
                      f"原始文本长度={len(text)}, "
                      f"前300字符: {repr(text[:300])}, "
                      f"后100字符: {repr(text[-100:]) if len(text) > 100 else '(同上)'}")
        return {}
    
    def _split_mixed_test_cases(self, test_cases):
        """
        预处理：检测并拆分被 LLM 混合在同一个 test_case 中的多个独立考点。
        
        判断逻辑：将 setup_sql + verify_sql 合并后，如果包含 2 个以上 CREATE TABLE，
        说明 LLM 把多个独立考点塞进了一个 test_case。此时按 CREATE TABLE 边界拆分为
        多个独立的 test_case，每个只包含一组 CREATE TABLE + 后续 DML/SELECT。
        
        仅对 setup_sql 和 verify_sql 合并后的语句进行检测。如果只有 1 个或 0 个
        CREATE TABLE，不做任何处理。
        """
        import re
        
        result = []
        for tc in test_cases:
            if not isinstance(tc, dict):
                result.append(tc)
                continue
            
            setup_sql = tc.get('setup_sql', '').strip()
            verify_sql = tc.get('verify_sql', '').strip()
            
            # 合并所有 SQL 语句
            all_stmts = []
            if setup_sql:
                for s in self._split_sql_statements(setup_sql):
                    if s.strip():
                        all_stmts.append(s.strip())
            if verify_sql:
                for s in self._split_sql_statements(verify_sql):
                    if s.strip():
                        all_stmts.append(s.strip())
            
            # 检测 CREATE TABLE 数量
            create_table_pattern = re.compile(
                r'(?i)^\s*CREATE\s+(?:TEMPORARY\s+)?TABLE\b', re.IGNORECASE
            )
            create_indices = [i for i, s in enumerate(all_stmts) if create_table_pattern.match(s)]
            
            if len(create_indices) <= 1:
                # 0 或 1 个 CREATE TABLE，不需要拆分
                result.append(tc)
                continue
            
            # 需要拆分：按 CREATE TABLE 位置切割为多个独立块
            logger.info(f"[阶段2-B] 检测到 test_case '{tc.get('test_point', '')}' "
                       f"包含 {len(create_indices)} 个 CREATE TABLE，自动拆分")
            
            # 按 CREATE TABLE 边界切割
            blocks = []
            for j, start_idx in enumerate(create_indices):
                end_idx = create_indices[j + 1] if j + 1 < len(create_indices) else len(all_stmts)
                block_stmts = all_stmts[start_idx:end_idx]
                blocks.append(block_stmts)
            
            # 收集 CREATE TABLE 之前的语句（如 DROP TABLE IF EXISTS 等前置清理）
            pre_stmts = all_stmts[:create_indices[0]] if create_indices[0] > 0 else []
            
            # 为每个块生成独立的 test_case
            for j, block_stmts in enumerate(blocks):
                # 从 CREATE TABLE 语句中提取表名作为测试点名
                create_stmt = block_stmts[0]
                table_match = re.search(
                    r'(?i)CREATE\s+(?:TEMPORARY\s+)?TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?[`"\']?(\w+)',
                    create_stmt
                )
                table_name = table_match.group(1) if table_match else f'测试点{j+1}'
                
                # 分离 setup（CREATE/INSERT/DROP）和 verify（SELECT 等查询）
                setup_parts = []
                verify_parts = []
                for s in block_stmts:
                    upper_s = s.strip().upper()
                    if upper_s.startswith('SELECT') or upper_s.startswith('SHOW') or upper_s.startswith('EXPLAIN'):
                        verify_parts.append(s)
                    else:
                        setup_parts.append(s)
                
                # 前置清理语句中与当前表相关的 DROP TABLE 也放进 setup
                if pre_stmts and table_name:
                    for ps in pre_stmts:
                        if table_name.lower() in ps.lower():
                            setup_parts.insert(0, ps)
                
                new_tc = {
                    'test_point': f"{tc.get('test_point', '')} - {table_name}" if tc.get('test_point') else table_name,
                    'setup_sql': ';\n'.join(setup_parts) + ';' if setup_parts else '',
                    'verify_sql': ';\n'.join(verify_parts) + ';' if verify_parts else '',
                    'expected_behavior': tc.get('expected_behavior', ''),
                }
                result.append(new_tc)
                logger.info(f"[阶段2-B] 拆分出子测试: {new_tc['test_point']} "
                           f"(setup={len(setup_parts)}条, verify={len(verify_parts)}条)")
        
        if len(result) != len(test_cases):
            logger.info(f"[阶段2-B] test_cases 拆分: {len(test_cases)} → {len(result)}")
        
        return result
    
    def _execute_test_cases(self, test_cases, sql_db_types=None):
        """
        逐条执行LLM生成的测试用例
        
        对每个 test case，将 setup_sql + verify_sql 合并为语句列表，
        调用 SQLValidator.execute_test_sql_on_all() 在所有目标数据库上执行。
        
        预处理：如果 LLM 把多个考点的 SQL 混在同一个 test_case 中（检测到多个
        CREATE TABLE），会自动拆分为独立的执行单元，每个单元在独立连接中执行。
        
        Args:
            test_cases: LLM 返回的测试用例列表
            sql_db_types: 指定数据库类型列表
            
        Returns:
            list: 带执行结果的测试用例列表
        """
        # 预处理：拆分被混合的 test_cases
        test_cases = self._split_mixed_test_cases(test_cases)
        
        from .sql_validator import get_sql_validator
        validator = get_sql_validator()
        
        executed = []
        for i, tc in enumerate(test_cases):
            # 防御: 如果 tc 不是 dict（LLM返回了格式异常的元素），跳过
            if not isinstance(tc, dict):
                logger.warning(f"[阶段2-B] test_cases[{i}] 不是dict而是 {type(tc).__name__}: {str(tc)[:200]}")
                executed.append({
                    'test_point': str(tc)[:80] if isinstance(tc, str) else f'测试点{i+1}',
                    'setup_sql': '',
                    'verify_sql': '',
                    'expected_behavior': '',
                    'db_results': {},
                    'error': f'LLM返回的test_case格式异常(期望dict,实际{type(tc).__name__})'
                })
                continue
            
            test_point = tc.get('test_point', f'测试点{i+1}')
            setup_sql = tc.get('setup_sql', '').strip()
            verify_sql = tc.get('verify_sql', '').strip()
            expected_behavior = tc.get('expected_behavior', '')
            
            logger.info(f"[阶段2-B] 执行测试用例 {i+1}/{len(test_cases)}: {test_point}")
            
            # 将 setup_sql 和 verify_sql 合并为语句列表
            sql_statements = []
            if setup_sql:
                # setup_sql 可能包含多条语句（用分号分隔）
                for stmt in self._split_sql_statements(setup_sql):
                    if stmt.strip():
                        sql_statements.append(stmt.strip())
            if verify_sql:
                for stmt in self._split_sql_statements(verify_sql):
                    if stmt.strip():
                        sql_statements.append(stmt.strip())
            
            if not sql_statements:
                logger.warning(f"[阶段2-B] 测试用例 {test_point} 没有有效的SQL语句，跳过")
                executed.append({
                    'test_point': test_point,
                    'setup_sql': setup_sql,
                    'verify_sql': verify_sql,
                    'expected_behavior': expected_behavior,
                    'db_results': {},
                    'error': '没有有效的SQL语句'
                })
                continue
            
            # 确保表名以 _mcp_test_ 开头（自动补全前缀）
            sql_statements = self._ensure_test_table_prefix(sql_statements)
            
            logger.info(f"[阶段2-B] SQL语句数: {len(sql_statements)}, 内容: {[s[:80] for s in sql_statements]}")
            
            # 执行
            try:
                db_results = validator.execute_test_sql_on_all(sql_statements, sql_db_types)
                
                # 统计结果
                success_count = sum(1 for r in db_results.values() if r.get('success'))
                total_count = len(db_results)
                logger.info(f"[阶段2-B] {test_point}: {success_count}/{total_count} 数据库执行成功")
                
                executed.append({
                    'test_point': test_point,
                    'setup_sql': setup_sql,
                    'verify_sql': verify_sql,
                    'expected_behavior': expected_behavior,
                    'db_results': db_results,
                    'error': None
                })
            except Exception as e:
                logger.error(f"[阶段2-B] 测试用例 {test_point} 执行异常: {e}")
                executed.append({
                    'test_point': test_point,
                    'setup_sql': setup_sql,
                    'verify_sql': verify_sql,
                    'expected_behavior': expected_behavior,
                    'db_results': {},
                    'error': str(e)
                })
        
        return executed
    
    @staticmethod
    def _split_sql_statements(sql_text):
        """
        智能分割SQL语句（按分号分割，但忽略引号内的分号）
        
        Args:
            sql_text: 可能包含多条SQL的文本
            
        Returns:
            list: SQL语句列表
        """
        statements = []
        current = []
        in_single_quote = False
        in_double_quote = False
        
        for char in sql_text:
            if char == "'" and not in_double_quote:
                in_single_quote = not in_single_quote
            elif char == '"' and not in_single_quote:
                in_double_quote = not in_double_quote
            elif char == ';' and not in_single_quote and not in_double_quote:
                stmt = ''.join(current).strip()
                if stmt:
                    statements.append(stmt)
                current = []
                continue
            current.append(char)
        
        # 最后一条（可能没有分号结尾）
        last = ''.join(current).strip()
        if last:
            statements.append(last)
        
        return statements
    
    @staticmethod
    def _ensure_test_table_prefix(sql_statements):
        """
        确保SQL中的表名以 _mcp_test_ 开头
        从所有DDL/DML语句中提取表名，如果没有 _mcp_test_ 前缀则自动添加，
        并统一替换所有SQL中对该表名的引用
        """
        import re
        
        rename_map = {}  # 原表名 -> 新表名
        
        # SQL 保留关键字（不能被当作表名替换）
        sql_keywords = {
            'CURRENT_TIMESTAMP', 'CURRENT_DATE', 'CURRENT_TIME', 'CURRENT_USER',
            'LOCALTIME', 'LOCALTIMESTAMP', 'SYSDATE', 'SYSTIMESTAMP', 'NOW',
            'NULL', 'TRUE', 'FALSE', 'DEFAULT', 'AUTO_INCREMENT', 'SERIAL',
            'PRIMARY', 'KEY', 'INDEX', 'UNIQUE', 'FOREIGN', 'REFERENCES',
            'CASCADE', 'RESTRICT', 'SET', 'VALUES', 'SELECT', 'FROM', 'WHERE',
            'INTO', 'TABLE', 'CREATE', 'DROP', 'ALTER', 'INSERT', 'UPDATE',
            'DELETE', 'AND', 'OR', 'NOT', 'IN', 'EXISTS', 'BETWEEN', 'LIKE',
            'ORDER', 'GROUP', 'BY', 'HAVING', 'LIMIT', 'OFFSET', 'JOIN',
            'LEFT', 'RIGHT', 'INNER', 'OUTER', 'ON', 'AS', 'IS', 'IF',
            'ELSE', 'THEN', 'WHEN', 'CASE', 'END', 'BEGIN', 'COMMIT',
            'ROLLBACK', 'CONSTRAINT', 'CHECK', 'GRANT', 'REVOKE',
            'INT', 'INTEGER', 'BIGINT', 'SMALLINT', 'TINYINT', 'FLOAT',
            'DOUBLE', 'DECIMAL', 'NUMERIC', 'NUMBER', 'VARCHAR', 'VARCHAR2',
            'CHAR', 'TEXT', 'BLOB', 'CLOB', 'LONGTEXT', 'MEDIUMTEXT',
            'DATE', 'TIME', 'DATETIME', 'TIMESTAMP', 'BOOLEAN', 'BOOL',
            'BYTEA', 'ENUM', 'JSON', 'JSONB', 'XML', 'IDENTITY', 'GENERATED',
            'ALWAYS', 'COLUMN', 'ADD', 'MODIFY', 'CHANGE', 'RENAME',
            'FULLTEXT', 'SPATIAL', 'IGNORE', 'REPLACE', 'DUPLICATE',
        }
        
        # 所有DDL/DML语句的表名提取正则（与 sql_validator 保持一致）
        table_patterns = [
            r'(?i)CREATE\s+(?:TEMPORARY\s+)?TABLE\s+(?:IF\s+NOT\s+EXISTS\s+)?[`"\']?(\w+)',
            r'(?i)ALTER\s+TABLE\s+[`"\']?(\w+)',
            r'(?i)DROP\s+TABLE\s+(?:IF\s+EXISTS\s+)?[`"\']?(\w+)',
            r'(?i)(?:INSERT|REPLACE)\s+(?:OR\s+\w+\s+)?(?:IGNORE\s+)?INTO\s+[`"\']?(\w+)',
            r'(?i)UPDATE\s+[`"\']?(\w+)',
            r'(?i)DELETE\s+FROM\s+[`"\']?(\w+)',
            r'(?i)TRUNCATE\s+(?:TABLE\s+)?[`"\']?(\w+)',
            r'(?i)CREATE\s+(?:FULLTEXT|SPATIAL|UNIQUE|CLUSTERED|NONCLUSTERED)?\s*INDEX\s+\S+\s+ON\s+[`"\']?(\w+)',
            r'(?i)DROP\s+INDEX\s+\S+\s+ON\s+[`"\']?(\w+)',
            r'(?i)LOAD\s+DATA\s+.*?INTO\s+TABLE\s+[`"\']?(\w+)',
        ]
        
        # 第一遍：从所有SQL中收集需要重命名的表名
        for sql in sql_statements:
            for pattern in table_patterns:
                m = re.search(pattern, sql)
                if m:
                    table_name = m.group(1)
                    # 跳过 SQL 关键字和已有前缀的表名
                    if table_name.upper() in sql_keywords:
                        continue
                    if not table_name.lower().startswith('_mcp_test_'):
                        new_name = f'_mcp_test_{table_name}'
                        rename_map[table_name] = new_name
        
        if not rename_map:
            return sql_statements
        
        # 第二遍：替换所有SQL中的表名引用（只替换不在 SQL 关键字集合中的）
        result = []
        for sql in sql_statements:
            new_sql = sql
            for old_name, new_name in rename_map.items():
                # 使用回调函数确保只替换确实是表名引用的位置，不误伤关键字
                def _replace_if_not_keyword(match, _old=old_name, _new=new_name):
                    # 如果匹配到的文本大小写与原表名一致，替换
                    return _new
                new_sql = re.sub(r'\b' + re.escape(old_name) + r'\b', _replace_if_not_keyword, new_sql)
                # 替换反引号包裹的表名
                new_sql = new_sql.replace(f'`{old_name}`', f'`{new_name}`')
            result.append(new_sql)
            if new_sql != sql:
                logger.info(f"[阶段2-B] 表名自动添加前缀: {sql[:80]} -> {new_sql[:80]}")
        
        return result
    
    def _process_sql_with_function_calling(self, llm, system_prompt, content, db_types_desc, sql_db_types):
        """
        模式A：原有的 Function Calling 流程
        LLM 通过 execute_sql_test 工具自动调用数据库验证
        """
        # 使用 Function Calling（execute_sql_test 工具）
        from .function_tools import get_tools_by_names
        sql_test_tools = get_tools_by_names(['execute_sql_test'])
        
        # 如果 execute_sql_test 不可用，回退到 execute_sql
        if not sql_test_tools:
            sql_test_tools = get_tools_by_names(['execute_sql'])
        
        user_message = (
            f'请分析以下数据库技术需求，设计测试SQL验证相关特性，并预判期望输出：\n\n'
            f'**需求内容**：{content}\n\n'
            f'**目标数据库**：{db_types_desc}'
        )
        
        messages = [
            {'role': 'system', 'content': system_prompt},
            {'role': 'user', 'content': user_message}
        ]
        
        # 初始化 process_log 用于前端展示过程
        process_log = {
            'stage': 'sql_validation',
            'stage_name': 'SQL测试验证',
            'system_prompt': system_prompt,
            'user_message': user_message,
            'tool_calls': [],
            'llm_raw_output': '',
            'error': None
        }
        
        try:
            fc_result = llm.chat_with_tools(messages, tools=sql_test_tools, max_rounds=3)
            
            answer = fc_result.get('content', '')
            tool_calls_log = fc_result.get('tool_calls_log', [])
            
            # 记录 LLM 原始输出
            process_log['llm_raw_output'] = answer
            
            # 记录工具调用详情
            for log_entry in tool_calls_log:
                process_log['tool_calls'].append({
                    'round': log_entry.get('round', 0),
                    'tool_name': log_entry.get('tool_name', ''),
                    'arguments': log_entry.get('arguments', ''),
                    'result_preview': log_entry.get('result', '')[:2000]
                })
            
            # 从工具调用日志中提取SQL测试验证结果
            sql_results = {}
            for log_entry in tool_calls_log:
                tool_name = log_entry.get('tool_name', '')
                if tool_name in ('execute_sql_test', 'execute_sql'):
                    try:
                        result_data = json.loads(log_entry.get('result', '{}'))
                        if 'test_results' in result_data:
                            # execute_sql_test 的结果格式
                            sql_results.update(result_data['test_results'])
                        else:
                            # execute_sql 的结果格式（兼容回退）
                            sql_results.update(result_data)
                    except (json.JSONDecodeError, TypeError):
                        pass
            
            # 如果有SQL结果但LLM没有生成总结，用格式化工具生成
            if sql_results and not answer:
                from .sql_validator import SQLValidator
                # 判断是测试模式结果还是只读模式结果
                first_result = next(iter(sql_results.values()), {})
                if 'results' in first_result:
                    # 测试模式结果
                    answer = SQLValidator.format_test_results_for_llm(sql_results)
                else:
                    answer = SQLValidator.format_results_for_llm(sql_results)
            
            if answer or sql_results:
                logger.info(f"[阶段2-A] Function Calling 完成: "
                           f"answer长度={len(answer)}, "
                           f"sql_results数量={len(sql_results)}, "
                           f"tool_calls轮次={len(tool_calls_log)}")
                return {
                    'answer': answer,
                    'sql_results': sql_results if sql_results else None,
                    'confidence': 0.8 if sql_results else 0.4,
                    'process_log': process_log
                }
            else:
                # LLM 既没有返回文本也没有调用工具 — 这不应该发生
                logger.warning(f"[阶段2-A] LLM 返回空结果! answer为空, sql_results为空, "
                              f"tool_calls_log={tool_calls_log}, fc_result keys={list(fc_result.keys())}")
                process_log['error'] = 'LLM返回空结果（无文本、无工具调用）'
                return {
                    'answer': 'SQL验证未能生成有效结果，LLM未返回内容也未调用工具。',
                    'sql_results': None,
                    'confidence': 0.1,
                    'process_log': process_log
                }
            
        except Exception as e:
            process_log['error'] = str(e)
            logger.error(f"[阶段2] SQL测试验证执行失败: {e}", exc_info=True)
            # 即使异常也返回 process_log，确保前端能看到处理过程
            return {
                'answer': f'SQL验证过程中发生错误: {e}',
                'sql_results': None,
                'confidence': 0,
                'process_log': process_log
            }
    
    def _process_rag_search(self, content):
        """
        阶段3替代方案: RAG 向量检索 + LLM 归纳
        
        从 pgvector 中检索相似文档块，用 LLM 归纳总结。
        返回格式与 _process_web_search_requirement 一致。
        
        Returns:
            dict: {answer, sources, confidence, process_log} 或 None
        """
        try:
            from .rag_service import search_similar_chunks, format_search_results_as_context
        except ImportError:
            logger.warning("[阶段3-RAG] rag_service 未安装，跳过")
            return None
        
        process_log = {
            'stage': 'rag_search',
            'stage_name': 'RAG向量检索',
            'search_query': content[:100],
            'search_results_count': 0,
            'system_prompt': '',
            'user_message': '',
            'llm_raw_output': '',
            'error': None
        }
        
        try:
            # 检索
            results = search_similar_chunks(content, top_k=8, threshold=0.45, username=self.username)
            process_log['search_results_count'] = len(results)
            
            if not results:
                logger.info("[阶段3-RAG] 向量检索无结果")
                return None
            
            logger.info(f"[阶段3-RAG] 检索到 {len(results)} 条结果, "
                       f"最高相似度={results[0]['similarity']}")
            
            # 构建 sources（与 web_search 格式兼容）
            sources = []
            for r in results:
                sources.append({
                    'title': f"{r['document']['product_name']} - {r['section']['title']}",
                    'snippet': r['content'][:300],
                    'url': f"文档: {r['document']['filename']} | 章节: {r['section']['full_path']}",
                    'similarity': r['similarity'],
                    'doc_type': r['document']['doc_type'],
                })
            
            # 用 LLM 归纳
            context = format_search_results_as_context(results, max_length=4000)
            
            llm = self._get_llm_service()
            if not llm:
                # 无 LLM，直接返回检索结果
                return {
                    'answer': context,
                    'sources': sources,
                    'confidence': results[0]['similarity'] if results else 0,
                    'process_log': process_log
                }
            
            system_prompt = (
                '你是一个专业的数据库技术顾问。根据提供的参考资料回答用户问题。\n'
                '要求：\n'
                '1. 回答必须基于参考资料，不要编造内容\n'
                '2. 引用信息时注明来源文档和章节\n'
                '3. 如果参考资料不足以完整回答，请明确指出哪些部分无法确认'
            )
            user_msg = f"请根据以下参考资料回答问题。\n\n**问题**：{content}\n\n**参考资料**：\n{context}"
            
            process_log['system_prompt'] = system_prompt
            process_log['user_message'] = user_msg[:500]
            
            llm_result = llm.chat_completion([
                {'role': 'system', 'content': system_prompt},
                {'role': 'user', 'content': user_msg}
            ])
            
            answer = llm_result.get('content', '') if isinstance(llm_result, dict) else str(llm_result)
            process_log['llm_raw_output'] = answer
            
            return {
                'answer': answer,
                'sources': sources,
                'confidence': min(results[0]['similarity'] + 0.1, 1.0) if results else 0.3,
                'process_log': process_log
            }
            
        except Exception as e:
            logger.error(f"[阶段3-RAG] 检索异常: {e}", exc_info=True)
            process_log['error'] = str(e)
            return None
    
    def _process_web_search_requirement(self, content, search_sites=None):
        """
        阶段3: 网络搜索 + LLM归纳总结
        
        Args:
            content: 需求文本
            search_sites: 可选的搜索站点限定列表（如 MySQL/PG/TDSQL 官方文档）
            
        Returns:
            dict: {answer, sources, confidence, process_log} 或 None
        """
        if not self.web_search_service.config:
            return None
        
        # 初始化 process_log
        process_log = {
            'stage': 'web_search',
            'stage_name': '网络搜索归纳',
            'search_query': '',
            'search_results_count': 0,
            'system_prompt': '',
            'user_message': '',
            'llm_raw_output': '',
            'error': None
        }
        
        try:
            search_query = self._extract_search_keywords(content)
            
            # 如果指定了搜索站点（如官方文档），加上 site: 限定
            if search_sites and isinstance(search_sites, list):
                site_query = ' OR '.join([f'site:{s}' for s in search_sites])
                search_query = f'{search_query} ({site_query})'
                logger.info(f"[阶段3] 限定搜索范围: {search_sites}")
            
            process_log['search_query'] = search_query
            
            results = self.web_search_service.search(search_query, num_results=5)
            
            if not results:
                return None
            
            process_log['search_results_count'] = len(results)
            
            llm = self._get_llm_service()
            if llm:
                context = '\n\n'.join([
                    f"[{r['title']}]\n{r['snippet']}\n来源: {r['url']}"
                    for r in results
                ])
                
                # 获取搜索归纳的 system prompt
                system_prompt = '你是一个专业的技术顾问，根据搜索结果回答用户问题。'
                use_custom_search_prompt = False
                try:
                    from .mcp_skills_config import SkillsConfigManager
                    template = SkillsConfigManager.get_prompt_for_scene('web_search_summary')
                    if template:
                        logger.info(f"[阶段3] 搜索归纳 Prompt 已加载: id={template.get('id')}, name={template.get('name')}")
                        rendered = SkillsConfigManager.render_prompt(template, {
                            'requirement': content,
                            'search_results': context
                        })
                        if rendered:
                            use_custom_search_prompt = True
                            # 自定义模板已包含需求和搜索结果，直接作为完整 user message
                            prompt = rendered
                            system_prompt = '你是一个专业的技术顾问，请根据提供的搜索结果给出准确、结构化的回答。'
                            logger.info(f"[阶段3] 使用自定义搜索归纳 Prompt (长度={len(prompt)})")
                        else:
                            logger.warning(f"[阶段3] 搜索归纳 Prompt 渲染失败，使用默认 Prompt")
                    else:
                        logger.info(f"[阶段3] 未找到 web_search_summary 场景的自定义 Prompt，使用默认 Prompt")
                except Exception as e:
                    logger.warning(f"[阶段3] 加载自定义搜索归纳 Prompt 异常: {e}")
                
                if not use_custom_search_prompt:
                    prompt = f"""根据以下搜索结果，回答用户的需求：

需求：{content}

搜索结果：
{context}

请提供简洁准确的回答："""
                
                process_log['system_prompt'] = system_prompt
                process_log['user_message'] = prompt
                
                try:
                    result = llm.chat_completion([
                        {'role': 'system', 'content': system_prompt},
                        {'role': 'user', 'content': prompt}
                    ])
                    
                    process_log['llm_raw_output'] = result['content']
                    
                    return {
                        'answer': result['content'],
                        'sources': results,
                        'confidence': 0.5,
                        'process_log': process_log
                    }
                except Exception as e:
                    process_log['error'] = str(e)
                    logger.error(f"[阶段3] LLM归纳失败: {e}")
            
            # LLM不可用，返回搜索结果摘要
            summary = '\n'.join([f"- {r['title']}: {r['snippet']}" for r in results[:3]])
            return {
                'answer': f"根据网络搜索结果：\n{summary}",
                'sources': results,
                'confidence': 0.3,
                'process_log': process_log
            }
            
        except Exception as e:
            process_log['error'] = str(e)
            logger.error(f"[阶段3] 网络搜索失败: {e}")
            # 即使异常也返回 process_log，确保前端能看到处理过程
            return {
                'answer': f'网络搜索过程中发生错误: {e}',
                'sources': [],
                'confidence': 0,
                'process_log': process_log
            }
    
    def _exact_match_in_documents(self, title, content, user_id, document_ids=None):
        """在文档中进行精确匹配"""
        # 构建查询条件
        conditions = ["dpr.username = %s", "dpr.status = 'completed'"]
        params = [user_id]
        
        if document_ids:
            placeholders = ','.join(['%s'] * len(document_ids))
            conditions.append(f"c.document_id IN ({placeholders})")
            params.extend(document_ids)
        
        # 先尝试标题精确匹配
        sql = f"""
            SELECT c.id as chapter_id, c.document_id, c.title as chapter_title, 
                   c.content, c.level, c.parent_id,
                   dpr.filename
            FROM chapters c
            JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            WHERE {' AND '.join(conditions)}
            AND c.title = %s
            LIMIT 5
        """
        params_with_title = params + [title]
        
        results = fetch_all(sql, params_with_title)
        
        if results:
            best_match = results[0]
            # 获取章节路径
            path = self._get_chapter_path(best_match['chapter_id'])
            best_match['path'] = path
            best_match['similarity'] = 1.0
            return best_match
        
        # 尝试内容模糊匹配
        search_term = self._clean_text(title)[:30]  # 取前30个字符搜索
        
        sql = f"""
            SELECT c.id as chapter_id, c.document_id, c.title as chapter_title, 
                   c.content, c.level, c.parent_id,
                   dpr.filename
            FROM chapters c
            JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            WHERE {' AND '.join(conditions)}
            AND (c.title LIKE %s OR c.content LIKE %s)
            LIMIT 20
        """
        params_with_like = params + [f'%{search_term}%', f'%{search_term}%']
        
        results = fetch_all(sql, params_with_like)
        
        if results:
            # 计算相似度并排序
            best_match = None
            best_similarity = 0
            
            for r in results:
                # 计算标题相似度
                title_sim = SequenceMatcher(None, 
                    self._clean_text(title), 
                    self._clean_text(r['chapter_title'] or '')
                ).ratio()
                
                # 计算内容相似度
                content_sim = SequenceMatcher(None,
                    self._clean_text(content)[:200],
                    self._clean_text(r['content'] or '')[:200]
                ).ratio()
                
                similarity = max(title_sim, content_sim * 0.8)
                
                if similarity > best_similarity and similarity >= self.FUZZY_MATCH_THRESHOLD:
                    best_similarity = similarity
                    best_match = r
                    best_match['similarity'] = similarity
            
            if best_match:
                path = self._get_chapter_path(best_match['chapter_id'])
                best_match['path'] = path
                return best_match
        
        return None
    
    # ==================== 旧方法已移除 ====================
    # _semantic_match_in_documents, _vector_search_match, _llm_semantic_match,
    # _search_from_web, _generate_llm_answer 已被新的三阶段流水线替代
    # ====================================================
    
    def _clean_text(self, text):
        """清理文本用于匹配"""
        if not text:
            return ""
        # 去掉标点符号和特殊字符
        cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)
        return cleaned.lower()
    
    def _extract_search_keywords(self, text):
        """提取搜索关键词"""
        # 简单实现：去掉停用词
        stopwords = {'的', '是', '在', '有', '和', '与', '了', '这', '那', '什么', 
                     '怎么', '如何', '为什么', '需要', '要求', '功能', '支持'}
        words = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]+', text)
        keywords = [w for w in words if w not in stopwords and len(w) > 1]
        return ' '.join(keywords[:10])  # 取前10个关键词
    
    def _get_chapter_path(self, chapter_id):
        """获取章节的完整路径"""
        path = []
        current_id = chapter_id
        
        while current_id:
            sql = """
                SELECT id, parent_id, level, title
                FROM chapters
                WHERE id = %s
            """
            chapter = fetch_one(sql, (current_id,))
            if chapter:
                path.insert(0, {
                    'id': chapter['id'],
                    'level': chapter['level'],
                    'title': chapter['title']
                })
                current_id = chapter['parent_id']
            else:
                break
        
        return path
    
    def _get_images_from_content(self, content):
        """从内容中提取图片ID并获取图片信息
        
        内容中可能包含 {{IMAGE_ID_xxx}} 格式的占位符，
        需要提取这些ID并从数据库获取对应的图片信息
        """
        if not content:
            return []
        
        # 提取所有图片ID
        image_ids = re.findall(r'\{\{IMAGE_ID_(\d+)\}\}', content)
        if not image_ids:
            return []
        
        # 去重
        image_ids = list(set(image_ids))
        logger.info(f"[图片提取] 从content中提取到 {len(image_ids)} 个图片ID: {image_ids}")
        
        try:
            # 将字符串ID转为整数，确保与数据库ID类型一致
            int_ids = [int(x) for x in image_ids]
            placeholders = ','.join(['%s'] * len(int_ids))
            sql = f"""
                SELECT id, image_url, image_path
                FROM document_images
                WHERE id IN ({placeholders})
            """
            images = fetch_all(sql, int_ids)
            logger.info(f"[图片提取] 查询 document_images 返回 {len(images) if images else 0} 条记录")
            
            result = []
            for img in (images or []):
                url = img.get('image_url') or img.get('image_path', '')
                logger.info(f"[图片提取] id={img['id']}, image_url={img.get('image_url')}, image_path={img.get('image_path')}, 最终url={url}")
                result.append({
                    'id': img['id'],
                    'image_url': url,
                    'image_path': img.get('image_path', '')
                })
            
            if not result and image_ids:
                logger.warning(f"[图片提取] 数据库中未找到图片记录! 查询ID: {int_ids}")
            
            return result
        except Exception as e:
            logger.warning(f"从内容获取图片失败: {e}")
            return []
    
    def _get_chapter_images(self, chapter_id):
        """获取章节关联的图片"""
        if not chapter_id:
            return []
        
        try:
            sql = """
                SELECT di.id, di.image_url, di.image_path
                FROM document_images di
                JOIN chapter_images ci ON di.id = ci.image_id
                WHERE ci.chapter_id = %s
                ORDER BY ci.position_in_chapter
            """
            images = fetch_all(sql, (chapter_id,))
            return [
                {
                    'id': img['id'],
                    'image_url': img.get('image_url') or img.get('image_path', ''),
                    'image_path': img.get('image_path', '')
                }
                for img in (images or [])
            ]
        except Exception as e:
            logger.warning(f"获取章节图片失败: {e}")
            return []
    
    def analyze_requirements_batch(self, requirements, user_id, document_ids=None, 
                                   enable_web_search=True, enable_sql_validation=True,
                                   sql_db_types=None, progress_callback=None):
        """
        批量分析需求（三阶段流水线）
        
        Args:
            requirements: 需求列表
            user_id: 用户ID
            document_ids: 文档ID列表
            enable_web_search: 是否启用网络搜索
            enable_sql_validation: 是否启用SQL验证
            sql_db_types: 指定验证的数据库类型列表
            progress_callback: 进度回调函数 callback(current, total, result)
        
        Returns:
            分析结果列表
        """
        results = []
        total = len(requirements)
        
        for i, req in enumerate(requirements):
            try:
                result = self.analyze_requirement(
                    req, user_id, document_ids,
                    enable_web_search=enable_web_search,
                    enable_sql_validation=enable_sql_validation,
                    sql_db_types=sql_db_types
                )
                result['index'] = i + 1
                results.append(result)
                
                if progress_callback:
                    progress_callback(i + 1, total, result)
                    
            except Exception as e:
                logger.error(f"分析需求 {i+1} 失败: {e}")
                results.append({
                    'index': i + 1,
                    'requirement': req.get('content', str(req)) if isinstance(req, dict) else str(req),
                    'answer': f'分析失败: {str(e)}',
                    'match_type': 'error',
                    'confidence': 0
                })
        
        return results
    
    def export_to_word(self, results, title='需求分析报告', requirement_tree=None,
                       format_config=None):
        """
        将分析结果导出为Word文档（招标应答格式）
        
        输出结构：
          文档标题 → 统计信息 → 
          每条需求：
            层级章节标题（Heading）→ "需求：xxx" → "应答：完全满足" → 
            "证明材料如下：" → 证明材料表格/内容
          元数据（路径、匹配类型、意图分类、考点等）输出为 Word 批注
        
        Args:
            results: 分析结果列表
            title: 文档标题
            requirement_tree: 需求树的字典表示（可选，用于层级输出）
            format_config: 格式配置（预留扩展接口）
        
        Returns:
            (filepath, filename) 临时文件路径和文件名
        """
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn, nsdecls
        from docx.oxml import parse_xml
        
        doc = Document()
        
        # ========== 文档标题 ==========
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # ========== 统计信息 ==========
        total = len(results)
        exact_count = sum(1 for r in results if r.get('match_type') == 'exact')
        sql_count = sum(1 for r in results if r.get('match_type') in ('sql_validation', 'combined'))
        web_count = sum(1 for r in results if r.get('match_type') == 'web_search')
        llm_count = sum(1 for r in results if r.get('match_type') == 'llm_generated')
        combined_count = sum(1 for r in results if r.get('match_type') == 'combined')
        
        stats_para = doc.add_paragraph()
        stats_para.add_run(f"总计 {total} 条需求：").bold = True
        stats_para.add_run(f"\n• 精确匹配: {exact_count} 条")
        stats_para.add_run(f"\n• SQL验证: {sql_count} 条")
        stats_para.add_run(f"\n• 网络搜索: {web_count} 条")
        stats_para.add_run(f"\n• LLM生成: {llm_count} 条")
        if combined_count:
            stats_para.add_run(f"\n• 综合(SQL+搜索): {combined_count} 条")
        
        doc.add_paragraph()
        
        # ========== 跟踪已输出的章节标题，避免重复 ==========
        # key: 编号字符串（如 "2", "2.3"），value: True
        emitted_headings = set()
        
        # ========== 遍历每条需求 ==========
        for result in results:
            index = result.get('index', 0)
            requirement = result.get('requirement', '')
            answer = result.get('answer', '无答案')
            match_type = result.get('match_type', 'none')
            confidence = result.get('confidence', 0)
            source = result.get('source', {})
            full_path = result.get('full_path', [])
            sql_results = result.get('sql_results')
            test_cases = result.get('test_cases')
            intent = result.get('intent')
            extracted_points = result.get('extracted_points', [])
            
            req_title = result.get('requirement_title', f'需求 {index}')
            
            # ========== 1. 层级章节标题 ==========
            # full_path 示例: [{'number': '2', 'title': '技术要求'}, {'number': '2.3', 'title': 'MySQL特性'}, ...]
            # 输出父级章节标题（仅首次出现时），最后一级作为本需求的子标题
            title_heading = None  # 保存本需求标题段落引用，用于后续添加失败批注
            if full_path and len(full_path) > 1:
                # 输出除最后一级外的所有父级标题（若尚未输出过）
                for i, path_item in enumerate(full_path[:-1]):
                    path_number = path_item.get('number', '')
                    path_title_text = path_item.get('title', '')
                    if path_number and path_number not in emitted_headings:
                        heading_level = min(i + 1, 4)  # Heading 1~4
                        doc.add_heading(f"{path_number} {path_title_text}", level=heading_level)
                        emitted_headings.add(path_number)
                
                # 最后一级作为本需求标题
                last_item = full_path[-1]
                last_number = last_item.get('number', '')
                heading_level = min(len(full_path), 4)
                title_heading = doc.add_heading(f"{last_number} {req_title}", level=heading_level)
            elif full_path and len(full_path) == 1:
                # 只有一级路径
                last_item = full_path[0]
                last_number = last_item.get('number', '')
                title_heading = doc.add_heading(f"{last_number} {req_title}", level=1)
            else:
                # 无路径信息，直接用标题
                title_heading = doc.add_heading(req_title, level=1)
            
            # ========== 2. 需求内容 ==========
            # 判断需求内容是否与标题重复（树解析时 content 常等于 title）
            # 重复情况：
            #   1) content == title
            #   2) content 以编号+title开头
            #   3) content 去掉编号后 == title
            #   4) content 包含 title（如 "2.5 数据类型\n2.5.1 支持..." 包含 "支持..."）
            #   5) title 包含 content 的核心内容
            req_stripped = requirement.strip()
            title_stripped = req_title.strip()
            
            # 去除 requirement 中可能包含的编号前缀（如 "2.3.1 xxx" → "xxx"）
            req_without_number = re.sub(r'^\d+(\.\d+)*\.?\s*', '', req_stripped)
            title_without_number = re.sub(r'^\d+(\.\d+)*\.?\s*', '', title_stripped)
            
            req_content_same_as_title = (
                req_stripped == title_stripped
                or req_without_number == title_stripped
                or title_without_number == req_without_number
                or (title_stripped and req_stripped.endswith(title_stripped)
                    and len(req_stripped) - len(title_stripped) < 20)
                # 需求内容包含标题（标题是需求文本的子串）
                or (title_without_number and len(title_without_number) > 5
                    and title_without_number in req_stripped)
                # 标题包含需求内容
                or (req_without_number and len(req_without_number) > 5
                    and req_without_number in title_stripped)
            )
            
            if req_content_same_as_title:
                # 需求内容与标题相同，不重复输出
                # 仍需创建段落作为批注的锚点
                req_para = doc.add_paragraph()
                req_anchor_run = req_para.add_run("")  # 空 run 作为批注锚点
            else:
                req_para = doc.add_paragraph()
                req_label_run = req_para.add_run("需求：")
                req_label_run.bold = True
                req_label_run.font.size = Pt(11)
                req_content_run = req_para.add_run(requirement)
                req_content_run.font.size = Pt(11)
            
            # ========== 3. 应答声明 ==========
            answer_para = doc.add_paragraph()
            answer_label_run = answer_para.add_run("应答：")
            answer_label_run.bold = True
            answer_label_run.font.size = Pt(11)
            # 统一应答为"完全满足"，不足之处在证明材料中标注"（请补充截图）"
            answer_text = "完全满足"
            answer_text_run = answer_para.add_run(answer_text)
            answer_text_run.font.size = Pt(11)
            answer_text_run.bold = True
            answer_text_run.font.color.rgb = RGBColor(0, 128, 0)
            
            # ========== 4. 批注：元数据信息 ==========
            # 将路径、匹配类型、意图分类、考点等以批注形式附加到"需求"段落
            comment_lines = []
            
            # 路径信息
            if full_path:
                path_str = ' > '.join([f"{p.get('number', '')} {p.get('title', '')}" for p in full_path])
                comment_lines.append(f"【路径】{path_str}")
            
            # 匹配类型
            match_type_text = {
                'exact': '精确匹配',
                'sql_validation': 'SQL验证',
                'web_search': '网络搜索',
                'combined': '综合(SQL+搜索)',
                'llm_generated': 'LLM生成',
                'none': '未匹配',
                'error': '分析失败'
            }.get(match_type, '未知')
            comment_lines.append(f"【匹配类型】{match_type_text} (置信度: {confidence:.0%})")
            
            # 意图分类
            if intent:
                intent_labels = {
                    'SQL_TEST': '🔍 SQL测试验证',
                    'DOC_PROOF': '📄 文档证明',
                    'UI_PROOF': '🖥️ 界面操作截图'
                }
                intent_label = intent_labels.get(intent, intent)
                comment_lines.append(f"【意图分类】{intent_label}")
            
            # 考点列表
            if extracted_points and isinstance(extracted_points, list):
                points_text = '、'.join(extracted_points[:10])
                comment_lines.append(f"【提取的考点】{points_text}")
            
            # 来源信息
            if source:
                source_type = source.get('type', '')
                if source_type == 'document':
                    path = source.get('path', [])
                    if path:
                        path_str_doc = ' -> '.join([p.get('title', '') for p in path])
                        comment_lines.append(f"【来源路径】{path_str_doc}")
                    comment_lines.append(f"【来源文件】{source.get('filename', '未知')}")
                elif source_type == 'web':
                    search_results = source.get('search_results', [])
                    if search_results:
                        links = [f"{sr.get('title', '')}: {sr.get('url', '')}" for sr in search_results[:3]]
                        comment_lines.append(f"【参考链接】\n" + '\n'.join(links))
            
            # 添加批注到"需求"段落
            if comment_lines and req_para.runs:
                try:
                    comment_text = '\n'.join(comment_lines)
                    doc.add_comment(
                        runs=req_para.runs[0],
                        text=comment_text,
                        author='需求分析系统',
                        initials='RA'
                    )
                except Exception as e:
                    # 批注添加失败时降级：不影响主流程
                    logger.warning(f"[export_to_word] 添加批注失败: {e}")
            
            # ========== 5. 证明材料 ==========
            
            # ----- DOC_PROOF / UI_PROOF：证明方式建议 -----
            if intent in ('DOC_PROOF', 'UI_PROOF'):
                evidence_label = doc.add_paragraph()
                evidence_label.add_run("证明材料如下：").bold = True
                
                # 创建证明材料表格（2列：要求内容 | 执行结果）
                evidence_table = doc.add_table(rows=1, cols=2)
                evidence_table.style = 'Table Grid'
                
                # 表头
                header_cells = evidence_table.rows[0].cells
                header_cells[0].text = '要求内容'
                header_cells[1].text = '执行结果'
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # DOC_PROOF: 展示搜索到的文档链接
                doc_proof_links = result.get('doc_proof_links')
                if intent == 'DOC_PROOF' and doc_proof_links:
                    for link in doc_proof_links:
                        row = evidence_table.add_row().cells
                        link_title = link.get('title', '')
                        link_url = link.get('url', '')
                        link_snippet = link.get('snippet', '')
                        row[0].text = f"{link_title}\n{link_snippet[:120]}" if link_snippet else link_title
                        row[1].text = link_url if link_url else '（请补充截图）'
                elif intent == 'DOC_PROOF':
                    suggestions = [
                        "产品官方文档 / 技术白皮书",
                        "行业标准认证文件（如 SQL:2003 兼容性声明）",
                        "产品架构设计文档",
                        "第三方评测报告"
                    ]
                    for suggestion in suggestions:
                        row = evidence_table.add_row().cells
                        row[0].text = suggestion
                        row[1].text = '（请补充截图）'
                else:  # UI_PROOF
                    suggestions = [
                        "管理控制台功能截图",
                        "操作步骤录屏或截图序列",
                        "产品配置界面截图",
                        "监控面板截图"
                    ]
                    for suggestion in suggestions:
                        row = evidence_table.add_row().cells
                        row[0].text = suggestion
                        row[1].text = '（请补充截图）'
                
                doc.add_paragraph()
            
            # ----- test_cases 展示（SQL_TEST 意图或无 intent 时） -----
            elif test_cases and isinstance(test_cases, list):
                evidence_label = doc.add_paragraph()
                evidence_label.add_run("证明材料如下：").bold = True
                
                # 创建证明材料表格（2列：要求内容 | 执行结果）
                evidence_table = doc.add_table(rows=1, cols=2)
                evidence_table.style = 'Table Grid'
                
                # 表头
                header_cells = evidence_table.rows[0].cells
                header_cells[0].text = '要求内容'
                header_cells[1].text = '执行结果'
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # 收集所有失败项
                failed_items = []
                
                # 每个测试点一行
                for tc_idx, tc in enumerate(test_cases):
                    test_point = tc.get('test_point', f'测试点{tc_idx+1}')
                    setup_sql = tc.get('setup_sql', '').strip()
                    verify_sql = tc.get('verify_sql', '').strip()
                    expected_behavior = tc.get('expected_behavior', '')
                    db_results = tc.get('db_results', {})
                    
                    row = evidence_table.add_row().cells
                    
                    # 左列：验证内容（测试点说明 + SQL）
                    verify_cell = row[0]
                    verify_cell.text = ''
                    desc_para = verify_cell.paragraphs[0]
                    desc_run = desc_para.add_run(f"{tc_idx+1}. {test_point}")
                    desc_run.bold = True
                    desc_run.font.size = Pt(10)
                    
                    # 右列：数据库执行结果（终端样式）
                    result_cell = row[1]
                    result_cell.text = ''
                    
                    has_db_output = False
                    for db_type, db_result in db_results.items():
                        db_name = db_result.get('db_name', db_type)
                        stmt_results = db_result.get('results', [])
                        
                        if not stmt_results and not db_result.get('error'):
                            continue
                        
                        has_db_output = True
                        
                        # 构建终端输出行
                        terminal_lines = []
                        if 'pg' in db_type.lower() or 'postgres' in db_type.lower() or 'oracle' in db_type.lower():
                            prompt_str = f"{db_name.lower().replace(' ', '_')}>"
                        else:
                            prompt_str = f"txsql>"
                        
                        for stmt_r in stmt_results:
                            sql_text = stmt_r.get('sql', '')
                            result_text = stmt_r.get('result', '')
                            error_text = stmt_r.get('error', '')
                            terminal_lines.append(f"{prompt_str} {sql_text}")
                            if result_text:
                                terminal_lines.append(result_text)
                            elif error_text:
                                terminal_lines.append(f"ERROR: {error_text}")
                            else:
                                terminal_lines.append("Query OK")
                        
                        if not stmt_results and db_result.get('error'):
                            terminal_lines.append(f"{prompt_str} -- 连接失败")
                            terminal_lines.append(f"ERROR: {db_result['error']}")
                        
                        terminal_lines.append(f"{prompt_str} ")
                        
                        if terminal_lines:
                            self._add_terminal_block_in_cell(
                                result_cell,
                                '\n'.join(terminal_lines),
                                db_name=db_name
                            )
                    
                    if not has_db_output:
                        no_result_para = result_cell.paragraphs[0] if result_cell.paragraphs else result_cell.add_paragraph()
                        no_result_para.text = '（请补充执行截图）'
                    
                    # 收集失败项
                    for db_type, db_result in db_results.items():
                        if not db_result.get('success'):
                            db_name = db_result.get('db_name', db_type)
                            error_msg = db_result.get('error', '')
                            if not error_msg:
                                for stmt_r in db_result.get('results', []):
                                    if stmt_r.get('error'):
                                        error_msg = stmt_r['error']
                            failed_items.append({
                                'test_point': test_point,
                                'db_name': db_name,
                                'error': error_msg
                            })
                
                # 表格后：失败项红字标注
                if failed_items:
                    doc.add_paragraph()
                    for fi in failed_items:
                        fail_para = doc.add_paragraph()
                        fail_run = fail_para.add_run(
                            f"❌ {fi['test_point']} - {fi['db_name']}: "
                            f"{fi['error'][:200] if fi['error'] else '验证失败'}"
                        )
                        fail_run.font.color.rgb = RGBColor(204, 0, 0)
                        fail_run.font.size = Pt(10)
                    
                    # 在标题处添加批注，汇总所有验证失败的考点
                    if title_heading and title_heading.runs:
                        try:
                            # 按考点去重汇总
                            failed_points_set = []
                            seen_points = set()
                            for fi in failed_items:
                                tp = fi['test_point']
                                if tp not in seen_points:
                                    seen_points.add(tp)
                                    failed_points_set.append(tp)
                            
                            fail_comment_lines = ['⚠️ 以下考点SQL验证失败，需要人工确认：']
                            for fp in failed_points_set:
                                # 找到该考点对应的所有失败数据库和错误
                                db_errors = [
                                    f"{fi['db_name']}: {fi['error'][:100] if fi['error'] else '验证失败'}"
                                    for fi in failed_items if fi['test_point'] == fp
                                ]
                                fail_comment_lines.append(f"• {fp}")
                                for de in db_errors:
                                    fail_comment_lines.append(f"  - {de}")
                            
                            doc.add_comment(
                                runs=title_heading.runs[0],
                                text='\n'.join(fail_comment_lines),
                                author='SQL验证',
                                initials='SV'
                            )
                        except Exception as e:
                            logger.warning(f"[export_to_word] 添加SQL失败批注失败: {e}")
                
                doc.add_paragraph()
            
            # ----- SQL验证结果表格（Function Calling 模式） -----
            elif sql_results and isinstance(sql_results, dict):
                evidence_label = doc.add_paragraph()
                evidence_label.add_run("证明材料如下：").bold = True
                
                # 创建表格
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                
                header_cells = table.rows[0].cells
                header_cells[0].text = '数据库类型'
                header_cells[1].text = '执行状态'
                header_cells[2].text = '支持情况'
                header_cells[3].text = '结果/错误'
                
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for db_type, db_result in sql_results.items():
                    row = table.add_row().cells
                    row[0].text = db_result.get('db_name', db_type)
                    row[1].text = '✅ 成功' if db_result.get('success') else '❌ 失败'
                    row[2].text = '支持' if db_result.get('supported') else '不支持'
                    error_or_result = db_result.get('result', '') or db_result.get('error', '')
                    row[3].text = error_or_result[:200] if error_or_result else '-'
                
                # 收集Function Calling模式的失败项，添加批注到标题
                fc_failed = []
                for db_type, db_result in sql_results.items():
                    if not db_result.get('success'):
                        db_name = db_result.get('db_name', db_type)
                        error_msg = db_result.get('error', '') or db_result.get('result', '')
                        fc_failed.append(f"• {db_name}: {error_msg[:100] if error_msg else '验证失败'}")
                
                if fc_failed and title_heading and title_heading.runs:
                    try:
                        fail_comment_text = '⚠️ 以下数据库SQL验证失败，需要人工确认：\n' + '\n'.join(fc_failed)
                        doc.add_comment(
                            runs=title_heading.runs[0],
                            text=fail_comment_text,
                            author='SQL验证',
                            initials='SV'
                        )
                    except Exception as e:
                        logger.warning(f"[export_to_word] 添加SQL失败批注失败: {e}")
                
                doc.add_paragraph()
            
            # ----- 无专项证明材料时：输出答案内容 -----
            else:
                has_evidence = False
                if answer and answer != '无答案':
                    evidence_label = doc.add_paragraph()
                    evidence_label.add_run("证明材料如下：").bold = True
                    
                    # 创建证明材料表格（2列）
                    evidence_table = doc.add_table(rows=1, cols=2)
                    evidence_table.style = 'Table Grid'
                    
                    header_cells = evidence_table.rows[0].cells
                    header_cells[0].text = '要求内容'
                    header_cells[1].text = '执行结果'
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    # 将答案内容放入验证内容列
                    row = evidence_table.add_row().cells
                    answer_images = source.get('images', []) if source else []
                    # 构建图片映射
                    img_map = {}
                    for img in (answer_images or []):
                        img_id = img.get('id')
                        if img_id:
                            img_map[str(img_id)] = img.get('image_path', '') or ''
                    
                    image_pattern = re.compile(r'\{\{IMAGE_ID_(\d+)\}\}')
                    
                    # 左列写答案文字
                    answer_cell = row[0]
                    answer_cell.text = ''
                    # 移除图片占位符后的纯文本
                    answer_text_clean = image_pattern.sub('', answer).strip()
                    if answer_text_clean:
                        answer_cell.paragraphs[0].text = answer_text_clean[:2000]
                    
                    # 右列写图片
                    img_cell = row[1]
                    img_cell.text = ''
                    has_images = False
                    for match in image_pattern.finditer(answer):
                        img_id = match.group(1)
                        img_path = img_map.get(img_id, '')
                        if img_path and os.path.exists(img_path):
                            try:
                                p = img_cell.add_paragraph() if has_images else img_cell.paragraphs[0]
                                run = p.add_run()
                                run.add_picture(img_path, width=Inches(3))
                                has_images = True
                            except Exception as e:
                                logger.warning(f"[export_to_word] 添加图片失败: {img_path}, 错误: {e}")
                    
                    if not has_images:
                        img_cell.paragraphs[0].text = '（请补充截图）'
                    
                    doc.add_paragraph()
        
        # ========== 保存到临时文件 ==========
        temp_dir = tempfile.gettempdir()
        filename = f"需求分析报告_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return filepath, filename
    
    @staticmethod
    def _add_terminal_code_block(doc, code_text, is_error=False):
        """
        在Word文档中添加终端代码块样式（深色背景 + 等宽字体）
        模拟终端/命令行截图效果
        
        Args:
            doc: Document 对象
            code_text: 代码文本
            is_error: 是否为错误信息（红色文字）
        """
        from docx.shared import Pt, RGBColor, Cm
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # 创建一个单行单列的表格作为代码块容器
        table = doc.add_table(rows=1, cols=1)
        table.style = 'Table Grid'
        
        cell = table.cell(0, 0)
        
        # 设置单元格深色背景 (#1E1E1E 深灰，模拟终端)
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
        cell._tc.get_or_add_tcPr().append(shading)
        
        # 设置单元格内边距
        tc_pr = cell._tc.get_or_add_tcPr()
        tc_margins = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="80" w:type="dxa"/>'
            f'  <w:left w:w="120" w:type="dxa"/>'
            f'  <w:bottom w:w="80" w:type="dxa"/>'
            f'  <w:right w:w="120" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tc_pr.append(tc_margins)
        
        # 清除默认段落
        cell.paragraphs[0].clear()
        
        # 按行添加文本
        lines = code_text.split('\n')
        for i, line in enumerate(lines):
            if i == 0:
                para = cell.paragraphs[0]
            else:
                para = cell.add_paragraph()
            
            # 设置段落间距为紧凑
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = Pt(14)
            
            run = para.add_run(line)
            run.font.name = 'Consolas'
            # 设置中文字体为等宽
            run.font.element.rPr.rFonts.set(
                '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 
                'SimSun'
            )
            run.font.size = Pt(9)
            
            if is_error:
                run.font.color.rgb = RGBColor(255, 100, 100)  # 红色错误文字
            elif line.strip().startswith('mysql>') or line.strip().startswith('postgres>'):
                run.font.color.rgb = RGBColor(80, 200, 120)   # 绿色提示符
            elif line.strip().startswith('ERROR'):
                run.font.color.rgb = RGBColor(255, 100, 100)  # 红色错误
            else:
                run.font.color.rgb = RGBColor(212, 212, 212)   # 浅灰色文字
    
    def _add_answer_with_images(self, doc, answer, images):
        """
        将答案内容添加到Word文档，处理 [表格]...[/表格] 和 {{IMAGE_ID_xxx}} 占位符
        
        Args:
            doc: Word Document 对象
            answer: 答案文本（可能包含 {{IMAGE_ID_xxx}} 和 [表格]...[/表格]）
            images: 图片列表 [{id, image_url, image_path}, ...]
        """
        import re
        from docx.shared import Inches
        
        if not answer:
            doc.add_paragraph('无答案')
            return
        
        # 构建图片ID -> 本地路径映射
        img_map = {}
        for img in (images or []):
            img_id = img.get('id')
            if img_id:
                img_map[str(img_id)] = img.get('image_path', '') or ''
        
        # 正则
        table_pattern = re.compile(r'\[表格\](.*?)\[/表格\]', re.DOTALL)
        image_pattern = re.compile(r'\{\{IMAGE_ID_(\d+)\}\}')
        
        last_end = 0
        for match in table_pattern.finditer(answer):
            # 表格之前的普通文本（含图片替换）
            before_text = answer[last_end:match.start()]
            if before_text.strip():
                self._add_text_with_images_to_doc(doc, before_text, img_map, image_pattern)
            
            # 表格内容
            table_content = match.group(1).strip()
            if table_content:
                table = doc.add_table(rows=1, cols=1)
                table.style = 'Table Grid'
                cell = table.cell(0, 0)
                self._add_cell_content_with_images_to_doc(cell, table_content, img_map, image_pattern)
            
            last_end = match.end()
        
        # 剩余的普通文本
        remaining = answer[last_end:]
        if remaining.strip():
            self._add_text_with_images_to_doc(doc, remaining, img_map, image_pattern)
    
    def _add_text_with_images_to_doc(self, doc, text, img_map, image_pattern):
        """将文本添加到文档，遇到 {{IMAGE_ID_xxx}} 替换为图片"""
        import os
        from docx.shared import Inches
        
        last_end = 0
        for match in image_pattern.finditer(text):
            before = text[last_end:match.start()]
            if before.strip():
                doc.add_paragraph(before.strip())
            
            img_id = match.group(1)
            img_path = img_map.get(img_id, '')
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logger.warning(f"[export_to_word] 添加图片失败: {img_path}, 错误: {e}")
                    doc.add_paragraph(f"[图片加载失败: {img_path}]")
            else:
                logger.warning(f"[export_to_word] 图片文件不存在: img_id={img_id}, path='{img_path}'")
                doc.add_paragraph(f"[图片未找到: ID={img_id}]")
            
            last_end = match.end()
        
        remaining = text[last_end:]
        if remaining.strip():
            doc.add_paragraph(remaining.strip())
    
    def _add_cell_content_with_images_to_doc(self, cell, text, img_map, image_pattern):
        """将文本添加到表格单元格，遇到 {{IMAGE_ID_xxx}} 替换为图片"""
        import os
        from docx.shared import Inches
        
        cell.text = ''
        
        last_end = 0
        for match in image_pattern.finditer(text):
            before = text[last_end:match.start()]
            if before.strip():
                cell.add_paragraph(before.strip())
            
            img_id = match.group(1)
            img_path = img_map.get(img_id, '')
            if img_path and os.path.exists(img_path):
                try:
                    p = cell.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(4.5))
                except Exception as e:
                    logger.warning(f"[export_to_word] 添加图片到表格失败: {img_path}, 错误: {e}")
                    cell.add_paragraph(f"[图片加载失败: {img_path}]")
            else:
                logger.warning(f"[export_to_word] 表格图片文件不存在: img_id={img_id}, path='{img_path}'")
                cell.add_paragraph(f"[图片未找到: ID={img_id}]")
            
            last_end = match.end()
        
        remaining = text[last_end:]
        if remaining.strip():
            cell.add_paragraph(remaining.strip())

    @staticmethod
    def _add_terminal_block_in_cell(parent_cell, code_text, db_name=''):
        """
        在已有的表格单元格内嵌入一个深色终端代码块
        
        实现方式：在单元格内创建一个嵌套表格（1行1列），设置深色背景。
        这样可以在一个大的"证明材料"框内同时包含：
        - 上方的白色背景说明文字
        - 下方的深色终端代码块
        
        Args:
            parent_cell: 父级单元格对象
            code_text: 终端输出文本
            db_name: 数据库名称（用于显示标题）
        """
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        
        # 在父单元格中添加一个嵌套表格
        # 先添加一个空段落作为间距
        spacer = parent_cell.add_paragraph()
        spacer.paragraph_format.space_before = Pt(4)
        spacer.paragraph_format.space_after = Pt(0)
        
        # 创建嵌套表格（python-docx 不直接支持在 cell 内嵌套表格，
        # 需要通过 XML 操作来实现——用段落+shading模拟即可）
        # 这里我们用段落 + 段落底纹来模拟深色背景
        
        lines = code_text.split('\n')
        for i, line in enumerate(lines):
            para = parent_cell.add_paragraph()
            
            # 设置段落底纹（深色背景）
            pPr = para._p.get_or_add_pPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E1E1E" w:val="clear"/>')
            pPr.append(shading)
            
            # 设置段落间距为紧凑
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = Pt(14)
            
            # 左侧缩进一点，模拟内边距
            para.paragraph_format.left_indent = Pt(6)
            
            run = para.add_run(line)
            run.font.name = 'Consolas'
            # 设置中文字体
            rPr = run.font.element.rPr
            if rPr is not None:
                rPr.rFonts.set(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia',
                    'SimSun'
                )
            run.font.size = Pt(8)
            
            # 根据内容设置颜色
            stripped = line.strip()
            if stripped.startswith('txsql>') or stripped.startswith('mysql>') or stripped.startswith('postgres>') or '>' in stripped[:20]:
                run.font.color.rgb = RGBColor(80, 200, 120)    # 绿色提示符行
            elif stripped.startswith('ERROR') or stripped.startswith('error'):
                run.font.color.rgb = RGBColor(255, 100, 100)   # 红色错误
            elif stripped.startswith('Query OK') or stripped.startswith('Rows matched'):
                run.font.color.rgb = RGBColor(180, 180, 180)   # 灰色状态行
            else:
                run.font.color.rgb = RGBColor(212, 212, 212)    # 浅灰色普通文字
    
    def parse_txt_as_tree(self, file_path):
        """
        解析TXT文件并返回完整的树结构（包括非叶子节点）
        
        Returns:
            dict: 树结构字典，包含所有节点和叶子节点
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
        
        root, all_nodes = self._build_requirement_tree(lines)
        leaves = root.get_leaf_nodes()
        
        return {
            'tree': root.to_dict(),
            'total_nodes': len(all_nodes),
            'leaf_count': len(leaves),
            'leaves': [
                {
                    'number': l.number,
                    'title': l.title,
                    'content': l.content,
                    'level': l.level,
                    'full_path': l.get_full_path()
                }
                for l in leaves
            ]
        }


def get_requirement_analyzer(llm_config_id=None, username=None):
    """获取需求分析器实例"""
    return RequirementAnalyzer(llm_config_id, username=username)
