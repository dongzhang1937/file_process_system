"""
RAG 向量检索服务
- 文档解析（docx → 章节 → 分块）
- Embedding 入库（pgvector）
- 语义检索（返回文档名/章节/段落信息）

三层结构：rag_documents → rag_sections → rag_chunks
PG 向量库：127.0.0.1 / t11 / t1（pgvector 扩展）
Embedding：sentence-transformers bge-m3（119.45.183.233:11435）
"""
from __future__ import annotations

import os
import re
import json
import hashlib
from typing import Any

import psycopg2
import psycopg2.extras
from config.logging_config import logger

# ==================== PG 向量库连接配置（从统一配置文件读取） ====================

def _get_rag_pg_config() -> dict[str, str | int]:
    """获取 RAG 向量信息库的 PG 连接配置"""
    from config.app_config import get_rag_pg_config
    return get_rag_pg_config()

# ==================== 分块参数 ====================

CHUNK_SIZE = 800       # 每个 chunk 的最大字符数（约 512 tokens）
CHUNK_OVERLAP = 100    # 滑动窗口重叠字符数


# 管理员用户名常量
ADMIN_USERNAME = 'asd'


def _get_pg_conn() -> Any:
    """获取 pgvector 数据库连接"""
    conn = psycopg2.connect(**_get_rag_pg_config())  # pyright: ignore[reportCallIssue,reportArgumentType]
    conn.autocommit = False
    return conn


def _get_file_hash(filepath: str) -> str:
    """计算文件 MD5 哈希"""
    h = hashlib.md5()
    with open(filepath, 'rb') as f:
        for chunk in iter(lambda: f.read(8192), b''):
            h.update(chunk)
    return h.hexdigest()


def _content_hash(text: str) -> str:
    """计算文本内容哈希"""
    return hashlib.sha256(text.encode('utf-8')).hexdigest()[:32]


# ==================== 文档解析 ====================

def parse_docx_to_sections(filepath: str) -> list[dict[str, Any]]:
    """
    解析 Word 文档为章节列表（保留层级结构）
    
    Returns:
        [
            {
                'title': '3.2 备份恢复',
                'section_number': '3.2',
                'heading_level': 2,
                'full_path': '3 运维管理 > 3.2 备份恢复',
                'content': '章节完整文本...',
                'order': 5,
                'children_indices': [6, 7]  # 子章节在列表中的索引
            },
            ...
        ]
    """
    from docx import Document
    
    doc = Document(filepath)
    sections = []
    
    # 维护标题路径栈：{level: title}
    path_stack: dict[int, str] = {}
    current_section: dict[str, Any] | None = None
    section_order = 0
    
    for para in doc.paragraphs:
        style_name = para.style.name if para.style and para.style.name else ''
        text = para.text.strip()
        
        if not text:
            continue
        
        # 判断是否为标题
        heading_level = 0
        if style_name.startswith('Heading'):
            try:
                heading_level = int(style_name.replace('Heading ', '').replace('Heading', '').strip())
            except (ValueError, AttributeError):
                heading_level = 0
        
        # 如果不是标题样式，尝试用正则检测（章节编号格式）
        if heading_level == 0 and re.match(r'^\d+(\.\d+)*\s+\S', text):
            # 根据编号层级推断标题等级
            num_match = re.match(r'^(\d+(?:\.\d+)*)\s', text)
            if num_match:
                num_parts = num_match.group(1).split('.')
                heading_level = len(num_parts)
        
        if heading_level > 0:
            # 保存上一个章节
            if current_section:
                sections.append(current_section)
            
            # 提取章节编号
            num_match = re.match(r'^(\d+(?:\.\d+)*)\s*(.*)', text)
            if num_match:
                section_number = num_match.group(1)
                title_text = text
            else:
                section_number = ''
                title_text = text
            
            # 更新路径栈
            path_stack[heading_level] = title_text
            # 清除更深层级
            for lvl in list(path_stack.keys()):
                if lvl > heading_level:
                    del path_stack[lvl]
            
            # 构建完整路径
            full_path = ' > '.join(
                path_stack[lvl] for lvl in sorted(path_stack.keys()) if lvl <= heading_level
            )
            
            section_order += 1
            current_section = {
                'title': title_text,
                'section_number': section_number,
                'heading_level': heading_level,
                'full_path': full_path,
                'content': '',
                'order': section_order,
            }
        elif current_section:
            # 普通段落，追加到当前章节
            if current_section['content']:
                current_section['content'] += '\n' + text
            else:
                current_section['content'] = text
        else:
            # 文档开头没有标题的段落，创建一个虚拟章节
            section_order += 1
            current_section = {
                'title': '文档概述',
                'section_number': '0',
                'heading_level': 1,
                'full_path': '文档概述',
                'content': text,
                'order': section_order,
            }
    
    # 别忘了最后一个章节
    if current_section:
        sections.append(current_section)
    
    # 同时解析表格内容，追加到最近的章节
    # （docx 的 paragraphs 不包含表格，需要单独处理）
    try:
        for table in doc.tables:
            table_text_parts = []
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_texts.append(cell_text)
                if row_texts:
                    table_text_parts.append(' | '.join(row_texts))
            
            if table_text_parts and sections:
                # 追加到最后一个章节
                table_content = '\n'.join(table_text_parts)
                sections[-1]['content'] += '\n[表格]\n' + table_content
    except Exception as e:
        logger.warning(f"解析表格时出错: {e}")
    
    logger.info(f"文档解析完成: {os.path.basename(filepath)}, 共 {len(sections)} 个章节")
    return sections


def split_section_to_chunks(section_content: str, section_title: str = '',
                             chunk_size: int = CHUNK_SIZE,
                             overlap: int = CHUNK_OVERLAP) -> list[dict[str, Any]]:
    """
    将章节文本分割为检索块（滑动窗口）
    
    策略：
    1. 优先按段落（双换行）切分
    2. 如果单段落 > chunk_size，用滑动窗口切分
    3. 小段落合并到一个 chunk 中直到达到 chunk_size
    
    Returns:
        [{'content': '块文本', 'chunk_index': 0}, ...]
    """
    if not section_content or not section_content.strip():
        return []
    
    content = section_content.strip()
    
    # 如果整个章节内容很短，直接作为一个 chunk
    if len(content) <= chunk_size:
        return [{'content': content, 'chunk_index': 0}]
    
    # 按段落分割（双换行 或 单换行）
    paragraphs = re.split(r'\n{2,}', content)
    if len(paragraphs) == 1:
        paragraphs = content.split('\n')
    
    chunks = []
    current_chunk = ''
    chunk_index = 0
    
    for para in paragraphs:
        para = para.strip()
        if not para:
            continue
        
        if len(para) > chunk_size:
            # 单段落太长，先保存当前积累的 chunk
            if current_chunk:
                chunks.append({'content': current_chunk.strip(), 'chunk_index': chunk_index})
                chunk_index += 1
                current_chunk = ''
            
            # 用滑动窗口切分长段落
            start = 0
            while start < len(para):
                end = start + chunk_size
                chunk_text = para[start:end]
                chunks.append({'content': chunk_text.strip(), 'chunk_index': chunk_index})
                chunk_index += 1
                start += chunk_size - overlap
        
        elif len(current_chunk) + len(para) + 1 > chunk_size:
            # 当前 chunk 已满，保存并开始新 chunk
            if current_chunk:
                chunks.append({'content': current_chunk.strip(), 'chunk_index': chunk_index})
                chunk_index += 1
            # 新 chunk 开始时带上上一个 chunk 尾部的 overlap
            if overlap > 0 and current_chunk:
                overlap_text = current_chunk[-overlap:]
                current_chunk = overlap_text + '\n' + para
            else:
                current_chunk = para
        else:
            # 追加到当前 chunk
            if current_chunk:
                current_chunk += '\n' + para
            else:
                current_chunk = para
    
    # 最后一个 chunk
    if current_chunk.strip():
        chunks.append({'content': current_chunk.strip(), 'chunk_index': chunk_index})
    
    return chunks


# ==================== Embedding + 入库 ====================

def _get_embedding_service(username: str = ADMIN_USERNAME) -> Any:
    """
    获取 Embedding 服务
    统一从数据库 embedding_configs 表读取默认配置，不硬编码。
    如果未配置，会抛出明确的错误提示用户去页面配置。
    """
    from .embedding_service import EmbeddingService
    
    service = EmbeddingService(username=username)
    if service.provider is None:
        raise RuntimeError(
            "Embedding 服务未配置或配置无效。"
            "请在系统设置中配置 Embedding 模型（提供商、模型名、API地址）并设为默认。"
        )
    return service


def embed_and_store_document(filepath: str, force: bool = False, username: str = ADMIN_USERNAME) -> dict[str, Any]:
    """
    完整流程：解析文档 → 分块 → 生成向量 → 存入 pgvector
    
    Args:
        filepath: docx 文件路径
        force: 强制重新处理（忽略文件哈希去重）
        username: 所属用户名
    
    Returns:
        {'success': bool, 'document_id': int, 'sections': int, 'chunks': int, 'message': str}
    """
    filename = os.path.basename(filepath)
    file_hash = _get_file_hash(filepath)
    
    conn = _get_pg_conn()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    doc_id: int | None = None
    
    try:
        # 检查当前用户是否已处理过同一文件，避免跨用户互相覆盖
        if not force:
            cur.execute(
                "SELECT id, status FROM rag_documents WHERE file_hash = %s AND username = %s",
                (file_hash, username),
            )
            existing = cur.fetchone()
            if existing and existing['status'] == 'done':
                return {
                    'success': True,
                    'document_id': existing['id'],
                    'sections': 0, 'chunks': 0,
                    'message': f'文档已存在（id={existing["id"]}），跳过。使用 force=True 强制重新处理。'
                }
            elif existing:
                # 之前处理失败或进行中，仅删除当前用户自己的旧数据后重来
                cur.execute("DELETE FROM rag_documents WHERE id = %s", (existing['id'],))
                conn.commit()
        else:
            # 强制模式：仅删除当前用户自己的旧数据
            cur.execute(
                "DELETE FROM rag_documents WHERE file_hash = %s AND username = %s",
                (file_hash, username),
            )
            conn.commit()
        
        # 从文件名提取产品名和文档类型
        product_name, doc_type = _extract_product_info(filename)
        
        # 插入文档记录
        cur.execute("""
            INSERT INTO rag_documents (filename, filepath, file_hash, product_name, doc_type, 
                                       embedding_model, username, status)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'processing')
            RETURNING id
        """, (filename, filepath, file_hash, product_name, doc_type, 'bge-m3:latest', username))
        doc_id = cur.fetchone()['id']
        conn.commit()
        
        logger.info(f"[RAG] 开始处理文档: {filename} (doc_id={doc_id})")
        
        # 1. 解析文档为章节
        sections = parse_docx_to_sections(filepath)
        if not sections:
            cur.execute("UPDATE rag_documents SET status='error', error_message='无法解析出章节' WHERE id=%s", (doc_id,))
            conn.commit()
            return {'success': False, 'document_id': doc_id, 'sections': 0, 'chunks': 0,
                    'message': '无法从文档中解析出章节'}
        
        # 2. 获取 embedding 服务
        embed_service = _get_embedding_service(username=username)
        
        # 3. 逐章节处理：分块 → embedding → 入库
        total_chunks = 0
        total_sections = 0
        
        for sec in sections:
            # 插入章节记录
            cur.execute("""
                INSERT INTO rag_sections (document_id, title, section_number, heading_level,
                                          full_path, content, content_length, section_order)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s)
                RETURNING id
            """, (
                doc_id, sec['title'], sec['section_number'], sec['heading_level'],
                sec['full_path'], sec['content'], len(sec['content']), sec['order']
            ))
            section_id = cur.fetchone()['id']
            total_sections += 1
            
            # 分块
            chunks = split_section_to_chunks(sec['content'], sec['title'])
            
            if not chunks:
                continue
            
            logger.info(f"[RAG] 章节 '{sec['title'][:40]}' 切分为 {len(chunks)} 个 chunks，开始 embedding...")
            
            # 批量 embedding（每批 50 条，利用批量接口加速）
            batch_size = 50
            section_chunk_count = 0
            
            for i in range(0, len(chunks), batch_size):
                batch = chunks[i:i + batch_size]
                texts = [c['content'] for c in batch]
                
                try:
                    embeddings = embed_service.embed_texts(texts)
                except Exception as e:
                    logger.error(f"[RAG] Embedding 失败 (section={sec['title'][:30]}): {e}")
                    continue
                
                if i % 100 == 0 and i > 0:
                    logger.info(f"[RAG]   进度: {i}/{len(chunks)} chunks")
                
                for chunk_data, embedding in zip(batch, embeddings):
                    c_hash = _content_hash(chunk_data['content'])
                    metadata = {
                        'section_title': sec['title'],
                        'section_number': sec['section_number'],
                        'full_path': sec['full_path'],
                        'heading_level': sec['heading_level'],
                    }
                    
                    # 转为 pgvector 格式字符串
                    vec_str = '[' + ','.join(str(v) for v in embedding) + ']'
                    
                    cur.execute("""
                        INSERT INTO rag_chunks (document_id, section_id, chunk_index, content,
                                                content_length, content_hash, embedding, metadata)
                        VALUES (%s, %s, %s, %s, %s, %s, %s::vector, %s::jsonb)
                    """, (
                        doc_id, section_id, chunk_data['chunk_index'],
                        chunk_data['content'], len(chunk_data['content']),
                        c_hash, vec_str, json.dumps(metadata, ensure_ascii=False)
                    ))
                    section_chunk_count += 1
                    total_chunks += 1
            
            # 更新章节的 chunk 数量
            cur.execute("UPDATE rag_sections SET chunk_count = %s WHERE id = %s",
                       (section_chunk_count, section_id))
        
        # 更新文档状态
        cur.execute("""
            UPDATE rag_documents 
            SET status='done', total_sections=%s, total_chunks=%s, updated_at=CURRENT_TIMESTAMP
            WHERE id=%s
        """, (total_sections, total_chunks, doc_id))
        conn.commit()
        
        logger.info(f"[RAG] 文档处理完成: {filename}, {total_sections} 章节, {total_chunks} chunks")
        return {
            'success': True,
            'document_id': doc_id,
            'sections': total_sections,
            'chunks': total_chunks,
            'message': f'成功：{total_sections} 个章节，{total_chunks} 个检索块'
        }
        
    except Exception as e:
        conn.rollback()
        logger.error(f"[RAG] 文档处理异常: {filename}: {e}", exc_info=True)
        try:
            if doc_id is not None:
                cur.execute("UPDATE rag_documents SET status='error', error_message=%s WHERE id=%s",
                           (str(e)[:500], doc_id))
                conn.commit()
        except Exception:
            pass
        return {'success': False, 'document_id': None, 'sections': 0, 'chunks': 0,
                'message': f'处理失败: {e}'}
    finally:
        cur.close()
        conn.close()


def embed_directory(dir_path: str, force: bool = False, username: str = ADMIN_USERNAME) -> dict[str, Any]:
    """
    批量处理目录下所有 docx 文件
    
    Args:
        dir_path: 目录路径
        force: 强制重新处理
        username: 所属用户名
    
    Returns:
        {'total': int, 'success': int, 'skipped': int, 'failed': int, 'details': [...]} 
    """
    results: dict[str, Any] = {'total': 0, 'success': 0, 'skipped': 0, 'failed': 0, 'details': []}
    
    for root, _dirs, files in os.walk(dir_path):
        for fname in files:
            if not fname.lower().endswith('.docx') or fname.startswith('~$'):
                continue
            
            filepath = os.path.join(root, fname)
            results['total'] += 1
            
            logger.info(f"[RAG] 处理 {results['total']}: {fname}")
            result = embed_and_store_document(filepath, force=force, username=username)
            
            detail = {'filename': fname, **result}
            results['details'].append(detail)
            
            if result['success']:
                if result['chunks'] == 0 and '已存在' in result.get('message', ''):
                    results['skipped'] += 1
                else:
                    results['success'] += 1
            else:
                results['failed'] += 1
    
    logger.info(f"[RAG] 批量处理完成: 总计 {results['total']}, "
               f"成功 {results['success']}, 跳过 {results['skipped']}, 失败 {results['failed']}")
    return results


def _extract_product_info(filename: str) -> tuple[str, str]:
    """从文件名提取产品名和文档类型"""
    name = filename.replace('.docx', '').replace('.DOCX', '')
    
    doc_types = [
        '白皮书', '产品白皮书', '用户指南', '运维手册', '运维管理指南', '部署手册',
        '安装部署', '开发手册', '开发', 'API参考', 'API 参考', 'API 接口',
        '参考手册', '快速部署', '升级指南', '应急指南', '故障恢复', '应急方案',
        '版本说明书', '调优指南', '调优手册', '调优', '系统原理', '简介',
        '产品简介', '产品介绍', '扩容指南', '验收测试', '功能测试',
        '常见问题', '已知问题', '错误信息', '参数说明', '告警参考',
        '命令参考', '日志参考', '监控指标', '配置文件', '生命周期',
        '购买指南', '试用', '安全', '工具', '对接调测', '典型故障',
        '售后手册', '故障应急', '安全能力', '容灾与高可用', '物理节点变更',
        'Release Notes', 'Release Note', 'JDBC', '驱动说明',
    ]
    
    doc_type = '其他'
    for dt in doc_types:
        if dt in name:
            doc_type = dt
            break
    
    product_name = name
    product_patterns = [
        r'(腾讯云TDSQL\s*(?:MySQL|PG)?\s*(?:版|企业版|基础版|生态工具)?(?:[^_]*?))',
        r'(数据库智能管家\s*DBbrain[^_]*)',
        r'(数据库迁移服务\s*DBbridge[^_]*)',
    ]
    for p in product_patterns:
        m = re.match(p, name)
        if m:
            product_name = m.group(1).strip()
            break
    
    return product_name, doc_type


# ==================== 向量检索 ====================

def search_similar_chunks(query: str, top_k: int = 10, threshold: float = 0.5,
                          product_filter: str | None = None, username: str = ADMIN_USERNAME) -> list[dict[str, Any]]:
    """
    语义检索：返回最相关的文档块，带完整来源信息
    
    Args:
        query: 查询文本
        top_k: 返回数量
        threshold: 相似度阈值（0-1，余弦相似度）
        product_filter: 可选的产品名过滤
        username: 用户名，普通用户仅检索自己的文档，管理员检索全部
    
    Returns:
        [...]
    """
    # 生成查询向量
    embed_service = _get_embedding_service(username=username)
    query_embedding = embed_service.embed_text(query)
    vec_str = '[' + ','.join(str(v) for v in query_embedding) + ']'
    
    conn = _get_pg_conn()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    
    try:
        # 构建查询（使用 pgvector 的余弦距离运算符 <=>）
        # 余弦距离 = 1 - 余弦相似度，所以 distance 越小越相似
        where_clauses = ["d.status = 'done'"]
        params: list[object] = [vec_str]

        if product_filter:
            where_clauses.append("d.product_name ILIKE %s")
            params.append(f'%{product_filter}%')

        if username != ADMIN_USERNAME:
            where_clauses.append("(d.username = %s OR d.username = %s)")
            params.append(username)
            params.append(ADMIN_USERNAME)

        where_clauses.append("(c.embedding <=> %s::vector) < %s")
        params.extend([vec_str, 1.0 - threshold, vec_str, top_k])

        sql = f"""
            SELECT 
                c.id AS chunk_id,
                c.content,
                c.chunk_index,
                c.metadata,
                1 - (c.embedding <=> %s::vector) AS similarity,
                d.id AS doc_id,
                d.filename,
                d.product_name,
                d.doc_type,
                s.id AS section_id,
                s.title AS section_title,
                s.section_number,
                s.heading_level,
                s.full_path,
                s.content AS section_content
            FROM rag_chunks c
            JOIN rag_documents d ON c.document_id = d.id
            JOIN rag_sections s ON c.section_id = s.id
            WHERE {' AND '.join(where_clauses)}
            ORDER BY c.embedding <=> %s::vector ASC
            LIMIT %s
        """

        cur.execute(sql, params)
        rows = cur.fetchall()
        
        results = []
        for row in rows:
            results.append({
                'content': row['content'],
                'similarity': round(float(row['similarity']), 4),
                'document': {
                    'id': row['doc_id'],
                    'filename': row['filename'],
                    'product_name': row['product_name'],
                    'doc_type': row['doc_type'],
                },
                'section': {
                    'id': row['section_id'],
                    'title': row['section_title'],
                    'section_number': row['section_number'],
                    'heading_level': row['heading_level'],
                    'full_path': row['full_path'],
                },
                'chunk_index': row['chunk_index'],
                # 父块内容（完整章节文本，用于上下文补充）
                'section_content': row['section_content'][:2000] if row['section_content'] else '',
            })
        
        logger.info(f"[RAG] 检索完成: query='{query[:50]}...', 返回 {len(results)} 条结果")
        return results
        
    except Exception as e:
        logger.error(f"[RAG] 检索异常: {e}", exc_info=True)
        return []
    finally:
        cur.close()
        conn.close()


def format_search_results_as_context(results: list[dict[str, Any]], max_length: int = 4000) -> str:
    """
    将检索结果格式化为 LLM 上下文（替代 web_search 的输出格式）
    
    Args:
        results: search_similar_chunks 的返回结果
        max_length: 上下文最大字符数
    
    Returns:
        格式化的文本，可直接作为 LLM 的参考资料
    """
    if not results:
        return "未检索到相关文档内容。"
    
    parts = []
    current_length = 0
    
    for i, r in enumerate(results):
        source = (
            f"📄 文档: {r['document']['filename']}\n"
            f"📂 章节: {r['section']['full_path']}\n"
            f"📊 相似度: {r['similarity']}"
        )
        content = r['content']
        
        block = f"[参考资料 {i+1}]\n{source}\n\n{content}\n"
        
        if current_length + len(block) > max_length:
            break
        
        parts.append(block)
        current_length += len(block)
    
    return '\n---\n'.join(parts)


# ==================== 管理接口 ====================

def get_rag_stats(username: str = ADMIN_USERNAME) -> dict[str, Any]:
    """获取 RAG 系统统计信息（普通用户仅统计自己的文档）"""
    conn = _get_pg_conn()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    
    try:
        doc_where = ""
        doc_params = ()
        done_where = "WHERE status = 'done'"
        done_params = ()
        if username != ADMIN_USERNAME:
            doc_where = "WHERE (username = %s OR username = %s)"
            doc_params = (username, ADMIN_USERNAME)
            done_where = "WHERE status = 'done' AND (username = %s OR username = %s)"
            done_params = (username, ADMIN_USERNAME)

        cur.execute(f"""
            SELECT 
                COUNT(*) AS total_documents,
                COUNT(*) FILTER (WHERE status = 'done') AS done_documents,
                COUNT(*) FILTER (WHERE status = 'error') AS error_documents,
                COUNT(*) FILTER (WHERE status = 'processing') AS processing_documents,
                COALESCE(SUM(total_sections), 0) AS total_sections,
                COALESCE(SUM(total_chunks), 0) AS total_chunks
            FROM rag_documents
            {doc_where}
        """, doc_params)
        stats = cur.fetchone()
        
        # 产品分布
        cur.execute(f"""
            SELECT product_name, COUNT(*) AS doc_count, SUM(total_chunks) AS chunk_count
            FROM rag_documents
            {done_where}
            GROUP BY product_name ORDER BY doc_count DESC
        """, done_params)
        products = cur.fetchall()
        
        return {
            'documents': dict(stats) if stats else {},
            'products': [dict(p) for p in products],
        }
    except Exception as e:
        logger.error(f"[RAG] 获取统计信息失败: {e}")
        return {'error': str(e)}
    finally:
        cur.close()
        conn.close()


def list_rag_documents(username: str = ADMIN_USERNAME) -> list[dict[str, Any]]:
    """列出 RAG 文档（管理员查看全部，普通用户查看自己的+admin的）"""
    conn = _get_pg_conn()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    
    try:
        if username == ADMIN_USERNAME:
            cur.execute("""
                SELECT id, filename, product_name, doc_type, total_sections, total_chunks,
                       status, error_message, username, created_at, updated_at
                FROM rag_documents
                ORDER BY created_at DESC
            """)
        else:
            cur.execute("""
                SELECT id, filename, product_name, doc_type, total_sections, total_chunks,
                       status, error_message, username, created_at, updated_at
                FROM rag_documents
                WHERE username = %s OR username = %s
                ORDER BY created_at DESC
            """, (username, ADMIN_USERNAME))
        docs = cur.fetchall()
        for d in docs:
            if d.get('created_at'):
                d['created_at'] = d['created_at'].isoformat()
            if d.get('updated_at'):
                d['updated_at'] = d['updated_at'].isoformat()
        return [dict(d) for d in docs]
    except Exception as e:
        logger.error(f"[RAG] 列出文档失败: {e}")
        return []
    finally:
        cur.close()
        conn.close()


def delete_rag_document(doc_id: int, username: str = ADMIN_USERNAME) -> bool:
    """删除指定文档及其所有章节和 chunks（普通用户只能删自己的）"""
    conn = _get_pg_conn()
    cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
    try:
        # 权限校验
        if username != ADMIN_USERNAME:
            cur.execute("SELECT username FROM rag_documents WHERE id = %s", (doc_id,))
            doc = cur.fetchone()
            if doc and doc.get('username', ADMIN_USERNAME) == ADMIN_USERNAME:
                return False  # 无权删除 admin 的文档
        
        cur2 = conn.cursor()
        cur2.execute("DELETE FROM rag_documents WHERE id = %s", (doc_id,))
        conn.commit()
        return cur2.rowcount > 0
    except Exception as e:
        conn.rollback()
        logger.error(f"[RAG] 删除文档失败: {e}")
        return False
    finally:
        cur.close()
        conn.close()
