"""
招标文档解析器模块
支持解析招标文档的章节结构、技术要求表格，并根据用户指令定位章节

Phase 1: Word文档章节解析（支持层级目录 + 表格解析）
Phase 2: 用户指令解析（"针对1.4.1,1.4.2作答"）
Phase 3: 三级匹配策略（精确匹配→语义匹配→网络搜索→LLM生成）
Phase 4: 作答结果整合与格式化
Phase 5: Word结果文档生成
"""
import os
import re
import json
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from docx import Document as DocumentLoader
from docx.document import Document as DocxDocument
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from difflib import SequenceMatcher

from config.logging_config import logger
from config.db_config import fetch_one, fetch_all


class BidDocumentParser:
    """招标文档解析器 - 专门处理招标/投标文档的章节结构"""
    
    # 章节编号正则模式 - 注意：只匹配真正的章节编号，不匹配列表项
    SECTION_NUMBER_PATTERNS = [
        # 1.4.1 或 1.4 格式（必须有点号分隔的多级编号）
        r'^(\d+\.\d+(?:\.\d+)*)\s*[\.、\s]?\s*(.+)',
        # 一、二、三 格式（中文章节）
        r'^([一二三四五六七八九十]+)[、\.]\s*(.+)',
        # 第一章、第二节 格式
        r'^第([一二三四五六七八九十\d]+)[章节条款]\s*(.+)',
    ]
    # 注意：不再包含 (1) 和 1) 格式，因为这些通常是列表项而不是章节标题
    
    # 技术要求表格的列名关键词
    TABLE_HEADER_KEYWORDS = ['序号', '技术要求', '技术规格', '功能要求', '参数', '指标', '规格', '要求']
    
    def __init__(self, file_path: str = None):
        """
        初始化解析器
        
        Args:
            file_path: Word文档路径
        """
        self.file_path = file_path
        self.doc = None
        self.document_structure = {}  # 解析后的文档结构
        self.section_index = {}  # 章节编号索引 {'1.4.1': {...}, '1.4.2': {...}}
        # 用于自动生成章节编号的计数器（支持10级标题）
        self.heading_counters = [0] * 10
        
        if file_path and os.path.exists(file_path):
            self.doc = DocumentLoader(file_path)
    
    def load_document(self, file_path: str):
        """加载文档"""
        self.file_path = file_path
        self.doc = DocumentLoader(file_path)
        self.document_structure = {}
        self.section_index = {}
        self.heading_counters = [0] * 10  # 重置计数器
    
    def parse_document_structure(self) -> Dict[str, Any]:
        """
        解析文档的完整目录结构
        
        Returns:
            {
                '1.4.1': {
                    'title': '国产数据库技术要求',
                    'number': '1.4.1',
                    'level': 3,
                    'content': '正文内容...',
                    'requirements': [
                        {'index': 1, 'text': '数据库软件需通过...', 'type': 'list'},
                    ],
                    'tables': [
                        {'headers': [...], 'rows': [...], 'type': 'requirement_table'}
                    ],
                    'parent_number': '1.4',
                    'children': ['1.4.1.1', '1.4.1.2']
                }
            }
        """
        if not self.doc:
            raise ValueError("请先加载文档")
        
        sections = {}
        current_section = None
        current_content = []
        
        # 遍历文档元素
        for element in self._iter_block_items(self.doc):
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if not text:
                    continue
                
                # 尝试识别章节标题
                section_info = self._parse_section_title(element)
                
                if section_info:
                    # 保存上一个章节的内容
                    if current_section:
                        current_section['content'] = '\n'.join(current_content)
                        current_content = []
                    
                    # 创建新章节
                    section_number = section_info['number']
                    current_section = {
                        'number': section_number,
                        'title': section_info['title'],
                        'level': section_info['level'],
                        'content': '',
                        'requirements': [],
                        'tables': [],
                        'parent_number': self._get_parent_number(section_number),
                        'children': [],
                        'raw_title': text
                    }
                    sections[section_number] = current_section
                    
                    # 更新父章节的children
                    parent_num = current_section['parent_number']
                    if parent_num and parent_num in sections:
                        if section_number not in sections[parent_num]['children']:
                            sections[parent_num]['children'].append(section_number)
                else:
                    # 普通段落，检查是否是需求项
                    if current_section:
                        req_item = self._parse_requirement_item(text)
                        if req_item:
                            current_section['requirements'].append(req_item)
                        current_content.append(text)
            
            elif isinstance(element, Table):
                if current_section:
                    table_data = self._parse_table(element)
                    if table_data:
                        current_section['tables'].append(table_data)
                        # 如果是需求表格，提取需求项
                        if table_data.get('type') == 'requirement_table':
                            table_reqs = self._extract_requirements_from_table(table_data)
                            current_section['requirements'].extend(table_reqs)
        
        # 保存最后一个章节
        if current_section:
            current_section['content'] = '\n'.join(current_content)
        
        self.document_structure = sections
        self.section_index = sections
        
        return sections
    
    def _iter_block_items(self, parent):
        """按文档流顺序遍历段落与表格"""
        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            parent_elm = parent
        
        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)
    
    def _parse_section_title(self, para: Paragraph) -> Optional[Dict]:
        """
        解析段落是否为章节标题
        
        支持两种情况：
        1. 文本以章节编号开头（如 "1.4.1 国产数据库技术要求"）
        2. 段落使用 Word 标题样式（如 Heading 1, Heading 2），自动生成编号
        
        Returns:
            {'number': '1.4.1', 'title': '国产数据库技术要求', 'level': 3} 或 None
        """
        text = para.text.strip()
        if not text:
            return None
        
        # 获取段落的样式和大纲级别
        style_name = para.style.name if para.style else ''
        outline_level = self._get_paragraph_outline_level(para)
        
        # 情况1：尝试匹配文本中的章节编号
        for pattern in self.SECTION_NUMBER_PATTERNS:
            match = re.match(pattern, text)
            if match:
                number = match.group(1)
                title = match.group(2).strip() if len(match.groups()) > 1 else ''
                
                # 计算层级
                if '.' in number:
                    level = len(number.split('.'))
                elif number.isdigit():
                    level = 1
                else:
                    # 中文数字
                    level = outline_level or 1
                
                return {
                    'number': number,
                    'title': title,
                    'level': level
                }
        
        # 情况2：如果是标题样式（Heading 1/2/3 或 标题 1/2/3），自动生成编号
        heading_level = self._get_heading_level_from_style(style_name)
        if heading_level is None and outline_level:
            heading_level = outline_level
        
        if heading_level and heading_level <= 9:
            # 更新计数器：当前级别+1，后续级别清零
            self.heading_counters[heading_level - 1] += 1
            for i in range(heading_level, 10):
                self.heading_counters[i] = 0
            
            # 生成编号（如 1.4.1）
            number_parts = [str(self.heading_counters[i]) for i in range(heading_level) if self.heading_counters[i] > 0]
            if number_parts:
                number = '.'.join(number_parts)
                
                return {
                    'number': number,
                    'title': text,
                    'level': heading_level
                }
        
        # 如果有大纲级别但没匹配到编号模式，尝试从文本提取（兼容旧逻辑）
        if outline_level:
            # 检查是否以章节编号格式开头（必须包含点号，如 1.2、1.2.3）
            # 不匹配 "1. xxx" 这种列表格式
            num_match = re.match(r'^(\d+\.\d+(?:\.\d+)*)\s*(.+)', text)
            if num_match:
                return {
                    'number': num_match.group(1).rstrip('.'),
                    'title': num_match.group(2).strip(),
                    'level': outline_level
                }
        
        return None
    
    def _get_heading_level_from_style(self, style_name: str) -> Optional[int]:
        """
        从样式名称获取标题级别
        
        支持：
        - Heading 1, Heading 2, Heading 3...
        - 标题 1, 标题 2, 标题 3...
        - TOC Heading 等
        """
        if not style_name:
            return None
        
        # 匹配 "Heading X" 或 "标题 X"
        patterns = [
            r'^Heading\s*(\d+)$',
            r'^标题\s*(\d+)$',
            r'^Title\s*(\d+)$',
        ]
        
        for pattern in patterns:
            match = re.match(pattern, style_name, re.IGNORECASE)
            if match:
                return int(match.group(1))
        
        return None
    
    def _get_paragraph_outline_level(self, para: Paragraph) -> Optional[int]:
        """获取段落的大纲级别"""
        try:
            p_pr = para._element.pPr
            if p_pr is not None:
                outline_lvl = p_pr.find(qn("w:outlineLvl"))
                if outline_lvl is not None:
                    return int(outline_lvl.get(qn("w:val"))) + 1
        except Exception:
            pass
        
        # 从样式名推断
        style_name = getattr(para.style, "name", "") if para.style else ""
        if any(h in style_name for h in ["Heading", "标题", "Title"]):
            match = re.search(r"(\d+)", style_name)
            return int(match.group(1)) if match else 1
        
        return None
    
    def _get_parent_number(self, section_number: str) -> Optional[str]:
        """获取父章节编号"""
        if '.' in section_number:
            parts = section_number.rsplit('.', 1)
            return parts[0]
        return None
    
    def _parse_requirement_item(self, text: str) -> Optional[Dict]:
        """
        解析需求项（编号列表形式）
        
        匹配格式：
        - 1. 数据库软件需通过...
        - (1) 承诺应答产品...
        - ① 支持...
        """
        patterns = [
            r'^(\d+)[\.、\)]\s*(.+)',  # 1. 或 1、或 1)
            r'^\((\d+)\)\s*(.+)',      # (1)
            r'^([①②③④⑤⑥⑦⑧⑨⑩])\s*(.+)',  # 圆圈数字
            r'^[•·]\s*(.+)',            # 项目符号
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                groups = match.groups()
                if len(groups) == 2:
                    return {
                        'index': groups[0],
                        'text': groups[1],
                        'type': 'list'
                    }
                elif len(groups) == 1:
                    return {
                        'index': '',
                        'text': groups[0],
                        'type': 'list'
                    }
        
        return None
    
    def _parse_table(self, table: Table) -> Optional[Dict]:
        """
        解析表格
        
        Returns:
            {
                'type': 'requirement_table' | 'data_table',
                'headers': ['序号', '技术要求', '技术规格'],
                'rows': [{'序号': '1', '技术要求': 'SQL语法', '技术规格': '支持SQL2003标准'}]
            }
        """
        if not table.rows:
            return None
        
        # 获取表头
        headers = []
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell_text = self._get_cell_text(cell)
            headers.append(cell_text)
        
        # 判断是否为需求表格
        is_requirement_table = any(
            keyword in ''.join(headers) 
            for keyword in self.TABLE_HEADER_KEYWORDS
        )
        
        # 解析数据行
        rows = []
        for row_idx, row in enumerate(table.rows[1:], start=1):
            row_data = {}
            for col_idx, cell in enumerate(row.cells):
                header = headers[col_idx] if col_idx < len(headers) else f'col_{col_idx}'
                row_data[header] = self._get_cell_text(cell)
                row_data['index'] = row_idx
            rows.append(row_data)
        
        return {
            'type': 'requirement_table' if is_requirement_table else 'data_table',
            'headers': headers,
            'rows': rows
        }
    
    def _get_cell_text(self, cell) -> str:
        """获取单元格文本"""
        texts = []
        for para in cell.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        return '\n'.join(texts)
    
    def _extract_requirements_from_table(self, table_data: Dict) -> List[Dict]:
        """
        从表格数据中提取技术要求
        
        支持多种表格格式：
        1. 标准格式：序号 | 技术要求 | 技术规格
        2. 简单格式：序号 | 要求内容
        3. 参数格式：参数名称 | 参数值/规格
        
        Args:
            table_data: 表格数据 {'headers': [...], 'rows': [...]}
        
        Returns:
            需求列表 [{'index': 1, 'text': '技术要求内容', 'spec': '规格要求', 'type': 'table'}]
        """
        requirements = []
        headers = table_data.get('headers', [])
        rows = table_data.get('rows', [])
        
        if not rows:
            return []
        
        # 识别列映射关系
        header_lower = [h.strip().lower() for h in headers]
        header_map = {
            'index': None,      # 序号列
            'requirement': None,  # 技术要求列
            'spec': None         # 技术规格列
        }
        
        # 识别序号列
        for i, h in enumerate(headers):
            h_clean = h.strip()
            if h_clean in ['序号', '编号', '项', 'No', 'No.', '#']:
                header_map['index'] = i
                break
        
        # 识别技术要求列
        req_keywords = ['技术要求', '功能要求', '要求', '功能', '项目', '名称', '技术项']
        for kw in req_keywords:
            for i, h in enumerate(headers):
                if kw in h:
                    header_map['requirement'] = i
                    break
            if header_map['requirement'] is not None:
                break
        
        # 识别技术规格列
        spec_keywords = ['技术规格', '规格', '规格要求', '参数', '参数值', '技术参数', '指标', '指标要求']
        for kw in spec_keywords:
            for i, h in enumerate(headers):
                if kw in h:
                    header_map['spec'] = i
                    break
            if header_map['spec'] is not None:
                break
        
        # 如果没有找到明确的要求列，使用智能推断
        if header_map['requirement'] is None:
            # 尝试找第二列（通常序号后面是要求）
            if len(headers) >= 2 and header_map['index'] == 0:
                header_map['requirement'] = 1
            elif len(headers) >= 1:
                header_map['requirement'] = 0
        
        if header_map['spec'] is None and len(headers) >= 3:
            # 尝试找第三列作为规格
            if header_map['requirement'] is not None:
                for i in range(len(headers)):
                    if i != header_map['index'] and i != header_map['requirement']:
                        header_map['spec'] = i
                        break
        
        # 提取每一行的需求
        for row_idx, row in enumerate(rows):
            # 获取序号
            if header_map['index'] is not None:
                idx_key = headers[header_map['index']]
                index_val = row.get(idx_key, '') or row.get('index', row_idx + 1)
            else:
                index_val = row.get('序号', '') or row.get('index', row_idx + 1)
            
            # 清理序号（可能是带有箭头符号的）
            if isinstance(index_val, str):
                index_val = index_val.replace('↵', '').replace('←', '').strip()
            
            # 获取技术要求
            req_text = ''
            if header_map['requirement'] is not None:
                req_key = headers[header_map['requirement']]
                req_text = row.get(req_key, '')
            
            if not req_text:
                # 尝试其他可能的键
                for key in ['技术要求', '功能要求', '要求', '功能', '项目']:
                    if key in row and row[key]:
                        req_text = row[key]
                        break
            
            # 获取技术规格
            spec_text = ''
            if header_map['spec'] is not None:
                spec_key = headers[header_map['spec']]
                spec_text = row.get(spec_key, '')
            
            if not spec_text:
                # 尝试其他可能的键
                for key in ['技术规格', '规格', '规格要求', '参数', '参数值']:
                    if key in row and row[key]:
                        spec_text = row[key]
                        break
            
            # 清理文本中的特殊字符
            if req_text:
                req_text = req_text.replace('↵', '\n').replace('←', '').strip()
            if spec_text:
                spec_text = spec_text.replace('↵', '\n').replace('←', '').strip()
            
            # 组合成需求项（技术要求 + 技术规格 作为完整的要求描述）
            if req_text or spec_text:
                requirements.append({
                    'index': index_val,
                    'text': req_text,
                    'spec': spec_text,
                    'type': 'table',
                    'row_data': row
                })
        
        return requirements
    
    def get_section_by_number(self, section_number: str) -> Optional[Dict]:
        """
        根据章节编号获取章节内容
        
        Args:
            section_number: 章节编号，如 '1.4.1'
        """
        if not self.section_index:
            self.parse_document_structure()
        
        # 尝试直接匹配
        if section_number in self.section_index:
            return self.section_index[section_number]
        
        # 尝试模糊匹配（处理前导零等情况）
        normalized = self._normalize_section_number(section_number)
        for key in self.section_index:
            if self._normalize_section_number(key) == normalized:
                return self.section_index[key]
        
        return None
    
    def _normalize_section_number(self, number: str) -> str:
        """标准化章节编号"""
        # 去除前导零
        parts = number.split('.')
        normalized_parts = []
        for part in parts:
            try:
                normalized_parts.append(str(int(part)))
            except ValueError:
                normalized_parts.append(part)
        return '.'.join(normalized_parts)
    
    def get_sections_by_numbers(self, section_numbers: List[str]) -> Dict[str, Dict]:
        """
        批量获取多个章节
        
        Args:
            section_numbers: 章节编号列表，如 ['1.4.1', '1.4.2']
        """
        results = {}
        for num in section_numbers:
            section = self.get_section_by_number(num)
            if section:
                results[num] = section
        return results
    
    def get_all_requirements_from_section(self, section_number: str, include_content_parse: bool = True) -> List[Dict]:
        """
        获取指定章节的所有技术要求
        
        如果章节没有明确的requirements列表，会尝试从content中解析
        
        Args:
            section_number: 章节编号
            include_content_parse: 是否从content中解析需求（当requirements为空时）
        
        Returns:
            [
                {'index': 1, 'text': '...', 'spec': '...', 'type': 'table', 'section': '1.4.1'},
                {'index': 2, 'text': '...', 'type': 'list', 'section': '1.4.1'},
            ]
        """
        section = self.get_section_by_number(section_number)
        if not section:
            logger.warning(f"章节 {section_number} 未找到")
            return []
        
        requirements = []
        
        logger.info(f"章节 {section_number} 数据: requirements={len(section.get('requirements', []))}, tables={len(section.get('tables', []))}, content长度={len(section.get('content', ''))}")
        
        # 首先获取已解析的requirements（可能来自表格或列表）
        for req in section.get('requirements', []):
            req_copy = req.copy()
            req_copy['section'] = section_number
            req_copy['section_title'] = section.get('title', '')
            requirements.append(req_copy)
        
        logger.info(f"章节 {section_number} 从requirements获取: {len(requirements)} 条")
        
        # 如果没有明确的requirements，尝试从content中解析
        if not requirements and include_content_parse:
            content = section.get('content', '')
            if content:
                logger.info(f"章节 {section_number} 尝试从content解析，content前200字符: {content[:200]}...")
                parsed_reqs = self._parse_requirements_from_content(content)
                logger.info(f"章节 {section_number} 从content解析出: {len(parsed_reqs)} 条")
                for req in parsed_reqs:
                    req['section'] = section_number
                    req['section_title'] = section.get('title', '')
                    requirements.append(req)
        
        # 如果还没有，尝试从表格中重新提取（以防parse时漏掉）
        if not requirements:
            logger.info(f"章节 {section_number} 尝试从表格重新提取，共 {len(section.get('tables', []))} 个表格")
            for table_data in section.get('tables', []):
                table_reqs = self._extract_requirements_from_table(table_data)
                for req in table_reqs:
                    req['section'] = section_number
                    req['section_title'] = section.get('title', '')
                    requirements.append(req)
        
        # 如果仍然没有，把整个章节内容作为一个需求
        if not requirements:
            content = section.get('content', '') or section.get('raw_title', '') or section.get('title', '')
            logger.info(f"章节 {section_number} 使用整体content作为需求，content={content[:100] if content else 'None'}...")
            if content:
                requirements.append({
                    'index': 1,
                    'text': content,
                    'spec': '',
                    'type': 'content',
                    'section': section_number,
                    'section_title': section.get('title', '')
                })
            else:
                # 即使没有内容，也返回一个空需求，标记该章节被处理过
                logger.warning(f"章节 {section_number} 没有任何内容！")
                requirements.append({
                    'index': 1,
                    'text': f"章节 {section_number} {section.get('title', '')} (无具体内容)",
                    'spec': '',
                    'type': 'empty',
                    'section': section_number,
                    'section_title': section.get('title', '')
                })
        
        logger.info(f"章节 {section_number} 提取到 {len(requirements)} 条技术要求")
        return requirements
    
    def _parse_requirements_from_content(self, content: str) -> List[Dict]:
        """
        从章节内容中解析出具体的技术要求
        
        支持格式：
        - 1. 要求内容
        - (1) 要求内容
        - ① 要求内容
        - • 要求内容
        - 1) 要求内容
        - 无编号的段落文本（按换行分隔，每段作为一条需求）
        """
        if not content:
            return []
        
        requirements = []
        lines = content.split('\n')
        
        # 需求项的正则模式
        patterns = [
            r'^(\d+)[\.、\)]\s*(.+)',      # 1. 或 1、或 1)
            r'^\((\d+)\)\s*(.+)',           # (1)
            r'^([①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮])\s*(.+)',  # 圆圈数字
            r'^[•·▪➢►]\s*(.+)',             # 项目符号
            r'^[（\(](\d+)[）\)]\s*(.+)',   # 中文括号
        ]
        
        current_index = 0
        current_text = []
        has_numbered_items = False  # 标记是否有编号格式的项
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            matched = False
            for pattern in patterns:
                match = re.match(pattern, line)
                if match:
                    has_numbered_items = True
                    # 保存上一个需求
                    if current_text:
                        requirements.append({
                            'index': current_index,
                            'text': '\n'.join(current_text),
                            'spec': '',
                            'type': 'list'
                        })
                    
                    # 开始新需求
                    groups = match.groups()
                    if len(groups) == 2:
                        idx_str = groups[0]
                        # 处理圆圈数字
                        if idx_str in '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮':
                            current_index = '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮'.index(idx_str) + 1
                        else:
                            try:
                                current_index = int(idx_str)
                            except ValueError:
                                current_index = len(requirements) + 1
                        current_text = [groups[1]]
                    else:
                        current_index = len(requirements) + 1
                        current_text = [groups[0]]
                    
                    matched = True
                    break
            
            if not matched and current_text:
                # 追加到当前需求
                current_text.append(line)
        
        # 保存最后一个需求
        if current_text:
            requirements.append({
                'index': current_index,
                'text': '\n'.join(current_text),
                'spec': '',
                'type': 'list'
            })
        
        # 【关键改进】如果没有发现编号格式的项，则按换行将每段非空内容作为单独的需求
        if not has_numbered_items and not requirements:
            # 过滤掉开头的介绍性语句（如"...需要满足以下要求："）
            non_empty_lines = [line.strip() for line in lines if line.strip()]
            
            # 跳过第一行如果它是介绍性语句
            start_idx = 0
            if non_empty_lines and ('以下要求' in non_empty_lines[0] or 
                                     '如下要求' in non_empty_lines[0] or
                                     '具体要求' in non_empty_lines[0] or
                                     non_empty_lines[0].endswith('：') or
                                     non_empty_lines[0].endswith(':')):
                start_idx = 1
            
            for idx, line in enumerate(non_empty_lines[start_idx:], start=1):
                # 跳过太短的行（可能是残缺内容）
                if len(line) < 10:
                    continue
                requirements.append({
                    'index': idx,
                    'text': line,
                    'spec': '',
                    'type': 'paragraph'
                })
        
        return requirements


class UserInstructionParser:
    """用户指令解析器 - 解析用户的作答指令"""
    
    # 指令模式
    INSTRUCTION_PATTERNS = [
        # 针对文档中的1.4.1,1.4.2作答
        r'针对.*?(?:文档|文件)?.*?(?:中的|里的|的)?([0-9\.,、\s]+)作答',
        # 回答1.4.1和1.4.2
        r'回答\s*([0-9\.,、和与\s]+)',
        # 对1.4.1,1.4.2进行作答/回复
        r'对\s*([0-9\.,、和与\s]+)\s*(?:进行)?(?:作答|回复|回答)',
        # 1.4.1,1.4.2 章节
        r'([0-9]+(?:\.[0-9]+)+(?:\s*[,、和与]\s*[0-9]+(?:\.[0-9]+)+)*)\s*(?:章节|部分)?',
        # 章节1.4.1,1.4.2
        r'章节\s*([0-9\.,、和与\s]+)',
    ]
    
    def __init__(self):
        pass
    
    def parse_instruction(self, instruction: str) -> Dict[str, Any]:
        """
        解析用户指令
        
        Args:
            instruction: 用户输入的指令，如 "针对文档中的1.4.1,1.4.2作答"
        
        Returns:
            {
                'action': 'answer_sections',
                'section_numbers': ['1.4.1', '1.4.2'],
                'original_instruction': '...',
                'parsed': True
            }
        """
        result = {
            'action': None,
            'section_numbers': [],
            'document_name': None,
            'original_instruction': instruction,
            'parsed': False
        }
        
        # 尝试提取文档名
        doc_match = re.search(r'["""]([^"""]+)["""]|文档[《]([^》]+)[》]', instruction)
        if doc_match:
            result['document_name'] = doc_match.group(1) or doc_match.group(2)
        
        # 尝试匹配各种指令模式
        for pattern in self.INSTRUCTION_PATTERNS:
            match = re.search(pattern, instruction)
            if match:
                numbers_str = match.group(1)
                section_numbers = self._extract_section_numbers(numbers_str)
                if section_numbers:
                    result['action'] = 'answer_sections'
                    result['section_numbers'] = section_numbers
                    result['parsed'] = True
                    break
        
        # 如果没有匹配到模式，尝试直接提取章节编号
        if not result['parsed']:
            section_numbers = self._extract_section_numbers(instruction)
            if section_numbers:
                result['action'] = 'answer_sections'
                result['section_numbers'] = section_numbers
                result['parsed'] = True
        
        return result
    
    def _extract_section_numbers(self, text: str) -> List[str]:
        """
        从文本中提取所有章节编号
        
        支持格式：
        - 1.4.1,1.4.2
        - 1.4.1、1.4.2
        - 1.4.1和1.4.2
        - 1.4.1 1.4.2
        """
        # 匹配章节编号模式：数字.数字.数字...
        pattern = r'\d+(?:\.\d+)+'
        matches = re.findall(pattern, text)
        
        # 去重并保持顺序
        seen = set()
        unique_numbers = []
        for num in matches:
            normalized = self._normalize_number(num)
            if normalized not in seen:
                seen.add(normalized)
                unique_numbers.append(normalized)
        
        return unique_numbers
    
    def _normalize_number(self, number: str) -> str:
        """标准化章节编号"""
        # 去除末尾的点和空格
        number = number.strip().rstrip('.')
        # 去除前导零
        parts = number.split('.')
        normalized = []
        for part in parts:
            try:
                normalized.append(str(int(part)))
            except ValueError:
                normalized.append(part)
        return '.'.join(normalized)


class BidResponseGenerator:
    """招标应答生成器 - 整合解析和匹配，生成应答"""
    
    def __init__(self, document_parser: BidDocumentParser = None):
        self.document_parser = document_parser or BidDocumentParser()
        self.instruction_parser = UserInstructionParser()
    
    def process_instruction(self, instruction: str, file_path: str = None, doc_id: int = None) -> Dict:
        """
        处理用户指令，返回需要作答的章节和技术要求
        
        Args:
            instruction: 用户指令
            file_path: 文档文件路径（可选，用于直接解析文件）
            doc_id: 文档ID（可选，用于从数据库查询已解析的结构）
        
        Returns:
            {
                'success': True,
                'sections': {
                    '1.4.1': {
                        'title': '国产数据库技术要求',
                        'requirements': [...],
                        'tables': [...]
                    }
                },
                'total_requirements': 15,
                'instruction_info': {...}
            }
        """
        # 1. 解析用户指令
        instruction_info = self.instruction_parser.parse_instruction(instruction)
        
        if not instruction_info['parsed']:
            return {
                'success': False,
                'error': '无法解析指令，请使用格式如："针对文档中的1.4.1,1.4.2作答"',
                'instruction_info': instruction_info
            }
        
        section_numbers = instruction_info['section_numbers']
        
        if not section_numbers:
            return {
                'success': False,
                'error': '未找到有效的章节编号',
                'instruction_info': instruction_info
            }
        
        # 2. 获取章节内容
        sections = {}
        total_requirements = 0
        
        if file_path:
            # 从文件解析
            self.document_parser.load_document(file_path)
            self.document_parser.parse_document_structure()
            
            for num in section_numbers:
                section = self.document_parser.get_section_by_number(num)
                if section:
                    sections[num] = section
                    total_requirements += len(section.get('requirements', []))
        
        elif doc_id:
            # 从数据库查询已解析的章节
            for num in section_numbers:
                section = self._get_section_from_db(doc_id, num)
                if section:
                    sections[num] = section
                    total_requirements += len(section.get('requirements', []))
        
        if not sections:
            return {
                'success': False,
                'error': f'未找到指定的章节: {", ".join(section_numbers)}',
                'instruction_info': instruction_info
            }
        
        return {
            'success': True,
            'sections': sections,
            'total_requirements': total_requirements,
            'instruction_info': instruction_info
        }
    
    def _get_section_from_db(self, doc_id: int, section_number: str) -> Optional[Dict]:
        """从数据库获取已解析的章节"""
        # 尝试按标题匹配
        sql = """
            SELECT c.id, c.title, c.content, c.level, c.parent_id
            FROM chapters c
            WHERE c.document_id = %s 
            AND (c.title LIKE %s OR c.title LIKE %s)
        """
        pattern1 = f'{section_number}%'
        pattern2 = f'%{section_number}%'
        
        result = fetch_one(sql, (doc_id, pattern1, pattern2))
        
        if result:
            return {
                'number': section_number,
                'title': result['title'],
                'content': result['content'],
                'level': result['level'],
                'requirements': self._extract_requirements_from_content(result['content']),
                'tables': [],
                'db_id': result['id']
            }
        
        return None
    
    def _extract_requirements_from_content(self, content: str) -> List[Dict]:
        """从内容中提取需求项"""
        if not content:
            return []
        
        requirements = []
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检查是否是需求项
            req_patterns = [
                r'^(\d+)[\.、\)]\s*(.+)',
                r'^\((\d+)\)\s*(.+)',
            ]
            
            for pattern in req_patterns:
                match = re.match(pattern, line)
                if match:
                    requirements.append({
                        'index': match.group(1),
                        'text': match.group(2),
                        'type': 'list'
                    })
                    break
        
        # 检查是否有表格标记
        if '[表格]' in content:
            table_reqs = self._parse_table_from_text(content)
            requirements.extend(table_reqs)
        
        return requirements
    
    def _parse_table_from_text(self, content: str) -> List[Dict]:
        """从文本中解析表格（WordParser存储的格式）"""
        requirements = []
        
        # 查找表格内容
        table_match = re.search(r'\[表格\](.*?)\[/表格\]', content, re.DOTALL)
        if not table_match:
            return []
        
        table_content = table_match.group(1).strip()
        lines = table_content.split('\n')
        
        # 解析表头
        headers = []
        data_rows = []
        
        for line in lines:
            line = line.strip()
            if not line or not line.startswith('|'):
                continue
            
            cells = [c.strip() for c in line.split('|') if c.strip()]
            
            if not headers:
                headers = cells
            else:
                data_rows.append(cells)
        
        # 转换为需求项
        for idx, row in enumerate(data_rows):
            row_dict = {}
            for i, cell in enumerate(row):
                if i < len(headers):
                    row_dict[headers[i]] = cell
            
            req_text = row_dict.get('技术要求') or row_dict.get('功能要求') or row_dict.get('要求', '')
            if req_text:
                requirements.append({
                    'index': row_dict.get('序号', str(idx + 1)),
                    'text': req_text,
                    'spec': row_dict.get('技术规格') or row_dict.get('规格', ''),
                    'type': 'table',
                    'row_data': row_dict
                })
        
        return requirements


# ==================== Phase 3: 三级匹配策略 ====================

class BidAnswerMatcher:
    """招标作答匹配器 - 实现三级匹配策略"""
    
    # 相似度阈值
    EXACT_MATCH_THRESHOLD = 0.90
    FUZZY_MATCH_THRESHOLD = 0.60
    SEMANTIC_MATCH_THRESHOLD = 0.50
    
    def __init__(self, llm_service=None, web_search_service=None):
        """
        初始化匹配器
        
        Args:
            llm_service: LLM服务实例
            web_search_service: 网络搜索服务实例
        """
        self.llm_service = llm_service
        self.web_search_service = web_search_service
    
    def match_requirement(self, requirement_text: str, username: str, 
                          knowledge_doc_ids: List[int] = None,
                          enable_web_search: bool = True) -> Dict:
        """
        对单个技术要求进行三级匹配
        
        匹配优先级：
        1. 精确匹配（在知识库文档中找到完全或高度相似的内容）
        2. 语义匹配（使用LLM进行语义相似度分析）
        3. 网络搜索（从互联网获取答案）
        4. LLM生成（直接使用LLM生成答案）
        
        Args:
            requirement_text: 技术要求文本
            username: 用户名
            knowledge_doc_ids: 知识库文档ID列表
            enable_web_search: 是否启用网络搜索
        
        Returns:
            {
                'answer': '答案内容',
                'match_type': 'exact/semantic/web/llm_generated/none',
                'confidence': 0.0-1.0,
                'source': {...}
            }
        """
        result = {
            'answer': '',
            'match_type': 'none',
            'confidence': 0,
            'source': None
        }
        
        if not requirement_text or not requirement_text.strip():
            result['answer'] = '技术要求内容为空'
            return result
        
        # 1. 精确匹配
        exact_match = self._exact_match(requirement_text, username, knowledge_doc_ids)
        if exact_match:
            return exact_match
        
        # 2. 语义匹配
        if self.llm_service:
            semantic_match = self._semantic_match(requirement_text, username, knowledge_doc_ids)
            if semantic_match:
                return semantic_match
        
        # 3. 网络搜索
        if enable_web_search and self.web_search_service:
            web_result = self._web_search_match(requirement_text)
            if web_result:
                return web_result
        
        # 4. LLM直接生成
        if self.llm_service:
            llm_result = self._llm_generate(requirement_text)
            if llm_result:
                return llm_result
        
        # 所有方法都失败
        result['answer'] = '未能找到相关答案，请手动补充。'
        return result
    
    def _exact_match(self, requirement_text: str, username: str, 
                     doc_ids: List[int] = None) -> Optional[Dict]:
        """精确匹配 - 在知识库中查找相似内容"""
        # 构建查询条件
        conditions = ["dpr.username = %s", "dpr.status = 'completed'"]
        params = [username]
        
        if doc_ids:
            placeholders = ','.join(['%s'] * len(doc_ids))
            conditions.append(f"c.document_id IN ({placeholders})")
            params.extend(doc_ids)
        
        # 提取关键词用于搜索
        keywords = self._extract_keywords(requirement_text)
        if not keywords:
            return None
        
        # 构建搜索条件
        like_conditions = []
        for kw in keywords[:5]:  # 取前5个关键词
            like_conditions.append("(c.title LIKE %s OR c.content LIKE %s)")
            params.extend([f'%{kw}%', f'%{kw}%'])
        
        if like_conditions:
            conditions.append(f"({' OR '.join(like_conditions)})")
        
        sql = f"""
            SELECT c.id as chapter_id, c.document_id, c.title as chapter_title, 
                   c.content, c.level, c.parent_id,
                   dpr.filename
            FROM chapters c
            JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            WHERE {' AND '.join(conditions)}
            LIMIT 30
        """
        
        try:
            candidates = fetch_all(sql, params)
        except Exception as e:
            logger.error(f"精确匹配查询失败: {e}")
            return None
        
        if not candidates:
            return None
        
        # 计算相似度并排序
        best_match = None
        best_similarity = 0
        
        cleaned_req = self._clean_text(requirement_text)
        
        for candidate in candidates:
            content = candidate.get('content', '') or ''
            title = candidate.get('chapter_title', '') or ''
            
            # 计算标题相似度
            title_sim = SequenceMatcher(None, 
                cleaned_req[:100], 
                self._clean_text(title)
            ).ratio()
            
            # 计算内容相似度
            content_sim = SequenceMatcher(None,
                cleaned_req[:300],
                self._clean_text(content)[:300]
            ).ratio()
            
            # 综合相似度
            similarity = max(title_sim * 1.2, content_sim)  # 标题匹配加权
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_match = candidate
        
        if best_match and best_similarity >= self.FUZZY_MATCH_THRESHOLD:
            # 获取章节路径
            path = self._get_chapter_path(best_match['chapter_id'])
            
            content = best_match.get('content', '')
            
            # 从内容中提取图片（优先），如果没有则获取章节关联的图片
            images = self._get_images_from_content(content)
            if not images:
                images = self._get_chapter_images(best_match['chapter_id'])
            
            return {
                'answer': content,
                'match_type': 'exact' if best_similarity >= self.EXACT_MATCH_THRESHOLD else 'semantic',
                'confidence': min(best_similarity, 1.0),
                'source': {
                    'type': 'document',
                    'filename': best_match.get('filename', ''),
                    'chapter_id': best_match.get('chapter_id'),
                    'chapter_title': best_match.get('chapter_title', ''),
                    'path': path,
                    'images': images
                }
            }
        
        return None
    
    def _semantic_match(self, requirement_text: str, username: str,
                        doc_ids: List[int] = None) -> Optional[Dict]:
        """语义匹配 - 使用LLM分析语义相似度"""
        if not self.llm_service:
            return None
        
        # 获取候选章节
        conditions = ["dpr.username = %s", "dpr.status = 'completed'"]
        params = [username]
        
        if doc_ids:
            placeholders = ','.join(['%s'] * len(doc_ids))
            conditions.append(f"c.document_id IN ({placeholders})")
            params.extend(doc_ids)
        
        sql = f"""
            SELECT c.id as chapter_id, c.document_id, c.title as chapter_title, 
                   c.content, c.level,
                   dpr.filename
            FROM chapters c
            JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            WHERE {' AND '.join(conditions)}
            AND c.content IS NOT NULL AND c.content != ''
            LIMIT 40
        """
        
        try:
            candidates = fetch_all(sql, params)
        except Exception as e:
            logger.error(f"语义匹配查询失败: {e}")
            return None
        
        if not candidates:
            return None
        
        # 构建LLM提示
        candidate_texts = []
        for i, c in enumerate(candidates):
            title = c.get('chapter_title', '(无标题)')
            text = (c.get('content', '') or '')[:400]
            candidate_texts.append(f"[{i+1}] 标题: {title}\n内容: {text}")
        
        candidates_str = '\n\n'.join(candidate_texts[:20])  # 限制数量
        
        prompt = f"""请分析以下技术要求与候选文档内容的相关性，找出最能回答该技术要求的内容。

技术要求：
{requirement_text}

候选文档内容：
{candidates_str}

请返回JSON格式结果：
{{"best_match_index": 序号（1-{min(len(candidates), 20)}，无相关内容返回0）, "relevance_score": 0-1相关性分数, "answer_summary": "基于匹配内容的简要回答"}}

只返回JSON，不要其他内容。"""

        try:
            result = self.llm_service.chat_completion([
                {'role': 'system', 'content': '你是一个专业的技术文档匹配助手，负责分析技术要求与文档内容的相关性，并生成简洁的应答。'},
                {'role': 'user', 'content': prompt}
            ])
            
            response_text = result.get('content', '')
            json_match = re.search(r'\{[^}]+\}', response_text, re.DOTALL)
            
            if json_match:
                match_result = json.loads(json_match.group())
                best_idx = match_result.get('best_match_index', 0)
                score = match_result.get('relevance_score', 0)
                answer_summary = match_result.get('answer_summary', '')
                
                if best_idx > 0 and score >= self.SEMANTIC_MATCH_THRESHOLD and best_idx <= len(candidates):
                    best_match = candidates[best_idx - 1]
                    path = self._get_chapter_path(best_match['chapter_id'])
                    
                    # 使用LLM的摘要或原始内容
                    content = best_match.get('content', '')
                    answer = answer_summary if answer_summary else content
                    
                    # 从原始内容中提取图片（答案可能是摘要，不包含图片占位符）
                    images = self._get_images_from_content(content)
                    if not images:
                        images = self._get_chapter_images(best_match['chapter_id'])
                    
                    return {
                        'answer': answer,
                        'match_type': 'semantic',
                        'confidence': min(score, 1.0),
                        'source': {
                            'type': 'document',
                            'filename': best_match.get('filename', ''),
                            'chapter_id': best_match.get('chapter_id'),
                            'chapter_title': best_match.get('chapter_title', ''),
                            'path': path,
                            'images': images
                        }
                    }
        except Exception as e:
            logger.error(f"语义匹配LLM调用失败: {e}")
        
        return None
    
    def _web_search_match(self, requirement_text: str) -> Optional[Dict]:
        """网络搜索匹配"""
        if not self.web_search_service:
            return None
        
        try:
            # 提取搜索关键词
            keywords = self._extract_keywords(requirement_text)
            search_query = ' '.join(keywords[:8])
            
            results = self.web_search_service.search(search_query, num_results=5)
            
            if not results:
                return None
            
            # 如果有LLM，用它整合结果
            if self.llm_service:
                context = '\n\n'.join([
                    f"[{r.get('title', '')}]\n{r.get('snippet', '')}\n来源: {r.get('url', '')}"
                    for r in results
                ])
                
                prompt = f"""根据以下搜索结果，为技术要求生成专业的应答：

技术要求：{requirement_text}

搜索结果：
{context}

请提供简洁专业的应答："""

                try:
                    result = self.llm_service.chat_completion([
                        {'role': 'system', 'content': '你是一个专业的技术顾问，根据搜索结果生成招标技术要求的应答。'},
                        {'role': 'user', 'content': prompt}
                    ])
                    
                    return {
                        'answer': result.get('content', ''),
                        'match_type': 'web',
                        'confidence': 0.5,
                        'source': {
                            'type': 'web',
                            'search_results': results[:3]
                        }
                    }
                except Exception as e:
                    logger.error(f"网络搜索LLM整合失败: {e}")
            
            # 没有LLM则返回搜索摘要
            summary = '\n'.join([f"• {r.get('title', '')}: {r.get('snippet', '')}" for r in results[:3]])
            return {
                'answer': f"根据网络搜索结果：\n{summary}",
                'match_type': 'web',
                'confidence': 0.4,
                'source': {
                    'type': 'web',
                    'search_results': results
                }
            }
            
        except Exception as e:
            logger.error(f"网络搜索失败: {e}")
            return None
    
    def _llm_generate(self, requirement_text: str) -> Optional[Dict]:
        """LLM直接生成答案"""
        if not self.llm_service:
            return None
        
        try:
            prompt = f"""请为以下招标技术要求生成专业的应答内容：

技术要求：
{requirement_text}

应答要求：
1. 内容要专业、准确
2. 回答要简洁明了
3. 如果是需要承诺的要求，说明"我方承诺满足该要求"并补充说明
4. 如果是技术参数要求，说明我方产品的对应参数

请生成应答："""

            result = self.llm_service.chat_completion([
                {'role': 'system', 'content': '你是一个专业的投标技术顾问，擅长为招标文档的技术要求生成专业的应答内容。'},
                {'role': 'user', 'content': prompt}
            ])
            
            return {
                'answer': result.get('content', ''),
                'match_type': 'llm_generated',
                'confidence': 0.3,
                'source': {'type': 'llm'}
            }
        except Exception as e:
            logger.error(f"LLM生成失败: {e}")
            return None
    
    def _extract_keywords(self, text: str) -> List[str]:
        """提取关键词"""
        if not text:
            return []
        
        # 停用词
        stopwords = {'的', '是', '在', '有', '和', '与', '了', '这', '那', '什么', 
                     '怎么', '如何', '为什么', '需要', '要求', '功能', '支持', '应',
                     '能', '可以', '进行', '具有', '提供', '包括', '等', '及'}
        
        words = re.findall(r'[\u4e00-\u9fa5a-zA-Z0-9]+', text)
        keywords = [w for w in words if w not in stopwords and len(w) > 1]
        
        return keywords[:15]
    
    def _clean_text(self, text: str) -> str:
        """清理文本"""
        if not text:
            return ""
        cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)
        return cleaned.lower()
    
    def _get_chapter_path(self, chapter_id: int) -> List[Dict]:
        """获取章节路径"""
        path = []
        current_id = chapter_id
        
        while current_id:
            try:
                sql = "SELECT id, parent_id, level, title FROM chapters WHERE id = %s"
                chapter = fetch_one(sql, (current_id,))
                if chapter:
                    path.insert(0, {
                        'id': chapter['id'],
                        'level': chapter['level'],
                        'title': chapter['title']
                    })
                    current_id = chapter['parent_id']
                else:
                    break
            except Exception:
                break
        
        return path
    
    def _get_chapter_images(self, chapter_id: int) -> List[Dict]:
        """获取章节关联的图片"""
        try:
            sql = """
                SELECT di.id, di.image_url, di.image_path
                FROM document_images di
                JOIN chapter_images ci ON di.id = ci.image_id
                WHERE ci.chapter_id = %s
                ORDER BY ci.position_in_chapter
            """
            images = fetch_all(sql, (chapter_id,))
            return [
                {
                    'id': img['id'],
                    'image_url': img.get('image_url') or img.get('image_path', '')
                }
                for img in (images or [])
            ]
        except Exception as e:
            logger.warning(f"获取章节图片失败: {e}")
            return []
    
    def _get_images_from_content(self, content: str) -> List[Dict]:
        """从内容中提取图片ID并获取图片信息
        
        内容中可能包含 {{IMAGE_ID_xxx}} 格式的占位符，
        需要提取这些ID并从数据库获取对应的图片信息
        """
        if not content:
            return []
        
        # 提取所有图片ID
        image_ids = re.findall(r'\{\{IMAGE_ID_(\d+)\}\}', content)
        if not image_ids:
            return []
        
        # 去重
        image_ids = list(set(image_ids))
        
        try:
            placeholders = ','.join(['%s'] * len(image_ids))
            sql = f"""
                SELECT id, image_url, image_path
                FROM document_images
                WHERE id IN ({placeholders})
            """
            images = fetch_all(sql, image_ids)
            return [
                {
                    'id': img['id'],
                    'image_url': img.get('image_url') or img.get('image_path', '')
                }
                for img in (images or [])
            ]
        except Exception as e:
            logger.warning(f"从内容获取图片失败: {e}")
            return []


# ==================== Phase 4 & 5: 作答结果整合与Word生成 ====================

class BidAnswerGenerator:
    """招标作答生成器 - 整合作答结果并生成Word文档"""
    
    def __init__(self, llm_config_id: int = None):
        """
        初始化生成器
        
        Args:
            llm_config_id: LLM配置ID
        """
        self.llm_config_id = llm_config_id
        self.llm_service = None
        self.web_search_service = None
        self.matcher = None
    
    def _init_services(self):
        """延迟初始化服务"""
        if self.matcher is not None:
            return
        
        try:
            from .llm_service import get_llm_service
            self.llm_service = get_llm_service(self.llm_config_id)
        except Exception as e:
            logger.warning(f"LLM服务初始化失败: {e}")
            self.llm_service = None
        
        try:
            from .web_search import WebSearchService
            self.web_search_service = WebSearchService()
            if not self.web_search_service.config:
                self.web_search_service = None
        except Exception as e:
            logger.warning(f"网络搜索服务初始化失败: {e}")
            self.web_search_service = None
        
        self.matcher = BidAnswerMatcher(self.llm_service, self.web_search_service)
    
    def answer_requirements(self, requirements: List[Dict], username: str,
                           knowledge_doc_ids: List[int] = None,
                           enable_web_search: bool = True,
                           progress_callback=None) -> List[Dict]:
        """
        批量处理技术要求并生成答案
        
        Args:
            requirements: 技术要求列表 [{'section_number': '1.4.1', 'content': '...', ...}]
            username: 用户名
            knowledge_doc_ids: 知识库文档ID列表
            enable_web_search: 是否启用网络搜索
            progress_callback: 进度回调 callback(current, total, result)
        
        Returns:
            作答结果列表
        """
        self._init_services()
        
        results = []
        total = len(requirements)
        
        for i, req in enumerate(requirements):
            try:
                # 提取要求内容
                req_content = req.get('content', '') or req.get('text', '')
                spec = req.get('spec', '')
                
                # 组合完整的要求文本
                full_text = req_content
                if spec:
                    full_text += f"\n技术规格: {spec}"
                
                # 执行匹配
                match_result = self.matcher.match_requirement(
                    full_text,
                    username,
                    knowledge_doc_ids,
                    enable_web_search
                )
                
                # 组合结果
                result = {
                    'section_number': req.get('section_number', ''),
                    'section_title': req.get('section_title', ''),
                    'requirement_index': req.get('index', i + 1),
                    'requirement': req_content,
                    'spec': spec,
                    'answer': match_result.get('answer', ''),
                    'match_type': match_result.get('match_type', 'none'),
                    'confidence': match_result.get('confidence', 0),
                    'source': match_result.get('source')
                }
                
                results.append(result)
                
                if progress_callback:
                    progress_callback(i + 1, total, result)
                    
            except Exception as e:
                logger.error(f"处理需求 {i+1} 失败: {e}")
                results.append({
                    'section_number': req.get('section_number', ''),
                    'section_title': req.get('section_title', ''),
                    'requirement_index': req.get('index', i + 1),
                    'requirement': req.get('content', ''),
                    'spec': req.get('spec', ''),
                    'answer': f'处理失败: {str(e)}',
                    'match_type': 'error',
                    'confidence': 0,
                    'source': None
                })
        
        return results
    
    def export_to_word(self, results: List[Dict], title: str = '招标技术要求应答书',
                       bid_doc_info: Dict = None) -> Tuple[str, str]:
        """
        将作答结果导出为Word文档
        
        Args:
            results: 作答结果列表
            title: 文档标题
            bid_doc_info: 招标文档信息 {'filename': '...', 'doc_id': ...}
        
        Returns:
            (文件路径, 文件名)
        """
        doc = DocumentLoader()
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加文档信息
        info_para = doc.add_paragraph()
        info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        if bid_doc_info:
            info_para.add_run(f"招标文档: {bid_doc_info.get('filename', '')}\n")
        
        doc.add_paragraph()
        
        # 添加统计信息
        total = len(results)
        stats = {
            'exact': sum(1 for r in results if r.get('match_type') == 'exact'),
            'semantic': sum(1 for r in results if r.get('match_type') == 'semantic'),
            'web': sum(1 for r in results if r.get('match_type') == 'web'),
            'llm_generated': sum(1 for r in results if r.get('match_type') == 'llm_generated'),
            'none': sum(1 for r in results if r.get('match_type') in ['none', 'error'])
        }
        
        stats_para = doc.add_paragraph()
        stats_run = stats_para.add_run("应答统计\n")
        stats_run.bold = True
        stats_para.add_run(f"• 总计: {total} 条\n")
        stats_para.add_run(f"• 精确匹配: {stats['exact']} 条\n")
        stats_para.add_run(f"• 语义匹配: {stats['semantic']} 条\n")
        stats_para.add_run(f"• 网络搜索: {stats['web']} 条\n")
        stats_para.add_run(f"• LLM生成: {stats['llm_generated']} 条\n")
        stats_para.add_run(f"• 未匹配: {stats['none']} 条")
        
        doc.add_paragraph()
        doc.add_paragraph('─' * 50)
        
        # 按章节分组
        grouped_results = {}
        for r in results:
            section_num = r.get('section_number', '其他')
            if section_num not in grouped_results:
                grouped_results[section_num] = {
                    'title': r.get('section_title', ''),
                    'items': []
                }
            grouped_results[section_num]['items'].append(r)
        
        # 按章节顺序输出
        sorted_sections = sorted(grouped_results.keys(), 
                                 key=lambda x: [int(p) if p.isdigit() else p for p in x.split('.')])
        
        for section_num in sorted_sections:
            group = grouped_results[section_num]
            
            # 章节标题
            section_heading = doc.add_heading(f"{section_num} {group['title']}", level=1)
            
            # 每个技术要求
            for item in group['items']:
                req_idx = item.get('requirement_index', '')
                requirement = item.get('requirement', '')
                spec = item.get('spec', '')
                answer = item.get('answer', '')
                match_type = item.get('match_type', 'none')
                confidence = item.get('confidence', 0)
                
                # 匹配类型标签
                match_labels = {
                    'exact': '精确匹配',
                    'semantic': '语义匹配',
                    'web': '网络搜索',
                    'llm_generated': 'LLM生成',
                    'none': '未匹配',
                    'error': '处理失败'
                }
                
                # 技术要求标题
                req_heading = doc.add_heading(f"要求 {req_idx}", level=2)
                
                # 技术要求内容
                req_para = doc.add_paragraph()
                req_run = req_para.add_run("【技术要求】")
                req_run.bold = True
                req_para.add_run(f"\n{requirement}")
                if spec:
                    req_para.add_run(f"\n规格: {spec}")
                
                # 匹配信息
                match_para = doc.add_paragraph()
                match_run = match_para.add_run(f"【匹配方式】{match_labels.get(match_type, '未知')} ")
                match_run.bold = True
                match_run.font.color.rgb = RGBColor(102, 126, 234)
                match_para.add_run(f"(置信度: {confidence:.0%})")
                
                # 应答内容
                answer_para = doc.add_paragraph()
                answer_run = answer_para.add_run("【应答内容】")
                answer_run.bold = True
                answer_para.add_run(f"\n{answer}")
                
                # 来源信息
                source = item.get('source')
                if source and source.get('type') == 'document':
                    source_para = doc.add_paragraph()
                    source_run = source_para.add_run("【来源】")
                    source_run.bold = True
                    source_run.font.size = Pt(10)
                    source_para.add_run(f" {source.get('filename', '')} - {source.get('chapter_title', '')}")
                
                doc.add_paragraph()  # 空行分隔
            
            doc.add_paragraph('─' * 30)
        
        # 保存文件
        temp_dir = tempfile.gettempdir()
        filename = f"招标应答_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return filepath, filename
    
    def export_to_word_table_format(self, results: List[Dict], title: str = '招标技术要求应答表',
                                    bid_doc_info: Dict = None) -> Tuple[str, str]:
        """
        将作答结果导出为表格格式的Word文档
        
        Args:
            results: 作答结果列表
            title: 文档标题
            bid_doc_info: 招标文档信息
        
        Returns:
            (文件路径, 文件名)
        """
        doc = DocumentLoader()
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加时间和来源
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        if bid_doc_info:
            info_para.add_run(f"\n招标文档: {bid_doc_info.get('filename', '')}")
        
        doc.add_paragraph()
        
        # 按章节分组创建表格
        grouped_results = {}
        for r in results:
            section_num = r.get('section_number', '其他')
            if section_num not in grouped_results:
                grouped_results[section_num] = {
                    'title': r.get('section_title', ''),
                    'items': []
                }
            grouped_results[section_num]['items'].append(r)
        
        sorted_sections = sorted(grouped_results.keys(),
                                 key=lambda x: [int(p) if p.isdigit() else p for p in x.split('.')])
        
        for section_num in sorted_sections:
            group = grouped_results[section_num]
            
            # 章节标题
            doc.add_heading(f"{section_num} {group['title']}", level=1)
            
            # 创建表格
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # 表头
            header_cells = table.rows[0].cells
            headers = ['序号', '技术要求', '应答内容', '匹配方式']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                header_cells[i].paragraphs[0].runs[0].bold = True
            
            # 数据行
            for item in group['items']:
                row_cells = table.add_row().cells
                
                row_cells[0].text = str(item.get('requirement_index', ''))
                row_cells[1].text = item.get('requirement', '')
                if item.get('spec'):
                    row_cells[1].text += f"\n规格: {item['spec']}"
                row_cells[2].text = item.get('answer', '')
                
                match_labels = {
                    'exact': '精确匹配',
                    'semantic': '语义匹配',
                    'web': '网络搜索',
                    'llm_generated': 'LLM生成',
                    'none': '未匹配',
                    'error': '失败'
                }
                row_cells[3].text = match_labels.get(item.get('match_type', 'none'), '未知')
            
            doc.add_paragraph()
        
        # 保存文件
        temp_dir = tempfile.gettempdir()
        filename = f"招标应答表_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return filepath, filename


# ==================== 便捷函数 ====================

def parse_bid_document(file_path: str) -> Dict:
    """解析招标文档"""
    parser = BidDocumentParser(file_path)
    return parser.parse_document_structure()


def parse_user_instruction(instruction: str) -> Dict:
    """解析用户指令"""
    parser = UserInstructionParser()
    return parser.parse_instruction(instruction)


def process_bid_response(instruction: str, file_path: str = None, doc_id: int = None) -> Dict:
    """处理招标应答请求"""
    generator = BidResponseGenerator()
    return generator.process_instruction(instruction, file_path, doc_id)


def answer_bid_requirements(requirements: List[Dict], username: str,
                            knowledge_doc_ids: List[int] = None,
                            llm_config_id: int = None,
                            enable_web_search: bool = True) -> List[Dict]:
    """
    批量处理招标技术要求并生成答案
    
    Args:
        requirements: 技术要求列表
        username: 用户名
        knowledge_doc_ids: 知识库文档ID列表
        llm_config_id: LLM配置ID
        enable_web_search: 是否启用网络搜索
    
    Returns:
        作答结果列表
    """
    generator = BidAnswerGenerator(llm_config_id)
    return generator.answer_requirements(requirements, username, knowledge_doc_ids, enable_web_search)


def export_bid_answers_to_word(results: List[Dict], title: str = '招标技术要求应答书',
                               bid_doc_info: Dict = None, format_type: str = 'default') -> Tuple[str, str]:
    """
    导出招标作答结果为Word文档
    
    Args:
        results: 作答结果列表
        title: 文档标题
        bid_doc_info: 招标文档信息
        format_type: 格式类型 'default' 或 'table'
    
    Returns:
        (文件路径, 文件名)
    """
    generator = BidAnswerGenerator()
    
    if format_type == 'table':
        return generator.export_to_word_table_format(results, title, bid_doc_info)
    else:
        return generator.export_to_word(results, title, bid_doc_info)
