"""
标书需求提取器模块
专门处理招标文档中的技术要求表格，支持合并单元格识别，
将每个要求整理成标题格式，以Word文档输出下载

功能特点：
1. 解析表格中的技术要求
2. 识别合并单元格（大标题 + 多个小标题）
3. 没有合并单元格则一行一条要求
4. 使用LLM进行智能解析和格式化
5. 输出结构化的Word文档
"""
import os
import re
import json
import tempfile
from datetime import datetime
from typing import Dict, List, Optional, Tuple, Any
from docx import Document as DocumentLoader
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

from config.logging_config import logger


# ==================== 标书应答专家 Prompt ====================

BID_EXPERT_SYSTEM_PROMPT = """你是一位资深的标书应答专家，拥有丰富的投标文档编写经验。

## 你的角色
- 专业的招投标技术顾问
- 精通各类招标文档的技术要求解读
- 善于将复杂的技术要求整理成清晰的结构化格式

## 你的任务
分析招标文档中的技术要求，将其整理成结构化的标题和内容格式。

## 处理规则

### 1. 表格识别规则
- **无合并单元格的表格**：每一行是一条独立的技术要求
- **有合并单元格的表格**：合并单元格的内容是大标题（分类），其下的多行是该分类下的小标题（具体要求）

### 2. 标题层级规则
- 一级标题：来自表格的分类标题（如"安全可靠测评"、"事务处理机制"）
- 二级标题：具体的技术要求条目
- 内容：技术规格说明、响应要求等

### 3. 特殊标记识别
- ★ 表示关键/必须满足的要求
- # 表示重要/参考要求
- 带有"应"、"需"、"必须"等字眼的是强制性要求

### 4. 输出格式
请以JSON格式返回解析结果：
```json
{
    "sections": [
        {
            "level": 1,
            "title": "大标题/分类名称",
            "is_key": true/false,  // 是否关键要求（有★标记）
            "children": [
                {
                    "level": 2,
                    "title": "具体要求标题",
                    "content": "详细的技术规格说明",
                    "is_key": true/false,
                    "requirement_type": "功能/性能/安全/接口/其他"
                }
            ]
        },
        {
            "level": 2,
            "title": "独立的技术要求（无分类）",
            "content": "详细说明",
            "is_key": false,
            "requirement_type": "其他"
        }
    ],
    "summary": {
        "total_requirements": 总条数,
        "key_requirements": 关键要求数,
        "categories": ["分类1", "分类2"]
    }
}
```
"""

BID_EXTRACT_PROMPT_TEMPLATE = """请分析以下招标文档章节中的技术要求，将其整理成结构化的标题格式。

## 章节信息
章节标题：{section_title}
章节编号：{section_number}

## 原始内容
{content}

## 表格数据
{table_data}

## 要求
1. 识别表格中是否有合并单元格：
   - 如果第一列（如"序号"列）有多行合并，说明是分类 + 子项结构
   - 如果没有合并，则每行是独立的要求
   
2. 将技术要求整理成标题格式：
   - 合并单元格的内容作为一级标题（分类）
   - 每行的"技术要求"内容作为二级标题
   - "技术规格"内容作为该标题下的详细说明
   
3. 识别关键要求（带★或#标记的）

请返回JSON格式的结构化结果。"""


class TableMergeAnalyzer:
    """表格合并单元格分析器"""
    
    def __init__(self, table: Table):
        self.table = table
        self.merge_map = {}  # 记录合并单元格的信息
        self._analyze_merge_cells()
    
    def _analyze_merge_cells(self):
        """分析表格中的合并单元格"""
        if not self.table or not self.table.rows:
            return
        
        for row_idx, row in enumerate(self.table.rows):
            for col_idx, cell in enumerate(row.cells):
                # 获取单元格的XML信息
                tc = cell._tc
                
                # 检查垂直合并
                v_merge = tc.find(qn('w:vMerge'))
                if v_merge is not None:
                    merge_val = v_merge.get(qn('w:val'))
                    if merge_val == 'restart':
                        # 这是合并单元格的起始
                        self.merge_map[(row_idx, col_idx)] = {
                            'type': 'vmerge_start',
                            'text': self._get_cell_text(cell)
                        }
                    else:
                        # 这是合并单元格的延续
                        self.merge_map[(row_idx, col_idx)] = {
                            'type': 'vmerge_continue'
                        }
                
                # 检查水平合并
                grid_span = tc.find(qn('w:gridSpan'))
                if grid_span is not None:
                    span_val = grid_span.get(qn('w:val'))
                    if span_val and int(span_val) > 1:
                        self.merge_map[(row_idx, col_idx)] = {
                            'type': 'hmerge',
                            'span': int(span_val),
                            'text': self._get_cell_text(cell)
                        }
    
    def _get_cell_text(self, cell) -> str:
        """获取单元格文本"""
        texts = []
        for para in cell.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        return '\n'.join(texts)
    
    def is_cell_merged_vertically(self, row_idx: int, col_idx: int) -> bool:
        """检查单元格是否是垂直合并的一部分"""
        key = (row_idx, col_idx)
        if key in self.merge_map:
            return self.merge_map[key]['type'] in ['vmerge_start', 'vmerge_continue']
        return False
    
    def is_merge_start(self, row_idx: int, col_idx: int) -> bool:
        """检查是否是合并单元格的起始"""
        key = (row_idx, col_idx)
        return key in self.merge_map and self.merge_map[key]['type'] == 'vmerge_start'
    
    def get_merge_group(self, start_row: int, col_idx: int) -> List[int]:
        """获取从指定行开始的垂直合并组的所有行索引"""
        rows = [start_row]
        for row_idx in range(start_row + 1, len(self.table.rows)):
            key = (row_idx, col_idx)
            if key in self.merge_map and self.merge_map[key]['type'] == 'vmerge_continue':
                rows.append(row_idx)
            else:
                break
        return rows


class BidRequirementExtractor:
    """标书需求提取器"""
    
    def __init__(self, llm_service=None):
        """
        初始化提取器
        
        Args:
            llm_service: LLM服务实例（可选，用于智能解析）
        """
        self.llm_service = llm_service
    
    def extract_from_document(self, file_path: str, section_filter: List[str] = None) -> Dict:
        """
        从文档中提取技术要求
        
        Args:
            file_path: Word文档路径
            section_filter: 要提取的章节编号列表（可选）
        
        Returns:
            {
                'sections': [...],  # 结构化的章节要求
                'summary': {...},   # 统计信息
                'raw_tables': [...]  # 原始表格数据
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        doc = DocumentLoader(file_path)
        result = {
            'sections': [],
            'summary': {
                'total_requirements': 0,
                'key_requirements': 0,
                'categories': []
            },
            'raw_tables': []
        }
        
        # 遍历文档，提取表格
        current_section = None
        
        for element in self._iter_block_items(doc):
            if isinstance(element, Paragraph):
                text = element.text.strip()
                if not text:
                    continue
                
                # 检查是否是章节标题
                section_info = self._parse_section_title(text, element)
                if section_info:
                    # 检查是否需要过滤
                    if section_filter:
                        if not self._match_section_filter(section_info['number'], section_filter):
                            current_section = None
                            continue
                    
                    current_section = {
                        'number': section_info['number'],
                        'title': section_info['title'],
                        'requirements': [],
                        'tables': []
                    }
                    result['sections'].append(current_section)
            
            elif isinstance(element, Table):
                if current_section is not None:
                    # 解析表格中的要求
                    table_requirements = self._extract_requirements_from_table(element)
                    if table_requirements:
                        current_section['requirements'].extend(table_requirements['requirements'])
                        current_section['tables'].append(table_requirements)
                        result['raw_tables'].append(table_requirements)
                        
                        # 更新统计
                        result['summary']['total_requirements'] += len(table_requirements['requirements'])
                        result['summary']['key_requirements'] += sum(
                            1 for r in table_requirements['requirements'] if r.get('is_key')
                        )
        
        # 提取所有分类
        categories = set()
        for section in result['sections']:
            for req in section.get('requirements', []):
                if req.get('category'):
                    categories.add(req['category'])
        result['summary']['categories'] = list(categories)
        
        return result
    
    def _iter_block_items(self, doc):
        """按文档流顺序遍历段落与表格"""
        from docx.document import Document as DocxDocument
        
        parent_elm = doc.element.body
        
        for child in parent_elm.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, doc)
            elif child.tag.endswith("}tbl"):
                yield Table(child, doc)
    
    def _parse_section_title(self, text: str, para: Paragraph) -> Optional[Dict]:
        """解析章节标题"""
        patterns = [
            r'^(\d+(?:\.\d+)*)\s*[\.、\s]?\s*(.+)',  # 1.4.1 xxx
            r'^第([一二三四五六七八九十\d]+)[章节条款]\s*(.+)',  # 第五章 xxx
        ]
        
        for pattern in patterns:
            match = re.match(pattern, text)
            if match:
                return {
                    'number': match.group(1),
                    'title': match.group(2).strip()
                }
        
        return None
    
    def _match_section_filter(self, section_number: str, filter_list: List[str]) -> bool:
        """检查章节编号是否匹配过滤条件"""
        for f in filter_list:
            if section_number == f or section_number.startswith(f + '.'):
                return True
        return False
    
    def _extract_requirements_from_table(self, table: Table) -> Optional[Dict]:
        """
        从表格中提取技术要求
        
        处理逻辑：
        1. 检测是否有合并单元格
        2. 有合并单元格：合并内容为大标题，其下为小标题
        3. 无合并单元格：每行一条要求
        
        Returns:
            {
                'headers': [...],
                'requirements': [
                    {
                        'index': 序号,
                        'category': 分类（大标题），
                        'title': 要求标题,
                        'content': 技术规格说明,
                        'is_key': 是否关键要求,
                        'raw_data': {...}
                    }
                ],
                'has_merge': 是否有合并单元格
            }
        """
        if not table.rows:
            return None
        
        # 获取表头
        headers = []
        header_row = table.rows[0]
        for cell in header_row.cells:
            cell_text = self._get_cell_text(cell)
            headers.append(cell_text)
        
        # 检查是否是技术要求表格
        req_keywords = ['序号', '技术要求', '技术规格', '功能要求', '参数', '指标']
        if not any(kw in ''.join(headers) for kw in req_keywords):
            return None
        
        # 分析合并单元格
        merge_analyzer = TableMergeAnalyzer(table)
        
        # 识别列映射
        col_map = self._identify_columns(headers)
        
        # 提取要求
        requirements = []
        has_merge = False
        current_category = None  # 当前分类（来自合并单元格）
        
        # 跳过表头行，从第二行开始
        for row_idx in range(1, len(table.rows)):
            row = table.rows[row_idx]
            
            # 检查第一列（通常是序号或分类）是否有合并
            first_col_idx = col_map.get('index', 0)
            
            # 检查是否是合并单元格的起始
            if merge_analyzer.is_merge_start(row_idx, first_col_idx):
                has_merge = True
                # 获取合并单元格的文本作为分类
                merge_text = self._get_cell_text(row.cells[first_col_idx])
                current_category = self._clean_category_text(merge_text)
                logger.info(f"检测到合并单元格，分类: {current_category}")
            elif merge_analyzer.is_cell_merged_vertically(row_idx, first_col_idx):
                # 这是合并单元格的延续，保持当前分类
                pass
            else:
                # 不是合并单元格，可能需要重置分类
                # 但如果当前行的第一列有内容，检查是否是新分类
                first_cell_text = self._get_cell_text(row.cells[first_col_idx]) if row.cells else ''
                if first_cell_text and not first_cell_text.isdigit():
                    # 非数字内容可能是分类
                    potential_category = self._clean_category_text(first_cell_text)
                    if potential_category and len(potential_category) > 2:
                        current_category = potential_category
            
            # 提取该行的要求内容
            row_data = {}
            for col_idx, cell in enumerate(row.cells):
                header = headers[col_idx] if col_idx < len(headers) else f'col_{col_idx}'
                row_data[header] = self._get_cell_text(cell)
            
            # 获取具体字段
            req_index = self._get_field_value(row_data, col_map, 'index', row_idx)
            req_title = self._get_field_value(row_data, col_map, 'requirement', '')
            req_spec = self._get_field_value(row_data, col_map, 'spec', '')
            
            # 清理数据
            req_index = str(req_index).replace('↵', '').replace('←', '').strip()
            req_title = req_title.replace('↵', '\n').replace('←', '').strip()
            req_spec = req_spec.replace('↵', '\n').replace('←', '').strip()
            
            # 检查是否是关键要求
            is_key = self._is_key_requirement(req_index, req_title, req_spec)
            
            # 跳过空行或只有序号的行
            if not req_title and not req_spec:
                continue
            
            # 如果标题为空但规格不为空，使用规格作为标题
            if not req_title and req_spec:
                req_title = req_spec[:50] + ('...' if len(req_spec) > 50 else '')
            
            requirements.append({
                'index': req_index,
                'category': current_category,
                'title': req_title,
                'content': req_spec,
                'is_key': is_key,
                'raw_data': row_data
            })
        
        return {
            'headers': headers,
            'requirements': requirements,
            'has_merge': has_merge
        }
    
    def _get_cell_text(self, cell) -> str:
        """获取单元格文本"""
        texts = []
        for para in cell.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        return '\n'.join(texts)
    
    def _identify_columns(self, headers: List[str]) -> Dict[str, int]:
        """识别列映射"""
        col_map = {
            'index': None,       # 序号列
            'requirement': None,  # 技术要求列
            'spec': None         # 技术规格列
        }
        
        for i, h in enumerate(headers):
            h_clean = h.strip().lower()
            
            if h_clean in ['序号', '编号', '项', 'no', 'no.', '#', '类别']:
                if col_map['index'] is None:
                    col_map['index'] = i
            
            if any(kw in h_clean for kw in ['技术要求', '功能要求', '要求', '功能', '项目', '名称']):
                if col_map['requirement'] is None:
                    col_map['requirement'] = i
            
            if any(kw in h_clean for kw in ['技术规格', '规格', '参数', '指标', '说明']):
                if col_map['spec'] is None:
                    col_map['spec'] = i
        
        # 智能推断
        if col_map['index'] is None:
            col_map['index'] = 0
        if col_map['requirement'] is None and len(headers) >= 2:
            col_map['requirement'] = 1
        if col_map['spec'] is None and len(headers) >= 3:
            col_map['spec'] = 2
        
        return col_map
    
    def _get_field_value(self, row_data: Dict, col_map: Dict, field: str, default='') -> str:
        """获取字段值"""
        if field == 'index':
            col_idx = col_map.get('index')
            for key in ['序号', '编号', '项', '类别']:
                if key in row_data:
                    return row_data[key]
        elif field == 'requirement':
            for key in ['技术要求', '功能要求', '要求', '功能', '项目', '名称']:
                if key in row_data:
                    return row_data[key]
        elif field == 'spec':
            for key in ['技术规格', '规格', '规格要求', '参数', '参数值', '指标', '说明']:
                if key in row_data:
                    return row_data[key]
        
        return default
    
    def _clean_category_text(self, text: str) -> str:
        """清理分类文本"""
        if not text:
            return ''
        # 去除特殊符号
        text = text.replace('★', '').replace('#', '').replace('↵', '').replace('←', '').strip()
        # 去除序号
        text = re.sub(r'^\d+[\.、\)\s]*', '', text)
        return text
    
    def _is_key_requirement(self, index: str, title: str, spec: str) -> bool:
        """判断是否是关键要求"""
        combined = f"{index}{title}{spec}"
        return '★' in combined or '▲' in combined
    
    def extract_with_llm(self, content: str, table_data: str, section_info: Dict) -> Dict:
        """
        使用LLM智能提取和结构化要求
        
        Args:
            content: 章节文本内容
            table_data: 表格数据的文本表示
            section_info: 章节信息 {'title': '...', 'number': '...'}
        
        Returns:
            结构化的要求数据
        """
        if not self.llm_service:
            logger.warning("未配置LLM服务，使用基础解析")
            return None
        
        prompt = BID_EXTRACT_PROMPT_TEMPLATE.format(
            section_title=section_info.get('title', ''),
            section_number=section_info.get('number', ''),
            content=content[:3000],  # 限制内容长度
            table_data=table_data[:3000]
        )
        
        try:
            result = self.llm_service.chat_completion([
                {'role': 'system', 'content': BID_EXPERT_SYSTEM_PROMPT},
                {'role': 'user', 'content': prompt}
            ])
            
            response_text = result.get('content', '')
            
            # 提取JSON
            json_match = re.search(r'\{[\s\S]*\}', response_text)
            if json_match:
                return json.loads(json_match.group())
        
        except Exception as e:
            logger.error(f"LLM提取失败: {e}")
        
        return None


class BidRequirementWordExporter:
    """标书需求Word导出器"""
    
    def __init__(self):
        pass
    
    def export_structured_requirements(self, data: Dict, title: str = '技术要求清单') -> Tuple[str, str]:
        """
        导出结构化的技术要求为Word文档
        
        Args:
            data: 提取的要求数据
            title: 文档标题
        
        Returns:
            (文件路径, 文件名)
        """
        doc = DocumentLoader()
        
        # 设置文档样式
        self._setup_document_style(doc)
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph()
        
        # 添加统计信息
        summary = data.get('summary', {})
        if summary:
            stats_para = doc.add_paragraph()
            stats_run = stats_para.add_run("📊 统计信息")
            stats_run.bold = True
            stats_run.font.size = Pt(14)
            
            stats_content = doc.add_paragraph()
            stats_content.add_run(f"• 总计要求: {summary.get('total_requirements', 0)} 条\n")
            stats_content.add_run(f"• 关键要求: {summary.get('key_requirements', 0)} 条\n")
            categories = summary.get('categories', [])
            if categories:
                stats_content.add_run(f"• 分类数量: {len(categories)} 个 ({', '.join(categories[:5])}{'...' if len(categories) > 5 else ''})")
            
            doc.add_paragraph()
            doc.add_paragraph('─' * 50)
        
        # 按章节输出要求
        sections = data.get('sections', [])
        req_global_index = 0
        
        for section in sections:
            section_number = section.get('number', '')
            section_title = section.get('title', '')
            requirements = section.get('requirements', [])
            
            if not requirements:
                continue
            
            # 章节标题
            section_heading = doc.add_heading(f"{section_number} {section_title}", level=1)
            
            # 按分类分组
            categorized = {}
            uncategorized = []
            
            for req in requirements:
                category = req.get('category')
                if category:
                    if category not in categorized:
                        categorized[category] = []
                    categorized[category].append(req)
                else:
                    uncategorized.append(req)
            
            # 输出有分类的要求
            for category, reqs in categorized.items():
                # 分类标题（一级）
                cat_heading = doc.add_heading(category, level=2)
                
                # 该分类下的要求（二级）
                for req in reqs:
                    req_global_index += 1
                    self._add_requirement_item(doc, req, req_global_index)
            
            # 输出无分类的要求
            for req in uncategorized:
                req_global_index += 1
                self._add_requirement_item(doc, req, req_global_index, is_independent=True)
            
            doc.add_paragraph()
        
        # 保存文件
        temp_dir = tempfile.gettempdir()
        filename = f"技术要求清单_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return filepath, filename
    
    def _setup_document_style(self, doc):
        """设置文档样式"""
        # 可以在这里自定义样式
        pass
    
    def _add_requirement_item(self, doc, req: Dict, index: int, is_independent: bool = False):
        """添加单个要求条目"""
        title = req.get('title', '')
        content = req.get('content', '')
        is_key = req.get('is_key', False)
        
        # 要求标题
        level = 2 if is_independent else 3
        
        # 构建标题文本
        title_text = f"{index}. {title}"
        if is_key:
            title_text = f"★ {title_text}"
        
        req_heading = doc.add_heading(title_text, level=level)
        
        # 要求内容/规格说明
        if content:
            content_para = doc.add_paragraph()
            content_para.paragraph_format.left_indent = Cm(1)
            content_para.add_run(content)
    
    def export_as_checklist(self, data: Dict, title: str = '技术要求检查清单') -> Tuple[str, str]:
        """
        导出为检查清单格式（带复选框）
        
        Args:
            data: 提取的要求数据
            title: 文档标题
        
        Returns:
            (文件路径, 文件名)
        """
        doc = DocumentLoader()
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        doc.add_paragraph()
        
        # 创建表格
        sections = data.get('sections', [])
        
        for section in sections:
            section_number = section.get('number', '')
            section_title = section.get('title', '')
            requirements = section.get('requirements', [])
            
            if not requirements:
                continue
            
            # 章节标题
            doc.add_heading(f"{section_number} {section_title}", level=1)
            
            # 创建检查表格
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            # 表头
            header_cells = table.rows[0].cells
            headers = ['序号', '分类', '技术要求', '技术规格', '响应情况']
            for i, header in enumerate(headers):
                header_cells[i].text = header
                for para in header_cells[i].paragraphs:
                    for run in para.runs:
                        run.bold = True
            
            # 数据行
            for req in requirements:
                row_cells = table.add_row().cells
                
                # 序号
                index_text = str(req.get('index', ''))
                if req.get('is_key'):
                    index_text = f"★ {index_text}"
                row_cells[0].text = index_text
                
                # 分类
                row_cells[1].text = req.get('category', '') or '-'
                
                # 技术要求
                row_cells[2].text = req.get('title', '')
                
                # 技术规格
                row_cells[3].text = req.get('content', '')
                
                # 响应情况（空白，待填写）
                row_cells[4].text = '□ 满足\n□ 部分满足\n□ 不满足'
            
            doc.add_paragraph()
        
        # 保存文件
        temp_dir = tempfile.gettempdir()
        filename = f"技术要求检查清单_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return filepath, filename


# ==================== 便捷函数 ====================

def extract_bid_requirements(file_path: str, section_filter: List[str] = None,
                            llm_service=None) -> Dict:
    """
    提取招标文档中的技术要求
    
    Args:
        file_path: 文档路径
        section_filter: 章节过滤列表
        llm_service: LLM服务实例（可选）
    
    Returns:
        结构化的要求数据
    """
    extractor = BidRequirementExtractor(llm_service)
    return extractor.extract_from_document(file_path, section_filter)


def export_requirements_to_word(data: Dict, format_type: str = 'structured',
                               title: str = None) -> Tuple[str, str]:
    """
    导出要求为Word文档
    
    Args:
        data: 要求数据
        format_type: 格式类型 'structured' | 'checklist'
        title: 文档标题
    
    Returns:
        (文件路径, 文件名)
    """
    exporter = BidRequirementWordExporter()
    
    if format_type == 'checklist':
        return exporter.export_as_checklist(data, title or '技术要求检查清单')
    else:
        return exporter.export_structured_requirements(data, title or '技术要求清单')


def extract_requirements_from_pdf(file_path: str, section_filter: List[str] = None,
                                  llm_config_id: int = None) -> Dict:
    """
    从PDF文件中提取技术要求
    
    使用 PDFParser 解析 PDF 内容，提取章节和表格，
    然后将其转换为结构化的技术要求格式。
    
    Args:
        file_path: PDF 文件路径
        section_filter: 要提取的章节过滤条件列表
        llm_config_id: LLM 配置 ID（可选）
    
    Returns:
        结构化的要求数据，格式与 extract_from_document 相同
    """
    from .pdf_parser import PDFParser, PDFChapter
    
    logger.info(f"[PDF提取] 开始解析PDF: {file_path}")
    
    # 使用 PDFParser 解析 PDF
    parser = PDFParser(file_path)
    if llm_config_id:
        logger.info(f"[PDF提取] 使用LLM解析PDF, llm_config_id={llm_config_id}")
        chapters, full_text = parser.parse_with_llm(llm_config_id=llm_config_id)
    else:
        chapters, full_text = parser.parse()
    
    logger.info(f"[PDF提取] 解析到 {len(chapters)} 个章节")
    
    result = {
        'sections': [],
        'summary': {
            'total_requirements': 0,
            'key_requirements': 0,
            'categories': []
        },
        'raw_tables': []
    }
    
    # 辅助函数：检查章节是否匹配过滤条件
    def match_section(chapter: PDFChapter, filters: List[str]) -> bool:
        if not filters:
            return True
        
        title = chapter.title if isinstance(chapter, PDFChapter) else chapter.get('title', '')
        parent = chapter.parent if isinstance(chapter, PDFChapter) else chapter.get('parent', '')
        
        for f in filters:
            f_clean = f.strip()
            f_lower = f_clean.lower()
            
            # 检查是否是 "父章节::子章节" 格式（精确匹配）
            if '::' in f_clean:
                parent_filter, child_filter = f_clean.split('::', 1)
                parent_filter = parent_filter.strip()
                child_filter = child_filter.strip()
                
                # 检查父章节是否匹配
                parent_matched = False
                if parent:
                    # 检查父章节中是否包含过滤关键字（如 "第五章"）
                    if parent_filter.lower() in parent.lower():
                        parent_matched = True
                    # 检查 "第X章" 格式匹配
                    parent_ch_match = re.search(r'第([一二三四五六七八九十\d]+)[章节篇部]', parent)
                    filter_ch_match = re.search(r'第([一二三四五六七八九十\d]+)[章节篇部]?', parent_filter)
                    if parent_ch_match and filter_ch_match:
                        if parent_ch_match.group(1) == filter_ch_match.group(1):
                            parent_matched = True
                
                # 检查子章节是否匹配
                child_matched = False
                if child_filter.lower() in title.lower():
                    child_matched = True
                
                # 父章节和子章节都匹配才返回 True
                if parent_matched and child_matched:
                    logger.debug(f"[match_section] 精确匹配成功: {parent_filter}::{child_filter} -> {parent} / {title}")
                    return True
                continue
            
            # 普通匹配（向后兼容）
            # 匹配标题
            if f_lower in title.lower():
                return True
            # 匹配父章节
            if parent and f_lower in parent.lower():
                return True
            # 匹配章节编号
            num_match = re.match(r'^(\d+(?:\.\d+)*)', title)
            if num_match:
                section_num = num_match.group(1)
                if f_lower.startswith(section_num) or section_num.startswith(f_lower.rstrip('.')):
                    return True
            # 匹配 "第X章" 格式
            title_ch_match = re.search(r'第([一二三四五六七八九十\d]+)[章节篇部]', title)
            filter_ch_match = re.search(r'第([一二三四五六七八九十\d]+)[章节篇部]?', f_clean)
            if title_ch_match and filter_ch_match:
                if title_ch_match.group(1) == filter_ch_match.group(1):
                    return True
        return False
    
    # 处理每个章节
    for chapter in chapters:
        ch_dict = chapter.to_dict() if isinstance(chapter, PDFChapter) else chapter
        
        # 检查是否需要过滤
        if section_filter and not match_section(chapter, section_filter):
            continue
        
        title = ch_dict.get('title', '')
        content = ch_dict.get('content', '')
        level = ch_dict.get('level', 1)
        parent = ch_dict.get('parent', '')
        
        if not content.strip():
            continue
        
        # 解析内容，提取表格和要求
        requirements = _extract_requirements_from_text(content, title)
        
        if requirements:
            section_data = {
                'number': _extract_section_number(title),
                'title': title,
                'parent': parent,
                'level': level,
                'requirements': requirements
            }
            result['sections'].append(section_data)
            result['summary']['total_requirements'] += len(requirements)
            
            # 统计关键要求
            for req in requirements:
                if req.get('is_key'):
                    result['summary']['key_requirements'] += 1
            
            # 添加分类
            category = parent if parent else title
            if category and category not in result['summary']['categories']:
                result['summary']['categories'].append(category)
    
    logger.info(f"[PDF提取] 完成，共提取 {result['summary']['total_requirements']} 条要求")
    
    return result


def _extract_section_number(title: str) -> str:
    """从标题中提取章节编号"""
    if not title:
        return ''
    
    # 匹配 "1.4.1" 或 "第五章" 格式
    num_match = re.match(r'^(\d+(?:\.\d+)*)', title)
    if num_match:
        return num_match.group(1)
    
    chapter_match = re.match(r'^第([一二三四五六七八九十\d]+)[章节条款]', title)
    if chapter_match:
        return f"第{chapter_match.group(1)}章"
    
    return ''


def _extract_requirements_from_text(content: str, section_title: str) -> List[Dict]:
    """
    从文本内容中提取技术要求
    
    支持识别：
    1. 表格格式（[表格开始]...[表格结束]）
    2. 编号列表格式（1. xxx 或 1）xxx）
    3. 段落格式
    """
    requirements = []
    
    if not content:
        return requirements
    
    # 检查是否包含表格
    table_pattern = r'\[表格开始\](.*?)\[表格结束\]'
    tables = re.findall(table_pattern, content, re.DOTALL)
    
    if tables:
        # 解析表格内容
        for table_content in tables:
            table_reqs = _parse_table_text(table_content)
            requirements.extend(table_reqs)
    
    # 解析编号列表
    lines = content.split('\n')
    current_category = None
    
    for line in lines:
        line = line.strip()
        if not line or line.startswith('[表格'):
            continue
        
        # 检测是否是分类标题（如 "安全可靠测评"、"事务处理机制"）
        if _is_category_title(line):
            current_category = line
            continue
        
        # 检测是否是要求条目
        req_match = re.match(r'^[\d]+[\.、\)）]\s*(.+)$', line)
        if req_match or _looks_like_requirement(line):
            req_text = req_match.group(1) if req_match else line
            is_key = '★' in line or '▲' in line
            
            req = {
                'title': req_text[:100],  # 限制标题长度
                'content': req_text,
                'is_key': is_key,
                'category': current_category,
                'requirement_type': _classify_requirement(req_text)
            }
            requirements.append(req)
    
    return requirements


def _parse_table_text(table_text: str) -> List[Dict]:
    """解析表格文本格式的内容"""
    requirements = []
    
    lines = table_text.strip().split('\n')
    if len(lines) < 2:
        return requirements
    
    # 尝试识别表头
    header_line = lines[0]
    headers = [h.strip() for h in re.split(r'\s{2,}|\t|\|', header_line) if h.strip()]
    
    # 解析每行数据
    current_category = None
    
    for line in lines[1:]:
        line = line.strip()
        if not line:
            continue
        
        # 分割列
        cols = [c.strip() for c in re.split(r'\s{2,}|\t|\|', line) if c.strip()]
        
        if len(cols) >= 2:
            # 检测是否是分类行（合并单元格效果：只有一个有意义的值）
            non_empty = [c for c in cols if c and c not in ['', '-', '—']]
            
            if len(non_empty) == 1 and len(non_empty[0]) < 30:
                # 可能是分类标题
                current_category = non_empty[0]
                continue
            
            # 获取要求内容
            req_title = cols[1] if len(cols) > 1 else cols[0]
            req_spec = cols[2] if len(cols) > 2 else ''
            is_key = '★' in line or '▲' in line
            
            req = {
                'title': req_title[:100],
                'content': f"{req_title}\n{req_spec}" if req_spec else req_title,
                'is_key': is_key,
                'category': current_category,
                'requirement_type': _classify_requirement(req_title)
            }
            requirements.append(req)
    
    return requirements


def _is_category_title(text: str) -> bool:
    """判断是否是分类标题"""
    if not text or len(text) > 30:
        return False
    
    # 分类标题通常是短文本，不包含标点
    if re.search(r'[。，、；：！？]', text):
        return False
    
    # 包含常见分类关键词
    category_keywords = ['能力', '功能', '性能', '安全', '接口', '测评', '机制', '管理', '支持']
    return any(kw in text for kw in category_keywords)


def _looks_like_requirement(text: str) -> bool:
    """判断文本是否看起来像技术要求"""
    if not text or len(text) < 5:
        return False
    
    # 包含技术要求常见关键词
    req_keywords = ['应', '需', '必须', '支持', '提供', '具备', '满足', '实现', '采用']
    return any(kw in text for kw in req_keywords)


def _classify_requirement(text: str) -> str:
    """对要求进行分类"""
    if not text:
        return '其他'
    
    text_lower = text.lower()
    
    if any(kw in text for kw in ['安全', '加密', '认证', '权限', '防护']):
        return '安全'
    if any(kw in text for kw in ['性能', '响应', '并发', '吞吐', '延迟', 'QPS']):
        return '性能'
    if any(kw in text for kw in ['接口', 'API', '协议', '对接', '集成']):
        return '接口'
    if any(kw in text for kw in ['功能', '操作', '管理', '配置', '界面']):
        return '功能'
    
    return '其他'
