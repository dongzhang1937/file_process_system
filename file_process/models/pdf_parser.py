"""
PDF 解析器模块
解析 PDF 文件中的文字、表格和层级结构，用于 LLM 对话中的文件上传
支持两种章节识别方式：
1. 正则表达式模式匹配（快速、离线）
2. LLM智能识别（准确、需要API调用）
"""
import os
import re
import json
from typing import List, Dict, Tuple, Optional
from dataclasses import dataclass, field
from config.logging_config import logger

# 尝试导入 PDF 解析库
try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False
    logger.warning("pdfplumber 未安装，PDF解析功能将受限")

try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    logger.warning("PyMuPDF 未安装，PDF层级解析功能将受限")


@dataclass
class PDFChapter:
    """PDF 章节结构"""
    level: int  # 层级 (1-6)
    title: str  # 标题
    content: str = ""  # 内容
    page_start: int = 0  # 起始页码
    page_end: int = 0  # 结束页码
    parent: str = ""  # 父章节标题（由 Gemini 返回）
    children: List['PDFChapter'] = field(default_factory=list)
    
    def to_dict(self) -> dict:
        return {
            'level': self.level,
            'title': self.title,
            'content': self.content,
            'page_start': self.page_start,
            'page_end': self.page_end,
            'parent': self.parent,
            'children': [child.to_dict() for child in self.children]
        }


class PDFParser:
    """PDF 解析器：提取文字、表格和层级结构
    
    使用场景：用户在 LLM 对话中上传 PDF 文件，然后针对内容提问
    解析目标：
    - 文字内容
    - 表格数据（转换为可读文本格式）
    - 文档层级结构（章节标题）
    - 忽略图片
    """
    
    # 常见章节标题模式（按优先级排序）
    # 注意：模式需要足够严格，避免把普通内容误识别为章节
    CHAPTER_PATTERNS = [
        # 中文章节模式 - 最高优先级
        (r'^第[一二三四五六七八九十百千万零0-9]+[章节篇部]\s+(.+)$', 1),  # 第一章 xxx（必须有空格和标题）
        (r'^第[一二三四五六七八九十百千万零0-9]+[章节篇部]$', 1),  # 第一章（独立成行）
        
        # 中文数字编号 - 需要有明确分隔符
        (r'^[一二三四五六七八九十]+[、．\.]\s*(.+)$', 2),  # 一、xxx 或 一．xxx
        (r'^[\(（][一二三四五六七八九十]+[\)）]\s*(.+)$', 3),  # （一）xxx
        
        # 数字编号 - 必须是 "数字.标题" 格式，标题长度有限制
        (r'^(\d+)[\.、]\s*([^\d\s].{1,50})$', 2),  # 1.基本要求 或 1、基本要求（标题2-50字符）
        (r'^(\d+)[\.、]\s*([^\d\s].{1,50})', 2),  # 更宽松：不要求行结尾
        (r'^(\d+\.\d+)\s+(.{2,50})$', 3),  # 1.1 xxx（需要空格分隔）
        (r'^(\d+\.\d+\.\d+)\s+(.{2,50})$', 4),  # 1.1.1 xxx
        (r'^(\d+\.\d+\.\d+\.\d+)\s+(.{2,50})$', 5),  # 1.1.1.1 xxx
        
        # 英文章节模式
        (r'^Chapter\s+(\d+)[\s\.\:：]+(.+)$', 1, re.IGNORECASE),
        (r'^Section\s+(\d+)[\s\.\:：]+(.+)$', 2, re.IGNORECASE),
    ]
    
    # 表格格式化模板
    TABLE_HEADER = "[表格开始]"
    TABLE_FOOTER = "[表格结束]"
    
    def __init__(self, file_path: str):
        """
        初始化 PDF 解析器
        
        Args:
            file_path: PDF 文件路径
        """
        self.file_path = file_path
        self.filename = os.path.basename(file_path)
        self.chapters: List[PDFChapter] = []
        self.full_text: str = ""
        self.page_count: int = 0
        self.metadata: dict = {}
        
        # 检查文件是否存在
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"PDF 文件不存在: {file_path}")
    
    def parse(self) -> Tuple[List[Dict], str]:
        """
        解析 PDF 文件
        
        Returns:
            (chapters, full_text): 章节结构列表和完整文本
        """
        if HAS_PDFPLUMBER:
            return self._parse_with_pdfplumber()
        elif HAS_PYMUPDF:
            return self._parse_with_pymupdf()
        else:
            raise ImportError("未安装 PDF 解析库，请安装 pdfplumber 或 PyMuPDF")
    
    def _parse_with_pdfplumber(self) -> Tuple[List[Dict], str]:
        """使用 pdfplumber 解析 PDF"""
        text_blocks = []
        tables_by_page = {}
        
        try:
            with pdfplumber.open(self.file_path) as pdf:
                self.page_count = len(pdf.pages)
                self.metadata = pdf.metadata or {}
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # 提取表格
                    tables = page.extract_tables()
                    if tables:
                        tables_by_page[page_num] = tables
                    
                    # 提取文本
                    text = page.extract_text() or ""
                    if text:
                        text_blocks.append({
                            'page': page_num,
                            'text': text,
                            'has_table': bool(tables)
                        })
        except Exception as e:
            logger.error(f"pdfplumber 解析失败: {e}")
            if HAS_PYMUPDF:
                return self._parse_with_pymupdf()
            raise
        
        # 构建完整文本（包含表格）
        full_text_parts = []
        for block in text_blocks:
            page_num = block['page']
            full_text_parts.append(f"--- 第 {page_num} 页 ---\n")
            full_text_parts.append(block['text'])
            
            # 插入该页的表格
            if page_num in tables_by_page:
                for table in tables_by_page[page_num]:
                    formatted_table = self._format_table(table)
                    full_text_parts.append(f"\n{formatted_table}\n")
            
            full_text_parts.append("\n")
        
        self.full_text = '\n'.join(full_text_parts)
        
        # 解析章节结构
        self.chapters = self._extract_chapters(text_blocks)
        
        return [c.to_dict() for c in self.chapters], self.full_text
    
    def _parse_with_pymupdf(self) -> Tuple[List[Dict], str]:
        """使用 PyMuPDF 解析 PDF"""
        text_blocks = []
        
        try:
            doc = fitz.open(self.file_path)
            self.page_count = len(doc)
            self.metadata = doc.metadata or {}
            
            for page_num, page in enumerate(doc, 1):
                # 提取文本
                text = page.get_text("text")
                if text:
                    text_blocks.append({
                        'page': page_num,
                        'text': text,
                        'has_table': False  # PyMuPDF 简单模式下不处理表格
                    })
                
                # 尝试提取表格（使用文本块分析）
                tables = self._extract_tables_from_text(text)
                if tables:
                    text_blocks[-1]['has_table'] = True
                    text_blocks[-1]['tables'] = tables
            
            doc.close()
        except Exception as e:
            logger.error(f"PyMuPDF 解析失败: {e}")
            raise
        
        # 构建完整文本
        full_text_parts = []
        for block in text_blocks:
            page_num = block['page']
            full_text_parts.append(f"--- 第 {page_num} 页 ---\n")
            full_text_parts.append(block['text'])
            full_text_parts.append("\n")
        
        self.full_text = '\n'.join(full_text_parts)
        
        # 解析章节结构
        self.chapters = self._extract_chapters(text_blocks)
        
        return [c.to_dict() for c in self.chapters], self.full_text
    
    def _format_table(self, table: List[List]) -> str:
        """
        格式化表格为可读文本
        
        Args:
            table: 二维表格数据
        
        Returns:
            格式化的表格文本
        """
        if not table:
            return ""
        
        lines = [self.TABLE_HEADER]
        
        for row in table:
            # 清理单元格内容
            cleaned_row = []
            for cell in row:
                cell_text = str(cell) if cell is not None else ""
                # 去除换行符，保持表格行结构
                cell_text = cell_text.replace('\n', ' ').strip()
                cleaned_row.append(cell_text)
            
            lines.append("| " + " | ".join(cleaned_row) + " |")
        
        lines.append(self.TABLE_FOOTER)
        return '\n'.join(lines)
    
    def _extract_tables_from_text(self, text: str) -> List[List[List[str]]]:
        """
        从文本中尝试识别表格结构（简单启发式方法）
        
        Args:
            text: 页面文本
        
        Returns:
            识别出的表格列表
        """
        # 简单的表格识别：查找具有规律性的行
        lines = text.split('\n')
        tables = []
        current_table = []
        
        for line in lines:
            # 如果行中包含多个制表符或连续空格，可能是表格行
            if '\t' in line or re.search(r'\s{3,}', line):
                # 分割单元格
                cells = re.split(r'\t|\s{3,}', line)
                cells = [c.strip() for c in cells if c.strip()]
                if len(cells) >= 2:
                    current_table.append(cells)
            else:
                if len(current_table) >= 2:
                    tables.append(current_table)
                current_table = []
        
        if len(current_table) >= 2:
            tables.append(current_table)
        
        return tables
    
    def _extract_chapters(self, text_blocks: List[Dict]) -> List[PDFChapter]:
        """
        从文本块中提取章节结构
        
        Args:
            text_blocks: 文本块列表
        
        Returns:
            章节结构列表
        """
        chapters = []
        current_chapter = None
        content_buffer = []
        
        for block in text_blocks:
            page_num = block['page']
            lines = block['text'].split('\n')
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                
                # 检查是否是章节标题
                chapter_info = self._match_chapter_title(line)
                
                if chapter_info:
                    # 保存之前的章节内容
                    if current_chapter:
                        current_chapter.content = '\n'.join(content_buffer).strip()
                        current_chapter.page_end = page_num
                        chapters.append(current_chapter)
                    
                    # 创建新章节
                    level, title = chapter_info
                    current_chapter = PDFChapter(
                        level=level,
                        title=title,
                        page_start=page_num
                    )
                    content_buffer = []
                else:
                    # 添加到当前章节内容
                    content_buffer.append(line)
        
        # 处理最后一个章节
        if current_chapter:
            current_chapter.content = '\n'.join(content_buffer).strip()
            current_chapter.page_end = self.page_count
            chapters.append(current_chapter)
        elif content_buffer:
            # 没有识别到章节，创建一个默认章节
            chapters.append(PDFChapter(
                level=1,
                title="正文",
                content='\n'.join(content_buffer).strip(),
                page_start=1,
                page_end=self.page_count
            ))
        
        # 构建层级关系
        return self._build_chapter_hierarchy(chapters)
    
    def _match_chapter_title(self, line: str) -> Optional[Tuple[int, str]]:
        """
        匹配章节标题
        
        Args:
            line: 文本行
        
        Returns:
            (level, title) 或 None，title 保留原始格式（包含编号）
        """
        # 过滤条件
        if len(line) > 80:  # 太长的行不是标题
            return None
        if len(line) < 2:  # 太短的行不是标题
            return None
        
        # 过滤掉页码、页眉页脚等
        if line.isdigit():  # 纯数字（页码）
            return None
        if '招标文件' in line or '投标文件' in line:  # 页眉
            return None
        
        for pattern_item in self.CHAPTER_PATTERNS:
            if len(pattern_item) == 3:
                pattern, level, flags = pattern_item
            else:
                pattern, level = pattern_item
                flags = 0
            
            match = re.match(pattern, line, flags)
            if match:
                # 返回完整的原始行作为标题（保留编号）
                return (level, line)
        
        return None
    
    def _build_chapter_hierarchy(self, chapters: List[PDFChapter]) -> List[PDFChapter]:
        """
        构建章节的层级关系
        
        Args:
            chapters: 扁平的章节列表
        
        Returns:
            具有层级关系的章节列表
        """
        if not chapters:
            return []
        
        # 简单实现：保持扁平结构，用 level 表示层级
        # 更复杂的实现可以构建树状结构
        return chapters
    
    def parse_with_llm(self, use_cache: bool = True, llm_config_id: int = None) -> Tuple[List[Dict], str]:
        """
        使用 LLM 智能识别章节结构
        
        优先使用 Gemini/智谱直接解析，
        最后回退到文本提取 + LLM 分析模式。
        
        Args:
            use_cache: 是否使用缓存（同一文件不重复调用LLM）
            llm_config_id: 指定的LLM配置ID（可选）
        
        Returns:
            (chapters, full_text): 章节结构列表和完整文本
        """
        logger.info(f"[parse_with_llm] 开始解析PDF: {self.filename}, llm_config_id={llm_config_id}")
        
        # 检查是否配置了支持 PDF 的模型（如 Gemini）
        try:
            from .llm_service import get_llm_service
            llm = get_llm_service(llm_config_id)
            
            # 打印当前 LLM 配置信息
            model_type = llm.config.get('model_type', 'unknown')
            model_name = llm.config.get('model_name', 'unknown')
            config_name = llm.config.get('config_name', 'unknown')
            logger.info(f"[parse_with_llm] 当前LLM配置: config_name={config_name}, model_type={model_type}, model_name={model_name}")
            logger.info(f"[parse_with_llm] supports_pdf() = {llm.supports_pdf()}")
            
            if llm.supports_pdf():
                if model_type == 'gemini':
                    logger.info(f"[parse_with_llm] 检测到 Gemini 模型，直接上传 PDF 解析")
                    return self._parse_with_gemini_pdf(llm)
                elif model_type == 'zhipu':
                    logger.info(f"[parse_with_llm] 检测到智谱模型，使用文件解析API")
                    return self._parse_with_zhipu_pdf(llm)
                else:
                    logger.info(f"[parse_with_llm] 模型类型 {model_type} 支持PDF，使用默认解析")
                    return self._parse_with_gemini_pdf(llm)
            else:
                logger.warning(f"[parse_with_llm] 当前模型类型 {model_type} 不支持直接解析PDF，将使用文本提取模式")
        except Exception as e:
            logger.warning(f"[parse_with_llm] 检查 PDF 支持失败: {e}，回退到文本模式")
            import traceback
            logger.warning(f"[parse_with_llm] 异常详情: {traceback.format_exc()}")
        
        # 其他模型：先提取文本再分析
        if HAS_PDFPLUMBER:
            logger.info(f"[parse_with_llm] 使用 pdfplumber 提取文本")
            text_blocks, full_text = self._extract_text_pdfplumber()
        elif HAS_PYMUPDF:
            logger.info(f"[parse_with_llm] 使用 PyMuPDF 提取文本")
            text_blocks, full_text = self._extract_text_pymupdf()
        else:
            raise ImportError("未安装 PDF 解析库，请安装 pdfplumber 或 PyMuPDF")
        
        self.full_text = full_text
        logger.info(f"[parse_with_llm] 文本提取完成，共 {len(full_text)} 字符，{len(text_blocks)} 个文本块")
        
        # 使用 LLM 识别章节
        logger.info(f"[parse_with_llm] 开始调用LLM识别章节...")
        chapters_data = self._identify_chapters_with_llm(full_text)
        
        if chapters_data:
            logger.info(f"[parse_with_llm] LLM识别成功，共 {len(chapters_data)} 个章节，开始构建章节结构")
            # 根据 LLM 返回的章节信息，构建章节结构
            self.chapters = self._build_chapters_from_llm(chapters_data, text_blocks)
            logger.info(f"[parse_with_llm] 章节结构构建完成，最终 {len(self.chapters)} 个章节")
        else:
            # LLM 识别失败，回退到正则匹配
            logger.warning("[parse_with_llm] LLM章节识别失败，回退到正则模式")
            self.chapters = self._extract_chapters(text_blocks)
        
        return [c.to_dict() for c in self.chapters], self.full_text
    
    def _parse_with_gemini_pdf(self, llm) -> Tuple[List[Dict], str]:
        """
        使用 Gemini 直接解析 PDF 文件
        
        Args:
            llm: LLMService 实例
        
        Returns:
            (chapters, full_text): 章节结构列表和完整文本
        """
        prompt = """请分析这个PDF文档，提取完整的章节结构和内容。

**重要：你必须提取每个章节的完整正文内容，不能只返回标题！**

请按照JSON格式返回：
{
  "chapters": [
    {
      "level": 1,
      "title": "第一章 xxx",
      "page": 1,
      "content": "这里是该章节的完整正文内容，包含所有段落、列表、表格数据等..."
    },
    {
      "level": 2,
      "title": "1. 基本要求",
      "page": 2,
      "parent": "第一章 xxx",
      "content": "这里是基本要求章节的所有正文内容，必须完整提取..."
    }
  ]
}

关键要求：
1. **content 字段必须包含该章节的所有正文内容**，不能为空或只有几十个字
2. 提取内容应该从章节标题后开始，直到下一个同级或更高级章节为止
3. 如果章节内容包含编号列表（如 1. 2. 3. 或 ① ② ③），必须完整保留
4. **【重要】如果章节内容包含表格，必须使用 Markdown 管道符格式完整保留表格的所有列和所有行**，格式如下：
   | 序号 | 重要性 | 指标项名称 | 技术指标要求 |
   | 1 | ★ | 安全可靠测评 | 产品应当符合安全可靠测评要求... |
   | 2 | # | 事务处理机制 | 支持事务隔离，可为实例粒度设置不同事务隔离级别... |
   - 表格的每一列都必须完整保留，不能省略或合并列
   - 如果同一行的内容跨多行显示，必须合并到同一行的对应列中
   - 不能把表格拆散成单独的文本行
5. level 表示层级深度：1=章（如"第一章"）、2=节（如"1.基本要求"）、3=小节
6. title 必须保留原文中的完整标题（包含编号）
7. 如果是子章节，在 parent 字段标注其父章节标题

请只返回JSON格式的结果，不要其他说明文字。"""

        try:
            logger.info(f"[Gemini PDF] 开始直接解析PDF文件: {self.file_path}")
            result = llm.analyze_pdf_with_gemini(self.file_path, prompt)
            content = result.get('content', '')
            
            logger.info(f"=" * 80)
            logger.info(f"[Gemini PDF] ========== 原始返回结果（完整）==========")
            # 打印完整结果，不截断
            for i in range(0, len(content), 2000):
                logger.info(content[i:i+2000])
            logger.info(f"[Gemini PDF] ========== 原始结果结束 ==========")
            logger.info(f"=" * 80)
            
            # 解析 JSON 响应
            chapters_data = self._parse_llm_response(content)
            
            if chapters_data:
                logger.info(f"[Gemini PDF] 成功解析 {len(chapters_data)} 个章节")
                
                # 构建章节对象
                self.chapters = []
                full_text_parts = []
                
                # 用于推断 parent 的辅助变量
                last_level1_title = ""
                last_level2_title = ""
                
                for i, ch_data in enumerate(chapters_data):
                    title = ch_data.get('title', f'章节{i+1}')
                    level = ch_data.get('level', 1)
                    ch_content = ch_data.get('content', '')
                    page = ch_data.get('page', 0)
                    parent = ch_data.get('parent', '')  # 提取 parent 字段
                    
                    # 如果 Gemini 没有返回 parent，尝试根据层级推断
                    if not parent and level > 1:
                        if level == 2 and last_level1_title:
                            parent = last_level1_title
                        elif level == 3 and last_level2_title:
                            parent = last_level2_title
                        elif level >= 2 and last_level1_title:
                            parent = last_level1_title
                    
                    # 更新层级跟踪
                    if level == 1:
                        last_level1_title = title
                        last_level2_title = ""
                    elif level == 2:
                        last_level2_title = title
                    
                    chapter = PDFChapter(
                        level=level,
                        title=title,
                        content=ch_content,
                        page_start=page,
                        page_end=page,
                        parent=parent  # 保存 parent
                    )
                    self.chapters.append(chapter)
                    
                    # 构建完整文本
                    full_text_parts.append(f"## {title}\n{ch_content}\n")
                    
                    logger.info(f"  [{i+1}] level={level}, title={title[:40]}, content_len={len(ch_content)}, parent={parent[:30] if parent else '(推断:{})'.format(parent or '无')}")
                
                self.full_text = '\n'.join(full_text_parts)
                self.page_count = max([ch.get('page', 0) for ch in chapters_data], default=0)
                
                return [c.to_dict() for c in self.chapters], self.full_text
            else:
                logger.warning("[Gemini PDF] 无法解析章节结构，尝试回退方案")
                raise Exception("Gemini 返回的内容无法解析为章节结构")
                
        except Exception as e:
            logger.error(f"[Gemini PDF] 解析失败: {e}，回退到文本提取模式")
            # 回退到传统方式
            if HAS_PDFPLUMBER:
                text_blocks, full_text = self._extract_text_pdfplumber()
            elif HAS_PYMUPDF:
                text_blocks, full_text = self._extract_text_pymupdf()
            else:
                raise
            
            self.full_text = full_text
            self.chapters = self._extract_chapters(text_blocks)
            return [c.to_dict() for c in self.chapters], self.full_text
    
    def _extract_chapters_from_markdown(self, markdown_content: str) -> List[PDFChapter]:
        """
        从 Markdown 内容中提取章节结构。
        
        识别 Markdown 标题（# ## ### 等）和中文编号标题，将内容按章节切分。
        
        Args:
            markdown_content: Markdown 格式的文档内容
            
        Returns:
            PDFChapter 列表
        """
        lines = markdown_content.split('\n')
        chapters = []
        current_title = None
        current_level = 1
        current_content_lines = []
        current_page = 0
        
        # 用于推断 parent 的辅助变量
        last_titles = {}  # {level: title}
        
        for line in lines:
            stripped = line.strip()
            
            # 检测 Markdown 标题（# 开头）
            md_heading = re.match(r'^(#{1,6})\s+(.+)$', stripped)
            if md_heading:
                level = len(md_heading.group(1))
                title = md_heading.group(2).strip()
            else:
                # 检测中文编号标题（用现有的 CHAPTER_PATTERNS）
                chapter_match = self._match_chapter_title(stripped) if stripped and 2 <= len(stripped) <= 80 else None
                if chapter_match:
                    level, title = chapter_match
                else:
                    # 不是标题，追加到当前内容
                    current_content_lines.append(line)
                    continue
            
            # 遇到新标题，保存前一个章节
            if current_title is not None:
                content = '\n'.join(current_content_lines).strip()
                parent = ''
                for lv in range(current_level - 1, 0, -1):
                    if lv in last_titles:
                        parent = last_titles[lv]
                        break
                
                chapters.append(PDFChapter(
                    level=current_level,
                    title=current_title,
                    content=content,
                    page_start=current_page,
                    page_end=current_page,
                    parent=parent
                ))
            
            # 更新当前章节
            current_title = title
            current_level = level
            current_content_lines = []
            last_titles[level] = title
            # 清除更低层级的 parent 记录
            for lv in list(last_titles.keys()):
                if lv > level:
                    del last_titles[lv]
        
        # 保存最后一个章节
        if current_title is not None:
            content = '\n'.join(current_content_lines).strip()
            parent = ''
            for lv in range(current_level - 1, 0, -1):
                if lv in last_titles:
                    parent = last_titles[lv]
                    break
            chapters.append(PDFChapter(
                level=current_level,
                title=current_title,
                content=content,
                page_start=current_page,
                page_end=current_page,
                parent=parent
            ))
        
        # 兜底：如果没识别到任何章节标题，创建一个默认章节
        if not chapters:
            logger.warning("[Markdown解析] 未检测到章节标题，创建默认章节")
            chapters.append(PDFChapter(
                level=1,
                title="正文",
                content=markdown_content,
                page_start=1,
                page_end=1,
                parent=""
            ))
        
        return chapters
    
    def _parse_with_zhipu_pdf(self, llm) -> Tuple[List[Dict], str]:
        """
        使用智谱文件解析API解析PDF文件
        
        Args:
            llm: LLMService 实例
        
        Returns:
            (chapters, full_text): 章节结构列表和完整文本
        """
        prompt = """请分析这个PDF文档，提取完整的章节结构和内容。

**重要：你必须提取每个章节的完整正文内容，不能只返回标题！**

请按照JSON格式返回：
{
  "chapters": [
    {
      "level": 1,
      "title": "第一章 xxx",
      "page": 1,
      "content": "这里是该章节的完整正文内容，包含所有段落、列表、表格数据等..."
    },
    {
      "level": 2,
      "title": "1. 基本要求",
      "page": 2,
      "parent": "第一章 xxx",
      "content": "这里是基本要求章节的所有正文内容，必须完整提取..."
    }
  ]
}

关键要求：
1. content 字段必须包含该章节的所有正文内容，不能为空或只有几十个字
2. level 表示层级深度：1=章（如"第一章"）、2=节（如"1.基本要求"）、3=小节
3. title 必须保留原文中的完整标题（包含编号）
4. 如果是子章节，在 parent 字段标注其父章节标题
5. **【重要】如果章节内容包含表格，必须使用 Markdown 管道符格式完整保留表格的所有列和所有行**，格式如下：
   | 序号 | 重要性 | 指标项名称 | 技术指标要求 |
   | 1 | ★ | 安全可靠测评 | 产品应当符合安全可靠测评要求... |
   | 2 | # | 事务处理机制 | 支持事务隔离，可为实例粒度设置不同事务隔离级别... |
   - 表格的每一列都必须完整保留，不能省略或合并列
   - 如果同一行的内容跨多行显示，必须合并到同一行的对应列中
   - 不能把表格拆散成单独的文本行

请只返回JSON格式的结果，不要其他说明文字。"""

        try:
            logger.info(f"[智谱 PDF] 开始使用文件解析API解析: {self.file_path}")
            result = llm.analyze_pdf_with_zhipu(self.file_path, prompt)
            content = result.get('content', '')
            
            logger.info(f"[智谱 PDF] LLM返回内容长度: {len(content)}")
            
            # 解析 JSON 响应
            chapters_data = self._parse_llm_response(content)
            
            if chapters_data:
                logger.info(f"[智谱 PDF] 成功解析 {len(chapters_data)} 个章节")
                
                self.chapters = []
                full_text_parts = []
                
                last_level1_title = ""
                last_level2_title = ""
                
                for i, ch_data in enumerate(chapters_data):
                    title = ch_data.get('title', f'章节{i+1}')
                    level = ch_data.get('level', 1)
                    ch_content = ch_data.get('content', '')
                    page = ch_data.get('page', 0)
                    parent = ch_data.get('parent', '')
                    
                    if not parent and level > 1:
                        if level == 2 and last_level1_title:
                            parent = last_level1_title
                        elif level == 3 and last_level2_title:
                            parent = last_level2_title
                        elif level >= 2 and last_level1_title:
                            parent = last_level1_title
                    
                    if level == 1:
                        last_level1_title = title
                        last_level2_title = ""
                    elif level == 2:
                        last_level2_title = title
                    
                    chapter = PDFChapter(
                        level=level,
                        title=title,
                        content=ch_content,
                        page_start=page,
                        page_end=page,
                        parent=parent
                    )
                    self.chapters.append(chapter)
                    full_text_parts.append(f"## {title}\n{ch_content}\n")
                    
                    logger.info(f"  [{i+1}] level={level}, title={title[:40]}, content_len={len(ch_content)}, parent={parent[:30] if parent else '无'}")
                
                self.full_text = '\n'.join(full_text_parts)
                self.page_count = max([ch.get('page', 0) for ch in chapters_data], default=0)
                
                return [c.to_dict() for c in self.chapters], self.full_text
            else:
                logger.warning("[智谱 PDF] 无法解析章节结构，回退到文本提取模式")
                raise Exception("智谱文件解析返回的内容无法解析为章节结构")
                
        except Exception as e:
            logger.error(f"[智谱 PDF] 解析失败: {e}，回退到文本提取模式")
            if HAS_PDFPLUMBER:
                text_blocks, full_text = self._extract_text_pdfplumber()
            elif HAS_PYMUPDF:
                text_blocks, full_text = self._extract_text_pymupdf()
            else:
                raise
            
            self.full_text = full_text
            self.chapters = self._extract_chapters(text_blocks)
            return [c.to_dict() for c in self.chapters], self.full_text
    
    def _extract_text_pdfplumber(self) -> Tuple[List[Dict], str]:
        """使用 pdfplumber 提取文本"""
        text_blocks = []
        tables_by_page = {}
        full_text_parts = []
        
        with pdfplumber.open(self.file_path) as pdf:
            self.page_count = len(pdf.pages)
            self.metadata = pdf.metadata or {}
            
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()
                if tables:
                    tables_by_page[page_num] = tables
                
                text = page.extract_text() or ""
                if text:
                    text_blocks.append({
                        'page': page_num,
                        'text': text,
                        'has_table': bool(tables)
                    })
        
        # 构建完整文本
        for block in text_blocks:
            page_num = block['page']
            full_text_parts.append(f"--- 第 {page_num} 页 ---\n")
            full_text_parts.append(block['text'])
            
            if page_num in tables_by_page:
                for table in tables_by_page[page_num]:
                    formatted_table = self._format_table(table)
                    full_text_parts.append(f"\n{formatted_table}\n")
            
            full_text_parts.append("\n")
        
        return text_blocks, '\n'.join(full_text_parts)
    
    def _extract_text_pymupdf(self) -> Tuple[List[Dict], str]:
        """使用 PyMuPDF 提取文本"""
        text_blocks = []
        full_text_parts = []
        
        doc = fitz.open(self.file_path)
        self.page_count = len(doc)
        self.metadata = doc.metadata or {}
        
        for page_num, page in enumerate(doc, 1):
            text = page.get_text("text")
            if text:
                text_blocks.append({
                    'page': page_num,
                    'text': text,
                    'has_table': False
                })
                full_text_parts.append(f"--- 第 {page_num} 页 ---\n")
                full_text_parts.append(text)
                full_text_parts.append("\n")
        
        doc.close()
        return text_blocks, '\n'.join(full_text_parts)
    
    def _identify_chapters_with_llm(self, full_text: str) -> Optional[List[Dict]]:
        """
        使用 LLM 识别章节结构
        
        Args:
            full_text: 完整文本内容
        
        Returns:
            章节信息列表 或 None（失败时）
        """
        try:
            from .llm_service import get_llm_service
            
            llm = get_llm_service()
            
            # 为了避免 token 过长，只取前 8000 字符用于章节识别
            sample_text = full_text[:8000] if len(full_text) > 8000 else full_text
            
            prompt = f"""请分析以下文档内容，识别其中的章节层级结构。

文档内容：
{sample_text}

请按照以下JSON格式返回章节结构（只返回JSON，不要其他内容）：
{{
  "chapters": [
    {{"level": 1, "title": "第一章 xxx", "keywords": ["关键词1", "关键词2"]}},
    {{"level": 2, "title": "1.基本要求", "parent": "第一章 xxx"}},
    {{"level": 2, "title": "2.技术规格", "parent": "第一章 xxx"}}
  ]
}}

注意：
1. level 表示层级深度：1=章、2=节、3=小节、4=条款
2. title 必须保留原文中的完整标题（包含编号）
3. 识别中文章节格式（如：第一章、第五章、一、二、1.、2.、1.1等）
4. keywords 是该章节的关键词，用于后续匹配
5. 如果是子章节，请在parent字段标注其父章节标题"""

            result = llm.chat_completion([
                {'role': 'system', 'content': '你是一个专业的文档结构分析助手，擅长识别文档的章节层级结构。'},
                {'role': 'user', 'content': prompt}
            ])
            
            content = result.get('content', '')
            logger.info(f"=" * 50)
            logger.info(f"[LLM章节识别] 原始返回结果:")
            logger.info(content)
            logger.info(f"=" * 50)
            
            # 解析 JSON 响应
            chapters_data = self._parse_llm_response(content)
            logger.info(f"[LLM章节识别] 解析后章节数量: {len(chapters_data) if chapters_data else 0}")
            if chapters_data:
                logger.info(f"[LLM章节识别] 完整章节列表:")
                for i, ch in enumerate(chapters_data):
                    logger.info(f"  [{i+1}] {ch}")
            else:
                logger.warning(f"[LLM章节识别] 解析失败，无法提取章节信息")
            return chapters_data
            
        except Exception as e:
            logger.error(f"LLM章节识别失败: {e}")
            return None
    
    def _parse_llm_response(self, content: str) -> Optional[List[Dict]]:
        """
        解析 LLM 返回的 JSON 响应
        
        Args:
            content: LLM 响应内容
        
        Returns:
            章节列表 或 None
        """
        try:
            # 尝试直接解析
            data = json.loads(content)
            if isinstance(data, dict) and 'chapters' in data:
                return data['chapters']
            elif isinstance(data, list):
                return data
        except json.JSONDecodeError:
            pass
        
        # 尝试提取 JSON 块
        json_patterns = [
            r'```json\s*([\s\S]*?)\s*```',
            r'```\s*([\s\S]*?)\s*```',
            r'\{[\s\S]*"chapters"[\s\S]*\}',
            r'\[[\s\S]*\]'
        ]
        
        for pattern in json_patterns:
            match = re.search(pattern, content)
            if match:
                try:
                    json_str = match.group(1) if '```' in pattern else match.group(0)
                    data = json.loads(json_str)
                    if isinstance(data, dict) and 'chapters' in data:
                        return data['chapters']
                    elif isinstance(data, list):
                        return data
                except (json.JSONDecodeError, IndexError):
                    continue
        
        logger.warning(f"无法解析LLM响应: {content[:200]}")
        return None
    
    def _build_chapters_from_llm(self, chapters_data: List[Dict], 
                                  text_blocks: List[Dict]) -> List[PDFChapter]:
        """
        根据 LLM 识别的章节信息构建章节结构
        
        Args:
            chapters_data: LLM返回的章节数据
            text_blocks: 文本块列表
        
        Returns:
            章节列表
        """
        chapters = []
        
        # 将所有文本合并，按行分割
        all_lines = []
        for block in text_blocks:
            page_num = block['page']
            for line in block['text'].split('\n'):
                line = line.strip()
                if line:
                    all_lines.append({'page': page_num, 'text': line})
        
        # 为每个章节找到对应的起始位置和内容
        for i, ch_data in enumerate(chapters_data):
            title = ch_data.get('title', '')
            level = ch_data.get('level', 1)
            keywords = ch_data.get('keywords', [])
            
            # 在文本中找到这个章节标题
            start_idx, start_page = self._find_chapter_position(title, keywords, all_lines)
            
            if start_idx == -1:
                logger.warning(f"LLM章节定位失败: {title}")
                continue
            
            logger.info(f"LLM章节定位成功: {title} -> 行{start_idx}, 页{start_page}")
            
            # 找到下一个章节的位置（作为当前章节的结束）
            if i + 1 < len(chapters_data):
                next_title = chapters_data[i + 1].get('title', '')
                next_keywords = chapters_data[i + 1].get('keywords', [])
                end_idx, end_page = self._find_chapter_position(next_title, next_keywords, all_lines, start_idx + 1)
                if end_idx == -1:
                    end_idx = len(all_lines)
                    end_page = self.page_count
            else:
                end_idx = len(all_lines)
                end_page = self.page_count
            
            # 提取章节内容
            content_lines = [all_lines[j]['text'] for j in range(start_idx + 1, end_idx)]
            content = '\n'.join(content_lines)
            
            chapter = PDFChapter(
                level=level,
                title=title,
                content=content,
                page_start=start_page,
                page_end=end_page
            )
            chapters.append(chapter)
        
        return chapters if chapters else self._extract_chapters(text_blocks)
    
    def _find_chapter_position(self, title: str, keywords: List[str], 
                               lines: List[Dict], start_from: int = 0) -> Tuple[int, int]:
        """
        在文本行中查找章节位置
        
        Args:
            title: 章节标题
            keywords: 关键词列表
            lines: 文本行列表
            start_from: 从哪一行开始搜索
        
        Returns:
            (行索引, 页码) 或 (-1, 0)
        """
        # 清理标题用于匹配
        clean_title = re.sub(r'\s+', '', title)
        
        for i in range(start_from, len(lines)):
            line_text = lines[i]['text']
            clean_line = re.sub(r'\s+', '', line_text)
            
            # 精确匹配
            if clean_title == clean_line:
                return i, lines[i]['page']
            
            # 标题包含在行中
            if clean_title in clean_line or clean_line in clean_title:
                return i, lines[i]['page']
            
            # 关键词匹配（至少匹配一个关键词）
            if keywords:
                for kw in keywords:
                    if kw in line_text:
                        # 额外验证：检查是否看起来像标题
                        if self._looks_like_title(line_text):
                            return i, lines[i]['page']
        
        return -1, 0
    
    def _looks_like_title(self, text: str) -> bool:
        """判断文本是否看起来像章节标题"""
        # 长度检查
        if len(text) > 80 or len(text) < 2:
            return False
        
        # 匹配常见标题模式
        title_patterns = [
            r'^第[一二三四五六七八九十百千万零0-9]+[章节篇部]',
            r'^[一二三四五六七八九十]+[、．\.]',
            r'^[\(（][一二三四五六七八九十]+[\)）]',
            r'^\d+\.\s*\S',
            r'^\d+\.\d+',
        ]
        
        for pattern in title_patterns:
            if re.match(pattern, text):
                return True
        
        return False
    
    def get_chapter_by_title(self, title: str) -> Optional[PDFChapter]:
        """
        根据标题查找章节
        
        Args:
            title: 章节标题（支持部分匹配）
        
        Returns:
            匹配的章节或 None
        """
        title_lower = title.lower()
        for chapter in self.chapters:
            if title_lower in chapter.title.lower():
                return chapter
        return None
    
    def get_chapters_by_range(self, start_chapter: str = None, 
                              end_chapter: str = None) -> List[PDFChapter]:
        """
        获取指定范围内的章节
        
        Args:
            start_chapter: 起始章节标题
            end_chapter: 结束章节标题
        
        Returns:
            章节列表
        """
        if not self.chapters:
            return []
        
        start_idx = 0
        end_idx = len(self.chapters)
        
        if start_chapter:
            for i, chapter in enumerate(self.chapters):
                if start_chapter.lower() in chapter.title.lower():
                    start_idx = i
                    break
        
        if end_chapter:
            for i, chapter in enumerate(self.chapters):
                if end_chapter.lower() in chapter.title.lower():
                    end_idx = i + 1
                    break
        
        return self.chapters[start_idx:end_idx]
    
    def get_content_for_qa(self, chapters: List[str] = None, 
                           max_length: int = 50000) -> str:
        """
        获取用于问答的内容
        
        Args:
            chapters: 指定的章节标题列表（为 None 则使用全部内容）
            max_length: 最大内容长度
        
        Returns:
            格式化的内容文本
        """
        if chapters:
            # 获取指定章节的内容
            content_parts = []
            for chapter_title in chapters:
                chapter = self.get_chapter_by_title(chapter_title)
                if chapter:
                    content_parts.append(f"## {chapter.title}\n{chapter.content}")
            
            content = '\n\n'.join(content_parts)
        else:
            # 使用完整内容
            content = self.full_text
        
        # 截断过长的内容
        if len(content) > max_length:
            content = content[:max_length] + "\n\n...[内容已截断]"
        
        return content
    
    def get_summary(self) -> dict:
        """
        获取 PDF 文件摘要信息
        
        Returns:
            摘要信息字典
        """
        return {
            'filename': self.filename,
            'page_count': self.page_count,
            'chapter_count': len(self.chapters),
            'chapters': [c.title for c in self.chapters],
            'total_chars': len(self.full_text),
            'metadata': self.metadata
        }


class DocumentParser:
    """统一文档解析器：支持 PDF 和 Word 文档
    
    - PDF 文档通过 LLM / 正则解析
    - Word 文档通过 python-docx 深度解析（段落 + 表格按原始顺序提取）
    """
    
    # Word 文档中常见的章节标题正则（与 PDFParser.CHAPTER_PATTERNS 保持一致）
    HEADING_PATTERNS = [
        (r'^第[一二三四五六七八九十百千万零0-9]+[章节篇部]\s*(.*)$', 1),
        (r'^[一二三四五六七八九十]+[、．\.]\s*(.+)$', 2),
        (r'^[\(（][一二三四五六七八九十]+[\)）]\s*(.+)$', 3),
        (r'^(\d+)[\.、]\s*([^\d\s].{1,50})$', 2),
        (r'^(\d+\.\d+)\s*(.{2,50})$', 3),
        (r'^(\d+\.\d+\.\d+)\s*(.{2,50})$', 4),
    ]
    
    @staticmethod
    def _iter_block_items(doc):
        """
        按文档中的原始顺序迭代段落和表格。
        
        python-docx 的 doc.paragraphs 只返回段落，doc.tables 只返回表格，
        二者都丢失了相对顺序。此方法从底层 XML body 中按序遍历，
        返回 ('para', Paragraph) 或 ('table', Table) 元组。
        """
        from docx.text.paragraph import Paragraph
        from docx.table import Table as DocxTable
        
        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                yield ('para', Paragraph(child, doc))
            elif tag == 'tbl':
                yield ('table', DocxTable(child, doc))
    
    @classmethod
    def _detect_heading(cls, text: str, style_name: str = '', is_bold: bool = False,
                        font_size_pt: float = 0) -> Optional[Tuple[int, str]]:
        """
        检测一段文本是否为章节标题。
        
        判断依据（按优先级）：
        1. Word 内置 Heading 样式
        2. Word 大纲级别 / TOC 标题样式
        3. 正则匹配中文/数字编号
        4. 全段加粗 + 较大字号（>=14pt）作为标题
        
        Returns:
            (level, title) 或 None
        """
        if not text or len(text) > 100 or len(text) < 2:
            return None
        
        # 1. Word 内置 Heading 样式
        style_lower = style_name.lower()
        if 'heading 1' in style_lower or 'heading1' in style_lower:
            return (1, text)
        elif 'heading 2' in style_lower or 'heading2' in style_lower:
            return (2, text)
        elif 'heading 3' in style_lower or 'heading3' in style_lower:
            return (3, text)
        elif 'heading 4' in style_lower or 'heading4' in style_lower:
            return (4, text)
        
        # 中文标题样式名
        if '标题 1' in style_name or '标题1' in style_name:
            return (1, text)
        elif '标题 2' in style_name or '标题2' in style_name:
            return (2, text)
        elif '标题 3' in style_name or '标题3' in style_name:
            return (3, text)
        
        # TOC / 目录样式
        if 'toc' in style_lower:
            return None  # 目录行不当标题
        
        # 2. 正则匹配
        for pattern, level in cls.HEADING_PATTERNS:
            if re.match(pattern, text):
                return (level, text)
        
        # 3. 全段加粗 + 大字号 → 当作标题
        if is_bold and font_size_pt >= 14 and len(text) <= 60:
            return (1, text)
        elif is_bold and font_size_pt >= 12 and len(text) <= 60:
            return (2, text)
        
        return None
    
    @staticmethod
    def _format_table(table) -> str:
        """
        将 Word 表格格式化为 Markdown 管道符格式。
        
        处理合并单元格（python-docx 中合并单元格会重复出现同一 cell 对象）。
        """
        lines = ['[表格开始]']
        
        for row_idx, row in enumerate(table.rows):
            cells = []
            prev_text = None
            for cell in row.cells:
                cell_text = cell.text.replace('\n', ' ').strip()
                # 处理水平合并：python-docx 中合并的单元格会重复同一 cell 对象
                if cell_text == prev_text and prev_text is not None:
                    continue  # 跳过合并的重复单元格
                cells.append(cell_text)
                prev_text = cell_text
            
            lines.append('| ' + ' | '.join(cells) + ' |')
            
            # 在第一行（表头）后添加分隔线
            if row_idx == 0:
                lines.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
        
        lines.append('[表格结束]')
        return '\n'.join(lines)
    
    @staticmethod
    def _get_para_properties(para) -> Tuple[bool, float]:
        """
        获取段落的格式属性：是否全段加粗、字号大小。
        
        Returns:
            (is_bold, font_size_pt)
        """
        is_bold = False
        font_size_pt = 0.0
        
        try:
            runs = para.runs
            if not runs:
                return is_bold, font_size_pt
            
            # 检查所有 run 是否都加粗
            bold_runs = sum(1 for r in runs if r.bold)
            if bold_runs == len(runs) and len(runs) > 0:
                is_bold = True
            
            # 获取第一个有字号信息的 run 的字号
            for r in runs:
                if r.font and r.font.size:
                    font_size_pt = r.font.size.pt
                    break
        except Exception:
            pass
        
        return is_bold, font_size_pt
    
    @classmethod
    def parse_docx(cls, docx_path: str) -> Tuple[List[Dict], str]:
        """
        使用 python-docx 深度解析 Word 文档。
        
        按文档中段落和表格的原始顺序遍历，完整提取：
        - 章节标题（通过样式名 / 加粗字号 / 正则编号识别）
        - 段落正文
        - 表格数据（Markdown 管道符格式）
        - 列表项（保持编号）
        
        Args:
            docx_path: Word 文档路径
            
        Returns:
            (chapters_list, full_text)  chapters_list 是 dict 列表
        """
        from docx import Document
        
        doc = Document(docx_path)
        
        chapters = []       # 最终的章节列表
        full_text_parts = []  # 全文拼接
        
        # 当前正在收集的章节状态
        current_title = None
        current_level = 1
        current_content_lines = []
        last_titles = {}    # {level: title} 用于推断 parent
        
        def _save_current_chapter():
            """将当前收集中的章节保存到列表"""
            nonlocal current_title, current_content_lines
            if current_title is not None:
                content = '\n'.join(current_content_lines).strip()
                parent = ''
                for lv in range(current_level - 1, 0, -1):
                    if lv in last_titles:
                        parent = last_titles[lv]
                        break
                chapters.append({
                    'level': current_level,
                    'title': current_title,
                    'content': content,
                    'page_start': 0,
                    'page_end': 0,
                    'parent': parent,
                    'children': []
                })
        
        logger.info(f"[WordParser] 开始解析: {docx_path}")
        element_count = {'para': 0, 'table': 0, 'heading': 0}
        
        for block_type, block in cls._iter_block_items(doc):
            if block_type == 'para':
                element_count['para'] += 1
                text = block.text.strip()
                if not text:
                    continue
                
                style_name = block.style.name if block.style else ''
                is_bold, font_size_pt = cls._get_para_properties(block)
                
                heading_info = cls._detect_heading(text, style_name, is_bold, font_size_pt)
                
                if heading_info:
                    element_count['heading'] += 1
                    # 遇到新标题，先保存上一个章节
                    _save_current_chapter()
                    
                    current_level, current_title = heading_info
                    current_content_lines = []
                    
                    # 更新层级跟踪
                    last_titles[current_level] = current_title
                    for lv in list(last_titles.keys()):
                        if lv > current_level:
                            del last_titles[lv]
                    
                    full_text_parts.append(f"\n{'#' * min(current_level, 6)} {current_title}\n")
                else:
                    # 普通段落内容
                    current_content_lines.append(text)
                    full_text_parts.append(text)
                    
            elif block_type == 'table':
                element_count['table'] += 1
                try:
                    table_text = cls._format_table(block)
                    current_content_lines.append(table_text)
                    full_text_parts.append(table_text)
                except Exception as e:
                    logger.warning(f"[WordParser] 解析表格失败: {e}")
        
        # 保存最后一个章节
        _save_current_chapter()
        
        # 兜底：如果没识别到任何标题
        if not chapters:
            full_text = '\n'.join(full_text_parts).strip()
            logger.warning("[WordParser] 未检测到章节标题，将整个文档作为一个章节")
            chapters.append({
                'level': 1,
                'title': '正文',
                'content': full_text,
                'page_start': 0,
                'page_end': 0,
                'parent': '',
                'children': []
            })
        
        full_text = '\n'.join(full_text_parts)
        logger.info(f"[WordParser] 解析完成: {element_count['para']} 个段落, "
                     f"{element_count['table']} 个表格, "
                     f"{element_count['heading']} 个标题, "
                     f"{len(chapters)} 个章节, "
                     f"{len(full_text)} 字符")
        
        return chapters, full_text
    
    @classmethod
    def parse(cls, file_path: str, use_llm: bool = True, llm_config_id: int = None) -> Tuple[List[Dict], str, dict]:
        """
        统一解析入口：根据文件类型自动选择解析方式
        
        - PDF → LLM / 正则
        - Word → python-docx 深度解析
        
        Args:
            file_path: 文件路径（.pdf 或 .docx）
            use_llm: 是否使用 LLM 智能解析（仅 PDF 有效）
            llm_config_id: LLM 配置 ID
            
        Returns:
            (chapters, full_text, summary)
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        if ext == '.pdf':
            parser = PDFParser(file_path)
            if use_llm:
                chapters, full_text = parser.parse_with_llm(llm_config_id=llm_config_id)
            else:
                chapters, full_text = parser.parse()
            summary = parser.get_summary()
            return chapters, full_text, summary
        
        elif ext in ('.docx', '.doc'):
            chapters, full_text = cls.parse_docx(file_path)
            summary = {
                'filename': os.path.basename(file_path),
                'page_count': 0,
                'chapter_count': len(chapters),
                'chapters': [c['title'] for c in chapters],
                'total_chars': len(full_text),
                'metadata': {}
            }
            return chapters, full_text, summary
        
        else:
            raise ValueError(f"不支持的文件类型: {ext}，仅支持 .pdf 和 .docx")

