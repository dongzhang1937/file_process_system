"""
文档对话模块 - 支持精准匹配和LLM匹配
"""
import os
import re
import json
import tempfile
from datetime import datetime
from flask import Blueprint, request, jsonify, render_template, session, send_file
from config.db_config import fetch_one, fetch_all, dml_sql
from config.logging_config import logger

chatdoc = Blueprint('chatdoc', __name__)

# 获取基础目录
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))


def clean_text_for_fuzzy(text):
    """
    清理文本用于模糊匹配：去掉标点符号和特殊字符
    """
    if not text:
        return ""
    # 去掉所有标点符号和特殊字符，只保留中文、英文、数字
    cleaned = re.sub(r'[^\u4e00-\u9fa5a-zA-Z0-9]', '', text)
    return cleaned.lower()


def get_chapter_with_children(chapter_id, document_id):
    """
    获取章节及其所有子章节（递归）
    """
    result = []
    
    # 获取当前章节
    chapter_sql = """
        SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index, 
               c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index
        FROM chapters c
        WHERE c.id = %s AND c.document_id = %s
    """
    chapter = fetch_one(chapter_sql, (chapter_id, document_id))
    if chapter:
        # 获取章节关联的图片
        images = get_chapter_images(chapter_id)
        chapter['images'] = images
        result.append(chapter)
        
        # 递归获取子章节
        children_sql = """
            SELECT id FROM chapters 
            WHERE parent_id = %s AND document_id = %s
            ORDER BY order_index
        """
        children = fetch_all(children_sql, (chapter_id, document_id))
        for child in children:
            result.extend(get_chapter_with_children(child['id'], document_id))
    
    return result


def get_chapter_path(chapter_id):
    """
    获取章节的完整层级路径（从根到当前章节）
    返回路径列表，每个元素包含 id, level, title
    """
    path = []
    current_id = chapter_id
    
    while current_id:
        sql = """
            SELECT id, parent_id, level, title
            FROM chapters
            WHERE id = %s
        """
        chapter = fetch_one(sql, (current_id,))
        if chapter:
            path.insert(0, {
                'id': chapter['id'],
                'level': chapter['level'],
                'title': chapter['title']
            })
            current_id = chapter['parent_id']
        else:
            break
    
    return path


def get_chapter_images(chapter_id):
    """
    获取章节关联的图片
    """
    sql = """
        SELECT di.id, di.image_name, di.image_path, di.image_url, di.image_type
        FROM document_images di
        INNER JOIN chapter_images ci ON di.id = ci.image_id
        WHERE ci.chapter_id = %s
        ORDER BY ci.position_in_chapter
    """
    return fetch_all(sql, (chapter_id,))


def build_chapter_number_index(document_id):
    """
    从数据库构建章节编号索引
    
    根据数据库中的树形结构（parent_id, level, order_index）生成章节编号（如 1.4.1, 1.4.2）
    
    Args:
        document_id: 文档ID
        
    Returns:
        dict: {章节编号: 章节数据}，如 {'1.4.1': {...}, '1.4.2': {...}}
    """
    # 获取所有章节
    sql = """
        SELECT id, parent_id, level, order_index, title, content
        FROM chapters
        WHERE document_id = %s
        ORDER BY level, order_index
    """
    chapters = fetch_all(sql, (document_id,))
    
    if not chapters:
        return {}
    
    # 构建 id -> chapter 映射
    id_to_chapter = {c['id']: c for c in chapters}
    
    # 构建 parent_id -> children 映射
    children_map = {}
    root_chapters = []
    for c in chapters:
        parent_id = c['parent_id']
        if parent_id is None:
            root_chapters.append(c)
        else:
            if parent_id not in children_map:
                children_map[parent_id] = []
            children_map[parent_id].append(c)
    
    # 对每个层级的子章节按 order_index 排序
    for parent_id in children_map:
        children_map[parent_id].sort(key=lambda x: x['order_index'])
    root_chapters.sort(key=lambda x: x['order_index'])
    
    # 递归生成编号
    number_index = {}
    
    def assign_numbers(chapter_list, prefix=''):
        for idx, chapter in enumerate(chapter_list, 1):
            if prefix:
                number = f"{prefix}.{idx}"
            else:
                number = str(idx)
            
            chapter['_number'] = number
            number_index[number] = chapter
            
            # 递归处理子章节
            children = children_map.get(chapter['id'], [])
            if children:
                assign_numbers(children, number)
    
    assign_numbers(root_chapters)
    
    logger.info(f"文档 {document_id} 构建章节索引完成，共 {len(number_index)} 个章节")
    return number_index


def get_chapter_by_number_from_db(document_id, section_number):
    """
    从数据库获取指定编号的章节
    
    Args:
        document_id: 文档ID
        section_number: 章节编号，如 '1.4.1'
        
    Returns:
        章节数据或 None
    """
    number_index = build_chapter_number_index(document_id)
    return number_index.get(section_number)


def search_chapters_exact(query, document_ids=None, search_scope='title'):
    """
    精确匹配搜索章节
    
    Args:
        query: 查询关键词
        document_ids: 文档ID列表（可选）
        search_scope: 搜索范围 - 'title'(只匹配标题) 或 'content'(匹配标题和内容)
    """
    if search_scope == 'content':
        # 同时匹配 title 和 content
        if document_ids:
            placeholders = ','.join(['%s'] * len(document_ids))
            sql = f"""
                SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                       c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                       dpr.filename as doc_filename,
                       CASE WHEN c.title = %s THEN 'title' ELSE 'content' END as match_field
                FROM chapters c
                LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
                WHERE (c.title = %s OR c.content LIKE %s) AND c.document_id IN ({placeholders})
                ORDER BY c.document_id, c.level, c.order_index
            """
            params = [query, query, f'%{query}%'] + list(document_ids)
        else:
            sql = """
                SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                       c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                       dpr.filename as doc_filename,
                       CASE WHEN c.title = %s THEN 'title' ELSE 'content' END as match_field
                FROM chapters c
                LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
                WHERE c.title = %s OR c.content LIKE %s
                ORDER BY c.document_id, c.level, c.order_index
            """
            params = [query, query, f'%{query}%']
    else:
        # 只匹配 title
        if document_ids:
            placeholders = ','.join(['%s'] * len(document_ids))
            sql = f"""
                SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                       c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                       dpr.filename as doc_filename,
                       'title' as match_field
                FROM chapters c
                LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
                WHERE c.title = %s AND c.document_id IN ({placeholders})
                ORDER BY c.document_id, c.level, c.order_index
            """
            params = [query] + list(document_ids)
        else:
            sql = """
                SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                       c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                       dpr.filename as doc_filename,
                       'title' as match_field
                FROM chapters c
                LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
                WHERE c.title = %s
                ORDER BY c.document_id, c.level, c.order_index
            """
            params = [query]
    
    return fetch_all(sql, params)


def search_chapters_fuzzy(query, document_ids=None, search_scope='title'):
    """
    模糊匹配搜索章节
    去掉标点符号和特殊字符后进行匹配
    
    Args:
        query: 查询关键词
        document_ids: 文档ID列表（可选）
        search_scope: 搜索范围 - 'title'(只匹配标题) 或 'content'(匹配标题和内容)
    """
    cleaned_query = clean_text_for_fuzzy(query)
    if not cleaned_query:
        return []
    
    # 获取章节
    if document_ids:
        placeholders = ','.join(['%s'] * len(document_ids))
        sql = f"""
            SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                   c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                   dpr.filename as doc_filename
            FROM chapters c
            LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            WHERE c.document_id IN ({placeholders})
            ORDER BY c.document_id, c.level, c.order_index
        """
        chapters = fetch_all(sql, document_ids)
    else:
        sql = """
            SELECT c.id, c.document_id, c.parent_id, c.level, c.order_index,
                   c.title, c.content, c.style_name, c.font_size, c.is_bold, c.paragraph_index,
                   dpr.filename as doc_filename
            FROM chapters c
            LEFT JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
            ORDER BY c.document_id, c.level, c.order_index
        """
        chapters = fetch_all(sql)
    
    # 根据搜索范围进行匹配
    results = []
    for chapter in chapters:
        cleaned_title = clean_text_for_fuzzy(chapter.get('title', ''))
        
        if search_scope == 'content':
            # 同时匹配 title 和 content
            cleaned_content = clean_text_for_fuzzy(chapter.get('content', ''))
            if cleaned_query in cleaned_title:
                chapter['match_field'] = 'title'
                results.append(chapter)
            elif cleaned_query in cleaned_content:
                chapter['match_field'] = 'content'
                results.append(chapter)
        else:
            # 只匹配 title
            if cleaned_query in cleaned_title:
                chapter['match_field'] = 'title'
                results.append(chapter)
    
    return results


@chatdoc.route('/chatdoc', methods=['GET', 'POST'])
def chat_page():
    """渲染文档对话页面"""
    return render_template('chat_doc.html')


@chatdoc.route('/api/chat/documents', methods=['GET'])
def get_available_documents():
    """获取可用于查询的文档列表"""
    try:
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        sql = """
            SELECT doc_id, filename, status, created_at
            FROM doc_process_records
            WHERE username = %s AND status = 'completed'
            ORDER BY created_at DESC
        """
        documents = fetch_all(sql, (username,))
        
        return jsonify({
            'success': True,
            'data': documents
        })
    except Exception as e:
        logger.error(f"获取文档列表失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/search', methods=['POST'])
def search_content():
    """
    搜索章节内容
    支持精确匹配和模糊匹配
    支持只查询标题或内容匹配（标题+内容）
    """
    try:
        data = request.json
        query = data.get('query', '').strip()
        match_type = data.get('match_type', 'exact')  # exact 或 fuzzy
        document_ids = data.get('document_ids', [])  # 可选，指定文档范围
        include_children = data.get('include_children', True)  # 是否包含子章节
        search_scope = data.get('search_scope', 'title')  # title 或 content
        
        if not query:
            return jsonify({
                'success': False,
                'error': '查询内容不能为空'
            }), 400
        
        # 根据匹配类型搜索
        if match_type == 'exact':
            chapters = search_chapters_exact(query, document_ids if document_ids else None, search_scope)
        else:
            chapters = search_chapters_fuzzy(query, document_ids if document_ids else None, search_scope)
        
        if not chapters:
            return jsonify({
                'success': True,
                'data': {
                    'query': query,
                    'match_type': match_type,
                    'search_scope': search_scope,
                    'results': [],
                    'message': '未找到匹配的内容'
                }
            })
        
        # 按文档分组
        doc_results = {}
        for chapter in chapters:
            doc_id = chapter['document_id']
            if doc_id not in doc_results:
                doc_results[doc_id] = {
                    'document_id': doc_id,
                    'filename': chapter.get('doc_filename', '未知文档'),
                    'chapters': []
                }
            
            # 获取章节的层级路径
            chapter['path'] = get_chapter_path(chapter['id'])
            
            # 获取章节图片
            images = get_chapter_images(chapter['id'])
            chapter['images'] = images
            
            # 如果需要包含子章节
            if include_children:
                children = get_chapter_with_children(chapter['id'], doc_id)
                # 去掉第一个（当前章节本身，避免重复）
                if len(children) > 1:
                    chapter['children'] = children[1:]
            
            doc_results[doc_id]['chapters'].append(chapter)
        
        return jsonify({
            'success': True,
            'data': {
                'query': query,
                'match_type': match_type,
                'search_scope': search_scope,
                'results': list(doc_results.values()),
                'total_matches': len(chapters),
                'document_count': len(doc_results)
            }
        })
        
    except Exception as e:
        logger.error(f"搜索失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/batch-search', methods=['POST'])
def batch_search():
    """
    批量搜索 - 支持多行查询
    每行一个查询条件
    """
    try:
        data = request.json
        queries = data.get('queries', [])  # 查询列表
        match_type = data.get('match_type', 'exact')
        document_ids = data.get('document_ids', [])
        include_children = data.get('include_children', True)
        search_scope = data.get('search_scope', 'title')  # title 或 content
        
        if not queries:
            return jsonify({
                'success': False,
                'error': '查询内容不能为空'
            }), 400
        
        all_results = []
        
        for query in queries:
            query = query.strip()
            if not query:
                continue
            
            # 根据匹配类型搜索
            if match_type == 'exact':
                chapters = search_chapters_exact(query, document_ids if document_ids else None, search_scope)
            else:
                chapters = search_chapters_fuzzy(query, document_ids if document_ids else None, search_scope)
            
            # 按文档分组
            doc_results = {}
            for chapter in chapters:
                doc_id = chapter['document_id']
                if doc_id not in doc_results:
                    doc_results[doc_id] = {
                        'document_id': doc_id,
                        'filename': chapter.get('doc_filename', '未知文档'),
                        'chapters': []
                    }
                
                # 获取章节的层级路径
                chapter['path'] = get_chapter_path(chapter['id'])
                
                # 获取章节图片
                images = get_chapter_images(chapter['id'])
                chapter['images'] = images
                
                # 如果需要包含子章节
                if include_children:
                    children = get_chapter_with_children(chapter['id'], doc_id)
                    if len(children) > 1:
                        chapter['children'] = children[1:]
                
                doc_results[doc_id]['chapters'].append(chapter)
            
            all_results.append({
                'query': query,
                'results': list(doc_results.values()),
                'match_count': len(chapters)
            })
        
        return jsonify({
            'success': True,
            'data': {
                'match_type': match_type,
                'search_scope': search_scope,
                'batch_results': all_results,
                'total_queries': len(queries)
            }
        })
        
    except Exception as e:
        logger.error(f"批量搜索失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/export-word', methods=['POST'])
def export_to_word():
    """
    导出搜索结果到Word文档
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        data = request.json
        results = data.get('results', [])  # 搜索结果
        title = data.get('title', '文档查询结果')
        
        if not results:
            return jsonify({
                'success': False,
                'error': '没有可导出的内容'
            }), 400
        
        # 创建Word文档
        doc = Document()
        
        # 添加标题
        heading = doc.add_heading(title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加生成时间
        doc.add_paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        
        # 全局章节计数器
        global_chapter_index = 0
        
        # 遍历结果
        for query_result in results:
            query = query_result.get('query', '')
            doc_results = query_result.get('results', [])
            
            if not doc_results:
                continue
            
            for doc_result in doc_results:
                filename = doc_result.get('filename', '未知文档')
                chapters = doc_result.get('chapters', [])
                
                for chapter in chapters:
                    global_chapter_index += 1
                    # 添加章节（带路径和序号）
                    add_chapter_with_path(doc, chapter, global_chapter_index, filename)
        
        # 保存到临时文件
        temp_dir = tempfile.gettempdir()
        filename = f"查询结果_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = os.path.join(temp_dir, filename)
        doc.save(filepath)
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except ImportError:
        logger.error("python-docx 库未安装")
        return jsonify({
            'success': False,
            'error': '服务器缺少 python-docx 库，无法导出Word文档'
        }), 500
    except Exception as e:
        logger.error(f"导出Word失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


def add_chapter_with_path(doc, chapter, index, filename):
    """
    添加章节到Word文档，显示完整路径和带序号的标题
    
    格式:
    [需删除]技术规范专用部分 -> 原生分布式数据库 -> 总体要求 -> 标题
    1 标题
      内容...
    
    Args:
        doc: Word文档对象
        chapter: 章节数据（包含 path, title, content, images, children）
        index: 章节序号（1, 2, 3...）
        filename: 来源文件名
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    title = chapter.get('title', '')
    content = chapter.get('content', '')
    images = chapter.get('images', [])
    children = chapter.get('children', [])
    path = chapter.get('path', [])
    
    # 1. 构建完整路径字符串（如: 技术规范专用部分 -> 原生分布式数据库 -> 总体要求 -> 标题）
    if path:
        path_titles = [p.get('title', '') for p in path]
        path_str = ' -> '.join(path_titles)
    else:
        path_str = title
    
    # 添加路径行（灰色小字，标记为[需删除]-文件名-路径）
    path_para = doc.add_paragraph()
    path_run = path_para.add_run(f"[需删除]-{filename}-{path_str}")
    path_run.font.size = Pt(9)
    path_run.font.color.rgb = RGBColor(128, 128, 128)  # 灰色
    
    # 2. 添加带序号的标题（如: 1 ▲拥有自主知识产权...）
    heading = doc.add_heading(f"{index} {title}", level=1)
    
    # 3. 添加内容
    if content:
        _add_content_with_tables(doc, content, images)
    else:
        # 没有内容时仍需添加图片
        for img in images:
            img_path = img.get('image_path', '')
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logger.warning(f"添加图片失败: {img_path}, 错误: {e}")
    
    # 4. 递归添加子章节（带层级序号 1.1, 1.1.1 等）
    if children:
        child_counters = [0] * 10
        for child in children:
            add_child_chapter_with_numbering(doc, child, prefix=str(index), 
                                              counters=child_counters, depth=0)
    
    # 添加空行分隔
    doc.add_paragraph()


def add_child_chapter_with_numbering(doc, chapter, prefix, counters, depth):
    """
    递归添加子章节，带层级序号（如 1.1, 1.1.1）
    """
    from docx.shared import Inches
    
    title = chapter.get('title', '')
    content = chapter.get('content', '')
    images = chapter.get('images', [])
    children = chapter.get('children', [])
    
    # 增加当前层级计数
    counters[depth] += 1
    # 重置更深层级的计数
    for i in range(depth + 1, len(counters)):
        counters[i] = 0
    
    # 构建完整序号（如 1.1.1）
    number_parts = [str(counters[i]) for i in range(depth + 1)]
    full_number = f"{prefix}.{'.'.join(number_parts)}"
    
    # 添加章节标题（带序号）
    if title:
        heading_level = min(depth + 2, 9)  # 从 level 2 开始
        doc.add_heading(f"{full_number} {title}", level=heading_level)
    
    # 添加内容
    if content:
        _add_content_with_tables(doc, content, images)
    else:
        for img in images:
            img_path = img.get('image_path', '')
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logger.warning(f"添加图片失败: {img_path}, 错误: {e}")
    
    # 递归添加子章节
    for child in children:
        add_child_chapter_with_numbering(doc, child, prefix=prefix, 
                                          counters=counters, depth=depth + 1)


def add_chapter_to_doc_with_numbering(doc, chapter, level=3, prefix="", counters=None, depth=0):
    """
    递归添加章节到Word文档，带层级序号（如 1.1.1, 1.1.2）
    
    Args:
        doc: Word文档对象
        chapter: 章节数据
        level: Word标题层级
        prefix: 序号前缀（如 "1.1"）
        counters: 各层级计数器列表
        depth: 当前递归深度
    """
    from docx.shared import Inches
    import re
    
    if counters is None:
        counters = [0] * 10
    
    title = chapter.get('title', '')
    content = chapter.get('content', '')
    images = chapter.get('images', [])
    children = chapter.get('children', [])
    
    # 增加当前层级计数
    counters[depth] += 1
    # 重置更深层级的计数
    for i in range(depth + 1, len(counters)):
        counters[i] = 0
    
    # 构建完整序号（如 1.1.1）
    number_parts = [str(counters[i]) for i in range(depth + 1)]
    full_number = f"{prefix}.{'.'.join(number_parts)}" if prefix else '.'.join(number_parts)
    
    # 添加章节标题（带序号）
    if title:
        heading_level = min(level, 9)
        doc.add_heading(f"{full_number} {title}", level=heading_level)
    
    # 添加内容（处理 [表格]...[/表格] 标签）
    if content:
        _add_content_with_tables(doc, content, images)
    else:
        # 没有内容时仍需添加图片
        for img in images:
            img_path = img.get('image_path', '')
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logger.warning(f"添加图片失败: {img_path}, 错误: {e}")
    
    # 递归添加子章节
    for child in children:
        add_chapter_to_doc_with_numbering(doc, child, level=level + 1, 
                                          prefix=prefix, counters=counters, depth=depth + 1)


def add_chapter_to_doc(doc, chapter, level=3):
    """
    递归添加章节到Word文档（旧版本，保留兼容性）
    """
    from docx.shared import Inches
    import re
    
    title = chapter.get('title', '')
    content = chapter.get('content', '')
    images = chapter.get('images', [])
    children = chapter.get('children', [])
    
    # 添加章节标题
    if title:
        # 限制最大层级为9
        heading_level = min(level, 9)
        doc.add_heading(title, level=heading_level)
    
    # 添加内容（处理 [表格]...[/表格] 标签）
    if content:
        _add_content_with_tables(doc, content, images)
    else:
        # 没有内容时仍需添加图片
        for img in images:
            img_path = img.get('image_path', '')
            if img_path and os.path.exists(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(5))
                except Exception as e:
                    logger.warning(f"添加图片失败: {img_path}, 错误: {e}")
    
    # 递归添加子章节
    for child in children:
        add_chapter_to_doc(doc, child, level=level + 1)


def _add_content_with_tables(doc, content, images):
    """
    将内容添加到 Word 文档，遇到 [表格]...[/表格] 就把里面的内容放进单行单列表格
    同时处理 {{IMAGE_ID_xxx}} 占位符替换为图片
    """
    from docx.shared import Inches, Pt
    from docx.oxml.ns import qn
    import re
    
    # 构建图片ID映射
    img_map = {}
    for img in images:
        img_id = img.get('id')
        if img_id:
            img_map[str(img_id)] = img.get('image_path', '')
    
    # 正则：匹配 [表格]...[/表格] 块
    table_pattern = re.compile(r'\[表格\](.*?)\[/表格\]', re.DOTALL)
    # 正则：匹配图片占位符
    image_pattern = re.compile(r'\{\{IMAGE_ID_(\d+)\}\}')
    
    last_end = 0
    for match in table_pattern.finditer(content):
        # 添加表格之前的普通文本
        before_text = content[last_end:match.start()]
        if before_text.strip():
            _add_text_with_images(doc, before_text, img_map, image_pattern)
        
        # 表格内容
        table_content = match.group(1).strip()
        if table_content:
            # 创建单行单列表格
            table = doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            cell = table.cell(0, 0)
            # 把表格内容（可能含图片占位符）写入单元格
            _add_cell_content_with_images(cell, table_content, img_map, image_pattern)
        
        last_end = match.end()
    
    # 添加最后剩余的普通文本
    remaining = content[last_end:]
    if remaining.strip():
        _add_text_with_images(doc, remaining, img_map, image_pattern)


def _add_text_with_images(doc, text, img_map, image_pattern):
    """
    将文本添加到文档，遇到 {{IMAGE_ID_xxx}} 替换为图片
    """
    from docx.shared import Inches
    
    last_end = 0
    for match in image_pattern.finditer(text):
        # 图片之前的文字
        before = text[last_end:match.start()]
        if before.strip():
            doc.add_paragraph(before.strip())
        
        # 插入图片
        img_id = match.group(1)
        img_path = img_map.get(img_id, '')
        if img_path and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(5))
            except Exception as e:
                logger.warning(f"添加图片失败: {img_path}, 错误: {e}")
        
        last_end = match.end()
    
    # 剩余文字
    remaining = text[last_end:]
    if remaining.strip():
        doc.add_paragraph(remaining.strip())


def _add_cell_content_with_images(cell, text, img_map, image_pattern):
    """
    将文本添加到表格单元格，遇到 {{IMAGE_ID_xxx}} 替换为图片
    """
    from docx.shared import Inches
    
    # 清空单元格默认段落
    cell.text = ''
    
    last_end = 0
    for match in image_pattern.finditer(text):
        # 图片之前的文字
        before = text[last_end:match.start()]
        if before.strip():
            p = cell.add_paragraph(before.strip())
        
        # 插入图片到单元格
        img_id = match.group(1)
        img_path = img_map.get(img_id, '')
        if img_path and os.path.exists(img_path):
            try:
                p = cell.add_paragraph()
                run = p.add_run()
                run.add_picture(img_path, width=Inches(4.5))
            except Exception as e:
                logger.warning(f"添加图片到表格失败: {img_path}, 错误: {e}")
        
        last_end = match.end()
    
    # 剩余文字
    remaining = text[last_end:]
    if remaining.strip():
        cell.add_paragraph(remaining.strip())



@chatdoc.route('/api/chat/upload-txt', methods=['POST'])
def upload_txt_for_search():
    """
    上传TXT或DOCX文件进行批量查询
    TXT: 每行一个查询条件
    DOCX: 解析段落和表格内容作为查询条件
    """
    try:
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': '未上传文件'
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': '未选择文件'
            }), 400
        
        filename = file.filename.lower()
        if not (filename.endswith('.txt') or filename.endswith('.docx')):
            return jsonify({
                'success': False,
                'error': '只支持TXT和DOCX文件'
            }), 400
        
        lines = []
        
        if filename.endswith('.txt'):
            # 读取TXT文件内容
            content = file.read().decode('utf-8')
            lines = [line.strip() for line in content.split('\n') if line.strip()]
        else:
            # 读取DOCX文件内容
            import tempfile
            import os
            from docx import Document
            
            # 保存临时文件
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
                file.save(tmp.name)
                tmp_path = tmp.name
            
            try:
                doc = Document(tmp_path)
                
                # 提取段落
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text and len(text) >= 2:  # 至少2个字符
                        lines.append(text)
                
                # 提取表格内容
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text = cell.text.strip()
                            if text and len(text) >= 2 and text not in lines:
                                lines.append(text)
            finally:
                # 清理临时文件
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)
        
        if not lines:
            return jsonify({
                'success': False,
                'error': '文件内容为空'
            }), 400
        
        return jsonify({
            'success': True,
            'data': {
                'queries': lines,
                'count': len(lines)
            }
        })
        
    except Exception as e:
        logger.error(f"上传TXT文件失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/select-document', methods=['POST'])
def select_document_for_conflict():
    """
    当多个文档都匹配时，用户选择使用哪个文档的结果
    """
    try:
        data = request.json
        selected_doc_id = data.get('document_id')
        query = data.get('query')
        match_type = data.get('match_type', 'exact')
        include_children = data.get('include_children', True)
        
        if not selected_doc_id or not query:
            return jsonify({
                'success': False,
                'error': '参数不完整'
            }), 400
        
        # 只在选定的文档中搜索
        if match_type == 'exact':
            chapters = search_chapters_exact(query, [selected_doc_id])
        else:
            chapters = search_chapters_fuzzy(query, [selected_doc_id])
        
        results = []
        for chapter in chapters:
            # 获取章节的层级路径
            chapter['path'] = get_chapter_path(chapter['id'])
            
            images = get_chapter_images(chapter['id'])
            chapter['images'] = images
            
            if include_children:
                children = get_chapter_with_children(chapter['id'], selected_doc_id)
                if len(children) > 1:
                    chapter['children'] = children[1:]
            
            results.append(chapter)
        
        return jsonify({
            'success': True,
            'data': {
                'query': query,
                'document_id': selected_doc_id,
                'chapters': results
            }
        })
        
    except Exception as e:
        logger.error(f"选择文档失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ============== LLM 匹配相关 ==============

@chatdoc.route('/api/chat/llm-search', methods=['POST'])
def llm_search():
    """
    LLM智能匹配搜索
    
    支持单条查询，先在文档中精确匹配，匹配不上则语义匹配，都不行则网络搜索
    """
    try:
        from .requirement_analyzer import get_requirement_analyzer
        from .llm_config import LLMConfigManager
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        query = data.get('query', '').strip()
        document_ids = data.get('document_ids', [])
        enable_web_search = data.get('enable_web_search', True)
        llm_config_id = data.get('llm_config_id')
        
        if not query:
            return jsonify({
                'success': False,
                'error': '查询内容不能为空'
            }), 400
        
        # 检查LLM配置
        if llm_config_id:
            config = LLMConfigManager.get_config(llm_config_id)
        else:
            config = LLMConfigManager.get_default_config()
        
        if not config:
            return jsonify({
                'success': False,
                'error': '未配置LLM，请先在设置中配置大模型'
            }), 400
        
        # 创建分析器并执行分析
        analyzer = get_requirement_analyzer(llm_config_id)
        result = analyzer.analyze_requirement(
            query, 
            username, 
            document_ids if document_ids else None,
            enable_web_search
        )
        
        return jsonify({
            'success': True,
            'data': result
        })
        
    except Exception as e:
        logger.error(f"LLM搜索失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/analyze-file', methods=['POST'])
def analyze_uploaded_file():
    """
    分析上传的需求文件
    
    支持上传 .docx 或 .txt 文件，解析其中的需求并逐条分析
    """
    try:
        from .requirement_analyzer import get_requirement_analyzer
        from .llm_config import LLMConfigManager
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        if 'file' not in request.files:
            return jsonify({
                'success': False,
                'error': '未上传文件'
            }), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({
                'success': False,
                'error': '未选择文件'
            }), 400
        
        # 检查文件类型
        allowed_extensions = {'.docx', '.txt'}
        ext = os.path.splitext(file.filename)[1].lower()
        if ext not in allowed_extensions:
            return jsonify({
                'success': False,
                'error': f'不支持的文件格式，仅支持: {", ".join(allowed_extensions)}'
            }), 400
        
        # 获取其他参数
        document_ids = request.form.getlist('document_ids[]')
        enable_web_search = request.form.get('enable_web_search', 'true').lower() == 'true'
        llm_config_id = request.form.get('llm_config_id')
        section_filter = request.form.get('section_filter')  # 章节过滤，如 "1.4.1,1.4.2"
        
        if llm_config_id:
            llm_config_id = int(llm_config_id)
        
        # 解析章节过滤参数
        section_filter_list = None
        if section_filter:
            section_filter_list = [s.strip() for s in section_filter.replace('、', ',').split(',') if s.strip()]
        
        # 检查LLM配置
        if llm_config_id:
            config = LLMConfigManager.get_config(llm_config_id)
        else:
            config = LLMConfigManager.get_default_config()
        
        if not config:
            return jsonify({
                'success': False,
                'error': '未配置LLM，请先在设置中配置大模型'
            }), 400
        
        # 保存上传的文件到临时目录
        import tempfile
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, f"req_upload_{username}_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}")
        file.save(temp_path)
        
        try:
            # 创建分析器
            analyzer = get_requirement_analyzer(llm_config_id)
            
            # 解析需求（应用章节过滤）
            requirements = analyzer.parse_requirements_from_file(temp_path, section_filter=section_filter_list)
            
            if not requirements:
                error_msg = '未能从文件中解析出需求'
                if section_filter_list:
                    error_msg += f'（过滤章节: {", ".join(section_filter_list)}）'
                return jsonify({
                    'success': False,
                    'error': error_msg
                }), 400
            
            # 返回解析出的需求，让前端确认后再分析
            return jsonify({
                'success': True,
                'data': {
                    'requirements': requirements,
                    'count': len(requirements),
                    'temp_file': temp_path,
                    'filename': file.filename,
                    'section_filter': section_filter_list
                }
            })
            
        except Exception as e:
            # 清理临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
            raise e
        
    except Exception as e:
        logger.error(f"分析上传文件失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/analyze-requirements', methods=['POST'])
def analyze_requirements():
    """
    批量分析需求
    
    对解析出的需求列表进行逐条分析
    """
    try:
        from .requirement_analyzer import get_requirement_analyzer
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        requirements = data.get('requirements', [])
        document_ids = data.get('document_ids', [])
        enable_web_search = data.get('enable_web_search', True)
        llm_config_id = data.get('llm_config_id')
        temp_file = data.get('temp_file')  # 临时文件路径
        
        if not requirements:
            return jsonify({
                'success': False,
                'error': '需求列表为空'
            }), 400
        
        # 创建分析器
        analyzer = get_requirement_analyzer(llm_config_id)
        
        # 批量分析
        results = analyzer.analyze_requirements_batch(
            requirements,
            username,
            document_ids if document_ids else None,
            enable_web_search
        )
        
        # 清理临时文件
        if temp_file and os.path.exists(temp_file):
            try:
                os.remove(temp_file)
            except:
                pass
        
        return jsonify({
            'success': True,
            'data': {
                'results': results,
                'total': len(results),
                'summary': {
                    'exact': sum(1 for r in results if r.get('match_type') == 'exact'),
                    'semantic': sum(1 for r in results if r.get('match_type') == 'semantic'),
                    'web': sum(1 for r in results if r.get('match_type') == 'web'),
                    'llm_generated': sum(1 for r in results if r.get('match_type') == 'llm_generated'),
                    'none': sum(1 for r in results if r.get('match_type') in ['none', 'error'])
                }
            }
        })
        
    except Exception as e:
        logger.error(f"批量分析需求失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/export-llm-results', methods=['POST'])
def export_llm_results():
    """
    导出LLM分析结果为Word文档
    """
    try:
        from .requirement_analyzer import get_requirement_analyzer
        
        data = request.json
        results = data.get('results', [])
        title = data.get('title', '需求分析报告')
        
        if not results:
            return jsonify({
                'success': False,
                'error': '没有可导出的结果'
            }), 400
        
        # 创建分析器并导出
        analyzer = get_requirement_analyzer()
        filepath, filename = analyzer.export_to_word(results, title)
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"导出LLM结果失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/llm-configs', methods=['GET'])
def get_llm_configs():
    """获取可用的LLM配置列表"""
    try:
        from .llm_config import LLMConfigManager
        
        configs = LLMConfigManager.list_configs()
        
        # 隐藏敏感信息
        for config in configs:
            if config.get('api_key'):
                config['api_key'] = config['api_key'][:8] + '****'
        
        return jsonify({
            'success': True,
            'data': configs
        })
        
    except Exception as e:
        logger.error(f"获取LLM配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/llm-config/check', methods=['GET'])
def check_llm_config():
    """检查是否有可用的LLM配置"""
    try:
        from .llm_config import LLMConfigManager
        
        config = LLMConfigManager.get_default_config()
        
        return jsonify({
            'success': True,
            'data': {
                'has_config': config is not None,
                'config_name': config.get('config_name') if config else None,
                'model_type': config.get('model_type') if config else None,
                'model_name': config.get('model_name') if config else None
            }
        })
        
    except Exception as e:
        logger.error(f"检查LLM配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== 招标文档智能应答API ====================

@chatdoc.route('/api/chat/parse-bid-instruction', methods=['POST'])
def parse_bid_instruction():
    """
    解析用户的招标作答指令
    
    请求参数:
        instruction: 用户指令，如 "针对文档中的1.4.1,1.4.2作答"
        doc_id: 文档ID（可选，用于从数据库获取已解析的章节）
        file_path: 文档路径（可选，用于直接解析文件）
    
    返回:
        解析出的章节编号和对应的技术要求
    """
    try:
        from .bid_document_parser import UserInstructionParser, BidResponseGenerator
        
        data = request.json
        instruction = data.get('instruction', '').strip()
        doc_id = data.get('doc_id')
        file_path = data.get('file_path')
        
        if not instruction:
            return jsonify({
                'success': False,
                'error': '请输入作答指令'
            }), 400
        
        # 解析用户指令
        parser = UserInstructionParser()
        instruction_info = parser.parse_instruction(instruction)
        
        if not instruction_info['parsed']:
            return jsonify({
                'success': False,
                'error': '无法解析指令，请使用格式如："针对文档中的1.4.1,1.4.2作答"',
                'instruction_info': instruction_info
            })
        
        # 如果提供了doc_id，尝试获取章节详情
        sections = {}
        if doc_id:
            generator = BidResponseGenerator()
            result = generator.process_instruction(instruction, doc_id=doc_id)
            if result['success']:
                sections = result['sections']
        
        return jsonify({
            'success': True,
            'data': {
                'instruction_info': instruction_info,
                'section_numbers': instruction_info['section_numbers'],
                'sections': sections
            }
        })
        
    except Exception as e:
        logger.error(f"解析招标指令失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/parse-bid-document', methods=['POST'])
def parse_bid_document_api():
    """
    解析招标文档的章节结构
    
    请求参数:
        doc_id: 文档ID（必需）
        
    返回:
        文档的完整章节结构
    """
    try:
        from .bid_document_parser import BidDocumentParser
        
        data = request.json
        doc_id = data.get('doc_id')
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '请提供文档ID'
            }), 400
        
        # 获取文档路径
        sql = """
            SELECT dpr.doc_id, dpr.filename, dpr.final_path, dpr.status
            FROM doc_process_records dpr
            WHERE dpr.doc_id = %s
        """
        doc_info = fetch_one(sql, (doc_id,))
        
        if not doc_info:
            return jsonify({
                'success': False,
                'error': '文档不存在'
            }), 404
        
        if doc_info['status'] != 'completed':
            return jsonify({
                'success': False,
                'error': '文档尚未处理完成'
            }), 400
        
        file_path = doc_info['final_path']
        if not file_path or not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': '文档文件不存在'
            }), 404
        
        # 解析文档结构
        parser = BidDocumentParser(file_path)
        structure = parser.parse_document_structure()
        
        # 转换为列表格式便于前端展示
        sections_list = []
        for num, section in structure.items():
            sections_list.append({
                'number': num,
                'title': section.get('title', ''),
                'level': section.get('level', 1),
                'requirements_count': len(section.get('requirements', [])),
                'has_tables': len(section.get('tables', [])) > 0,
                'children': section.get('children', [])
            })
        
        # 按章节编号排序
        sections_list.sort(key=lambda x: [int(p) if p.isdigit() else p for p in x['number'].split('.')])
        
        return jsonify({
            'success': True,
            'data': {
                'doc_id': doc_id,
                'filename': doc_info['filename'],
                'sections': sections_list,
                'total_sections': len(sections_list)
            }
        })
        
    except Exception as e:
        logger.error(f"解析招标文档失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/get-section-requirements', methods=['POST'])
def get_section_requirements():
    """
    获取指定章节的技术要求列表
    
    请求参数:
        doc_id: 文档ID（必需）
        section_numbers: 章节编号列表，如 ['1.4.1', '1.4.2']
    
    返回:
        各章节的技术要求详情（解析每一条具体要求）
    """
    try:
        from .bid_document_parser import BidDocumentParser
        
        data = request.json
        doc_id = data.get('doc_id')
        section_numbers = data.get('section_numbers', [])
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '请提供文档ID'
            }), 400
        
        if not section_numbers:
            return jsonify({
                'success': False,
                'error': '请提供章节编号'
            }), 400
        
        # 获取文档信息
        sql = """
            SELECT dpr.final_path, dpr.filename
            FROM doc_process_records dpr
            WHERE dpr.doc_id = %s AND dpr.status = 'completed'
        """
        doc_info = fetch_one(sql, (doc_id,))
        
        if not doc_info:
            return jsonify({
                'success': False,
                'error': '文档不存在或未处理完成'
            }), 404
        
        file_path = doc_info['final_path']
        
        # 检查文件是否存在
        import os
        if not file_path or not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': f'文档文件不存在: {file_path}'
            }), 404
        
        # 【核心修改】直接解析文档文件，而不是从数据库查
        logger.info(f"解析文档: {file_path}")
        parser = BidDocumentParser(file_path)
        parser.parse_document_structure()
        
        # 打印解析到的所有章节编号
        all_section_numbers = list(parser.section_index.keys())
        logger.info(f"文档解析到的章节编号: {all_section_numbers}")
        
        sections_data = {}
        total_requirements = 0
        missing_sections = []
        
        for num in section_numbers:
            logger.info(f"查找章节 {num}...")
            
            # 从解析结果中查找章节
            section = parser.get_section_by_number(num)
            
            if section:
                logger.info(f"章节 {num} 找到: title={section.get('title', '')}, requirements={len(section.get('requirements', []))}")
                
                # 使用 BidDocumentParser 的方法获取所有需求
                requirements = parser.get_all_requirements_from_section(num)
                
                logger.info(f"章节 {num} 获取到 {len(requirements)} 条需求")
                
                sections_data[num] = {
                    'number': num,
                    'title': section.get('title', ''),
                    'content': section.get('content', ''),
                    'requirements': requirements,
                    'tables': section.get('tables', []),
                    'requirements_count': len(requirements)
                }
                total_requirements += len(requirements)
            else:
                logger.warning(f"章节 {num} 在文档中未找到，可用章节: {all_section_numbers}")
                missing_sections.append(num)
        
        return jsonify({
            'success': True,
            'data': {
                'doc_id': doc_id,
                'filename': doc_info['filename'],
                'sections': sections_data,
                'total_requirements': total_requirements,
                'found_sections': list(sections_data.keys()),
                'missing_sections': missing_sections,
                'available_sections': all_section_numbers  # 返回所有可用章节供调试
            }
        })
        
    except Exception as e:
        logger.error(f"获取章节需求失败: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/answer-bid-requirements', methods=['POST'])
def answer_bid_requirements():
    """
    对招标文档的技术要求进行智能作答
    
    请求参数:
        doc_id: 招标文档ID（必需）
        section_numbers: 要作答的章节编号列表
        knowledge_doc_ids: 知识库文档ID列表（用于匹配答案）
        llm_config_id: LLM配置ID（可选）
        enable_web_search: 是否启用网络搜索（默认True）
        export_format: 导出格式 'json' | 'word' | 'word_table'（默认json）
    
    返回:
        各技术要求的作答结果（对每一条具体要求分别作答）
    """
    try:
        from .bid_document_parser import BidDocumentParser, BidAnswerGenerator
        import os
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        doc_id = data.get('doc_id')
        section_numbers = data.get('section_numbers', [])
        knowledge_doc_ids = data.get('knowledge_doc_ids', [])
        llm_config_id = data.get('llm_config_id')
        enable_web_search = data.get('enable_web_search', True)
        export_format = data.get('export_format', 'json')
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '请提供文档ID'
            }), 400
        
        if not section_numbers:
            return jsonify({
                'success': False,
                'error': '请提供要作答的章节编号'
            }), 400
        
        # 获取文档路径
        sql = """
            SELECT dpr.final_path, dpr.filename
            FROM doc_process_records dpr
            WHERE dpr.doc_id = %s AND dpr.status = 'completed'
        """
        doc_info = fetch_one(sql, (doc_id,))
        
        if not doc_info or not doc_info['final_path']:
            return jsonify({
                'success': False,
                'error': '文档不存在或未处理完成'
            }), 404
        
        file_path = doc_info['final_path']
        
        # 检查文件是否存在
        if not file_path or not os.path.exists(file_path):
            return jsonify({
                'success': False,
                'error': f'文档文件不存在: {file_path}'
            }), 404
        
        # 【核心修改】直接解析文档文件
        logger.info(f"answer_bid_requirements - 解析文档: {file_path}")
        parser = BidDocumentParser(file_path)
        parser.parse_document_structure()
        
        all_section_numbers = list(parser.section_index.keys())
        logger.info(f"文档解析到的章节编号: {all_section_numbers}")
        
        all_requirements = []
        
        for num in section_numbers:
            logger.info(f"从文档获取章节 {num}...")
            
            # 使用 BidDocumentParser 获取章节的所有需求
            requirements = parser.get_all_requirements_from_section(num)
            
            if requirements:
                logger.info(f"章节 {num} 获取到 {len(requirements)} 条需求")
                for req in requirements:
                    all_requirements.append({
                        'section_number': num,
                        'section_title': req.get('section_title', ''),
                        'index': req.get('index', ''),
                        'content': req.get('text', ''),
                        'spec': req.get('spec', ''),
                        'type': req.get('type', 'list')
                    })
            else:
                logger.warning(f"章节 {num} 未找到，可用章节: {all_section_numbers}")
        
        logger.info(f"从文档获取到 {len(all_requirements)} 条需求")
        
        if not all_requirements:
            return jsonify({
                'success': False,
                'error': f'未找到需要作答的技术要求，可用章节: {all_section_numbers}'
            }), 400
        
        logger.info(f"开始处理 {len(all_requirements)} 条技术要求")
        
        # 使用新的作答生成器，对每一条要求分别执行匹配
        generator = BidAnswerGenerator(llm_config_id)
        results = generator.answer_requirements(
            all_requirements,
            username,
            knowledge_doc_ids if knowledge_doc_ids else None,
            enable_web_search
        )
        
        # 统计
        summary = {
            'total': len(results),
            'exact': sum(1 for r in results if r['match_type'] == 'exact'),
            'semantic': sum(1 for r in results if r['match_type'] == 'semantic'),
            'web': sum(1 for r in results if r['match_type'] == 'web'),
            'llm_generated': sum(1 for r in results if r['match_type'] == 'llm_generated'),
            'none': sum(1 for r in results if r['match_type'] in ['none', 'error'])
        }
        
        # 如果需要导出Word
        if export_format in ['word', 'word_table']:
            bid_doc_info = {'filename': doc_info['filename'], 'doc_id': doc_id}
            format_type = 'table' if export_format == 'word_table' else 'default'
            filepath, filename = generator.export_to_word(
                results, 
                title='招标技术要求应答书',
                bid_doc_info=bid_doc_info
            ) if format_type == 'default' else generator.export_to_word_table_format(
                results,
                title='招标技术要求应答表',
                bid_doc_info=bid_doc_info
            )
            
            return send_file(
                filepath,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        
        return jsonify({
            'success': True,
            'data': {
                'results': results,
                'summary': summary,
                'doc_info': {
                    'doc_id': doc_id,
                    'filename': doc_info['filename']
                }
            }
        })
        
    except Exception as e:
        logger.error(f"招标需求作答失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/export-bid-answers', methods=['POST'])
def export_bid_answers():
    """
    导出招标作答结果为Word文档
    
    请求参数:
        results: 作答结果列表
        title: 文档标题
        doc_info: 文档信息
        format_type: 格式类型 'default' | 'table'
    """
    try:
        from .bid_document_parser import BidAnswerGenerator
        
        data = request.json
        results = data.get('results', [])
        title = data.get('title', '招标技术要求应答书')
        doc_info = data.get('doc_info')
        format_type = data.get('format_type', 'default')
        
        if not results:
            return jsonify({
                'success': False,
                'error': '没有可导出的结果'
            }), 400
        
        generator = BidAnswerGenerator()
        
        if format_type == 'table':
            filepath, filename = generator.export_to_word_table_format(results, title, doc_info)
        else:
            filepath, filename = generator.export_to_word(results, title, doc_info)
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        logger.error(f"导出招标作答失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== Embedding 向量化相关 API ====================

@chatdoc.route('/api/chat/embedding/vectorize', methods=['POST'])
def vectorize_document():
    """
    手动触发文档向量化
    
    请求参数:
        doc_id: 文档ID
        config_id: Embedding配置ID（可选）
    """
    try:
        from .embedding_service import get_vector_store
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        doc_id = data.get('doc_id')
        config_id = data.get('config_id')
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '请提供文档ID'
            }), 400
        
        # 验证文档权限
        doc_sql = """
            SELECT doc_id, filename, status 
            FROM doc_process_records 
            WHERE doc_id = %s AND username = %s
        """
        doc = fetch_one(doc_sql, (doc_id, username))
        
        if not doc:
            return jsonify({
                'success': False,
                'error': '文档不存在或无权访问'
            }), 404
        
        if doc['status'] != 'completed':
            return jsonify({
                'success': False,
                'error': '文档尚未处理完成'
            }), 400
        
        # 获取章节
        chapters_sql = """
            SELECT id, title, content, level, parent_id
            FROM chapters
            WHERE document_id = %s
        """
        chapters = fetch_all(chapters_sql, (doc_id,))
        
        if not chapters:
            return jsonify({
                'success': False,
                'error': '文档没有章节内容'
            }), 400
        
        # 执行向量化
        vector_store = get_vector_store(config_id)
        count = vector_store.add_document_embeddings(doc_id, chapters)
        
        return jsonify({
            'success': True,
            'data': {
                'doc_id': doc_id,
                'filename': doc['filename'],
                'vectorized_count': count,
                'total_chapters': len(chapters)
            }
        })
        
    except Exception as e:
        logger.error(f"文档向量化失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding/search', methods=['POST'])
def vector_search():
    """
    向量相似度搜索
    
    请求参数:
        query: 查询文本
        document_ids: 限定的文档ID列表（可选）
        top_k: 返回结果数量（默认10）
        threshold: 相似度阈值（默认0.5）
    """
    try:
        from .embedding_service import get_vector_store
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        query = data.get('query', '')
        document_ids = data.get('document_ids', [])
        top_k = data.get('top_k', 10)
        threshold = data.get('threshold', 0.5)
        
        if not query:
            return jsonify({
                'success': False,
                'error': '请提供查询文本'
            }), 400
        
        # 如果没有指定文档，获取用户的所有文档
        if not document_ids:
            docs_sql = """
                SELECT doc_id FROM doc_process_records 
                WHERE username = %s AND status = 'completed'
            """
            docs = fetch_all(docs_sql, (username,))
            document_ids = [d['doc_id'] for d in docs] if docs else []
        
        if not document_ids:
            return jsonify({
                'success': True,
                'data': {
                    'results': [],
                    'message': '没有可搜索的文档'
                }
            })
        
        # 执行搜索
        vector_store = get_vector_store()
        results = vector_store.search_similar(
            query=query,
            document_ids=document_ids,
            top_k=top_k,
            threshold=threshold
        )
        
        # 补充章节详细信息
        for r in results:
            if r.get('chapter_id'):
                chapter_sql = """
                    SELECT c.title, c.level, dpr.filename
                    FROM chapters c
                    JOIN doc_process_records dpr ON c.document_id = dpr.doc_id
                    WHERE c.id = %s
                """
                chapter = fetch_one(chapter_sql, (r['chapter_id'],))
                if chapter:
                    r['chapter_title'] = chapter['title']
                    r['chapter_level'] = chapter['level']
                    r['filename'] = chapter['filename']
        
        return jsonify({
            'success': True,
            'data': {
                'query': query,
                'results': results,
                'count': len(results)
            }
        })
        
    except Exception as e:
        logger.error(f"向量搜索失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding/stats', methods=['GET'])
def get_embedding_stats():
    """获取向量存储统计信息"""
    try:
        from .embedding_service import get_vector_store
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        doc_id = request.args.get('doc_id', type=int)
        
        vector_store = get_vector_store()
        stats = vector_store.get_embedding_stats(doc_id)
        
        return jsonify({
            'success': True,
            'data': stats
        })
        
    except Exception as e:
        logger.error(f"获取向量统计失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding/delete', methods=['POST'])
def delete_document_embeddings():
    """删除文档的向量数据"""
    try:
        from .embedding_service import get_vector_store
        
        user_info = session.get('user')
        username = user_info.get('username') if user_info else 'anonymous'
        
        data = request.json
        doc_id = data.get('doc_id')
        
        if not doc_id:
            return jsonify({
                'success': False,
                'error': '请提供文档ID'
            }), 400
        
        # 验证权限
        doc_sql = """
            SELECT doc_id FROM doc_process_records 
            WHERE doc_id = %s AND username = %s
        """
        doc = fetch_one(doc_sql, (doc_id, username))
        
        if not doc:
            return jsonify({
                'success': False,
                'error': '文档不存在或无权访问'
            }), 404
        
        vector_store = get_vector_store()
        deleted = vector_store.delete_document_embeddings(doc_id)
        
        return jsonify({
            'success': True,
            'data': {
                'doc_id': doc_id,
                'deleted_count': deleted
            }
        })
        
    except Exception as e:
        logger.error(f"删除向量数据失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


# ==================== Embedding 配置管理 API ====================

@chatdoc.route('/api/chat/embedding-configs', methods=['GET'])
def get_embedding_configs():
    """获取 Embedding 配置列表"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        configs = EmbeddingConfigManager.get_all_configs()
        
        return jsonify({
            'success': True,
            'data': configs
        })
        
    except Exception as e:
        logger.error(f"获取Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/<int:config_id>', methods=['GET'])
def get_embedding_config(config_id):
    """获取单个 Embedding 配置详情"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        config = EmbeddingConfigManager.get_config(config_id)
        
        if not config:
            return jsonify({
                'success': False,
                'error': '配置不存在'
            }), 404
        
        # 隐藏敏感信息
        if config.get('api_key'):
            config['api_key'] = config['api_key'][:8] + '****' if len(config['api_key']) > 8 else '****'
        
        return jsonify({
            'success': True,
            'data': config
        })
        
    except Exception as e:
        logger.error(f"获取Embedding配置详情失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config', methods=['POST'])
def create_embedding_config():
    """创建新的 Embedding 配置"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        data = request.json
        
        name = data.get('name', '').strip()
        provider = data.get('provider', '').strip()
        model_name = data.get('model_name', '').strip()
        dimensions = data.get('dimensions', 1536)
        api_key = data.get('api_key', '').strip() or None
        api_base = data.get('api_base', '').strip() or None
        is_default = data.get('is_default', False)
        extra_config = data.get('extra_config')
        
        if not name:
            return jsonify({'success': False, 'error': '请输入配置名称'}), 400
        if not provider:
            return jsonify({'success': False, 'error': '请选择提供商'}), 400
        if not model_name:
            return jsonify({'success': False, 'error': '请选择或输入模型名称'}), 400
        
        config_id = EmbeddingConfigManager.create_config(
            name=name,
            provider=provider,
            model_name=model_name,
            dimensions=dimensions,
            api_key=api_key,
            api_base=api_base,
            is_default=is_default,
            extra_config=extra_config
        )
        
        if config_id:
            return jsonify({
                'success': True,
                'data': {'id': config_id},
                'message': '配置创建成功'
            })
        else:
            return jsonify({
                'success': False,
                'error': '创建配置失败'
            }), 500
        
    except Exception as e:
        logger.error(f"创建Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/<int:config_id>', methods=['PUT'])
def update_embedding_config(config_id):
    """更新 Embedding 配置"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        data = request.json
        
        # 过滤可更新字段
        update_data = {}
        for field in ['name', 'provider', 'model_name', 'dimensions', 'api_key', 'api_base', 'is_default', 'extra_config']:
            if field in data:
                update_data[field] = data[field]
        
        success = EmbeddingConfigManager.update_config(config_id, **update_data)
        
        if success:
            return jsonify({
                'success': True,
                'message': '配置更新成功'
            })
        else:
            return jsonify({
                'success': False,
                'error': '更新失败，配置可能不存在'
            }), 404
        
    except Exception as e:
        logger.error(f"更新Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/<int:config_id>', methods=['DELETE'])
def delete_embedding_config(config_id):
    """删除 Embedding 配置"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        success = EmbeddingConfigManager.delete_config(config_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': '配置删除成功'
            })
        else:
            return jsonify({
                'success': False,
                'error': '删除失败，配置可能不存在'
            }), 404
        
    except Exception as e:
        logger.error(f"删除Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/<int:config_id>/set-default', methods=['POST'])
def set_default_embedding_config(config_id):
    """设置默认 Embedding 配置"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        success = EmbeddingConfigManager.set_default(config_id)
        
        if success:
            return jsonify({
                'success': True,
                'message': '已设为默认配置'
            })
        else:
            return jsonify({
                'success': False,
                'error': '设置失败，配置可能不存在'
            }), 404
        
    except Exception as e:
        logger.error(f"设置默认Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/test', methods=['POST'])
def test_embedding_config():
    """
    测试 Embedding 配置是否可用
    
    请求参数:
        config_id: 配置ID（测试已保存的配置）
        或
        provider, model_name, api_key, dimensions 等（测试新配置）
    """
    try:
        from .embedding_service import EmbeddingConfigManager
        
        data = request.json
        config_id = data.get('config_id')
        
        if config_id:
            # 测试已保存的配置
            result = EmbeddingConfigManager.test_config(config_id=config_id)
        else:
            # 测试传入的配置数据
            config_data = {
                'provider': data.get('provider'),
                'model_name': data.get('model_name'),
                'api_key': data.get('api_key'),
                'api_base': data.get('api_base'),
                'dimensions': data.get('dimensions', 1536),
                'extra_config': data.get('extra_config')
            }
            
            if not config_data['provider']:
                return jsonify({'success': False, 'error': '请选择提供商'}), 400
            if not config_data['model_name']:
                return jsonify({'success': False, 'error': '请选择模型'}), 400
            
            result = EmbeddingConfigManager.test_config(config_data=config_data)
        
        return jsonify(result)
        
    except Exception as e:
        logger.error(f"测试Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-providers', methods=['GET'])
def get_embedding_providers():
    """获取支持的 Embedding 提供商列表"""
    try:
        from .embedding_service import EmbeddingConfigManager
        
        providers = EmbeddingConfigManager.get_supported_providers()
        
        return jsonify({
            'success': True,
            'data': providers
        })
        
    except Exception as e:
        logger.error(f"获取Embedding提供商失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500


@chatdoc.route('/api/chat/embedding-config/check', methods=['GET'])
def check_embedding_config():
    """检查是否有可用的 Embedding 配置"""
    try:
        from .embedding_service import EmbeddingConfigManager, EmbeddingService
        
        config = EmbeddingConfigManager.get_default_config()
        
        # 尝试初始化服务
        service = EmbeddingService()
        has_provider = service.provider is not None
        
        return jsonify({
            'success': True,
            'data': {
                'has_config': config is not None,
                'has_working_provider': has_provider,
                'config_name': config.get('name') if config else None,
                'provider': config.get('provider') if config else None,
                'model_name': config.get('model_name') if config else None,
                'using_fallback': not has_provider,
                'fallback_message': '未配置或配置无效，使用简单词向量方案（效果有限）' if not has_provider else None
            }
        })
        
    except Exception as e:
        logger.error(f"检查Embedding配置失败: {e}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500
